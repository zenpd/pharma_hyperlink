"""Word (.docx) hyperlink injection.

Per ADR-0001, this module manipulates OOXML directly via python-docx. It
never mutates the input file — every call to `save()` writes a new copy
to the configured output path.

Design notes:
- Hyperlinks live in the `w:hyperlink` element which wraps one or more
  `w:r` (run) elements. To wrap part of an existing run, we split the
  run at the link boundary first, then wrap only the targeted span.
- External URLs use a relationship in the paragraph's part with a
  generated `r:id`. Internal anchors use `w:anchor` and require a
  pre-existing `w:bookmarkStart`/`w:bookmarkEnd` pair in the document.
- A "Hyperlink" character style is applied by reference (`w:rStyle`) so
  Word renders blue underlined text; the original run's other styling
  (font, size, color overrides) is preserved on the inner run.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from hyperlink_engine.config.logging_setup import get_logger
from hyperlink_engine.models import HyperlinkSpec, LinkKind, RunLocation

log = get_logger(__name__)


class DocxInjectionError(RuntimeError):
    """Raised when a hyperlink cannot be injected safely."""


class DocxLinker:
    """High-level hyperlink injector for .docx files.

    Workflow:
        linker = DocxLinker(source_path, output_path)
        linker.add_external_link(location, "https://example.com", "click here")
        linker.add_internal_link(location, "sec_2_5_3")
        linker.save()
    """

    def __init__(self, source_path: Path, output_path: Path) -> None:
        if not source_path.exists():
            raise FileNotFoundError(source_path)
        self.source_path = source_path
        self.output_path = output_path
        self._document = Document(str(source_path))
        self._pending: list[HyperlinkSpec] = []
        self._pending_bookmarks: list[tuple[RunLocation, str]] = []
        self._bookmark_names: set[str] = set()
        self._next_bookmark_id = 1000  # avoid collisions with any inherited IDs

    # ── Public API ──────────────────────────────────────────────────────

    def add_external_link(
        self,
        location: RunLocation,
        url: str,
        display_text: str | None = None,
    ) -> None:
        self._pending.append(
            HyperlinkSpec(
                location=location,
                kind=LinkKind.EXTERNAL_URL,
                target=url,
                display_text=display_text,
            )
        )

    def add_internal_link(
        self,
        location: RunLocation,
        anchor: str,
        display_text: str | None = None,
    ) -> None:
        self._pending.append(
            HyperlinkSpec(
                location=location,
                kind=LinkKind.INTERNAL_BOOKMARK,
                target=anchor,
                display_text=display_text,
            )
        )

    def add_link(self, spec: HyperlinkSpec) -> None:
        self._pending.append(spec)

    def add_many(self, specs: Iterable[HyperlinkSpec]) -> None:
        self._pending.extend(specs)

    def add_bookmark(self, location: RunLocation, name: str) -> None:
        """Queue a w:bookmarkStart/End pair around the given run span.

        Bookmark names must be unique within a document. The linker keeps
        a registry and silently no-ops duplicate names — first declaration
        wins.
        """
        if name in self._bookmark_names:
            return
        self._bookmark_names.add(name)
        self._pending_bookmarks.append((location, name))

    @property
    def pending_count(self) -> int:
        return len(self._pending) + len(self._pending_bookmarks)

    def save(self) -> Path:
        """Apply every queued hyperlink and write the output file."""
        # Apply in reverse order within each paragraph so character offsets
        # do not shift under us.
        by_para: dict[int, list[HyperlinkSpec]] = {}
        for spec in self._pending:
            by_para.setdefault(spec.location.paragraph_index, []).append(spec)

        injected = 0
        for para_index, specs in by_para.items():
            # Sort each paragraph's specs by run_index DESC, then char_start DESC
            specs.sort(
                key=lambda s: (s.location.run_index, s.location.char_start),
                reverse=True,
            )
            for spec in specs:
                self._inject_one(para_index, spec)
                injected += 1

        bookmarks_added = 0
        for location, name in self._pending_bookmarks:
            self._inject_bookmark(location, name)
            bookmarks_added += 1

        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self._document.save(str(self.output_path))
        log.info(
            "docx_linker_saved",
            source=str(self.source_path),
            output=str(self.output_path),
            links_injected=injected,
            bookmarks_added=bookmarks_added,
        )
        return self.output_path

    def _inject_bookmark(self, location: RunLocation, name: str) -> None:
        """Insert a paragraph-level w:bookmarkStart/End pair.

        We anchor the bookmark at the head of the paragraph rather than
        splitting runs at the precise character offset — bookmark scope
        only needs to *contain* the target so anchor resolution works,
        and paragraph-level anchors round-trip cleanly through Word.
        """
        if location.paragraph_index >= len(self._document.paragraphs):
            raise DocxInjectionError(
                f"Cannot place bookmark {name!r}: paragraph "
                f"{location.paragraph_index} out of range"
            )
        paragraph = self._document.paragraphs[location.paragraph_index]
        bm_id = str(self._next_bookmark_id)
        self._next_bookmark_id += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bm_id)
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bm_id)
        # Insert at the start of the paragraph; place end at the tail.
        paragraph._p.insert(0, start)
        paragraph._p.append(end)

    # ── Internals ───────────────────────────────────────────────────────

    def _inject_one(self, para_index: int, spec: HyperlinkSpec) -> None:
        if para_index >= len(self._document.paragraphs):
            raise DocxInjectionError(
                f"Paragraph index {para_index} out of range "
                f"(document has {len(self._document.paragraphs)} paragraphs)"
            )
        paragraph = self._document.paragraphs[para_index]
        if spec.location.run_index >= len(paragraph.runs):
            raise DocxInjectionError(
                f"Run index {spec.location.run_index} out of range in paragraph "
                f"{para_index} (has {len(paragraph.runs)} runs)"
            )

        run = paragraph.runs[spec.location.run_index]
        run_text = run.text
        start = spec.location.char_start
        end = spec.location.char_end

        if end > len(run_text):
            raise DocxInjectionError(
                f"Span [{start},{end}) exceeds run length {len(run_text)} in "
                f"paragraph {para_index}, run {spec.location.run_index}"
            )

        link_text = spec.display_text or run_text[start:end]

        # Split the run into up to three parts: before / link / after.
        # We keep the original run as "before" (preserving its styling) and
        # build sibling runs/hyperlinks for the rest.
        before_text = run_text[:start]
        after_text = run_text[end:]

        link_run = self._build_styled_run(run, link_text)
        hyperlink_elem = self._build_hyperlink_element(
            paragraph=paragraph,
            spec=spec,
            inner_run_element=link_run,
        )

        # Mutate the original run to hold only the "before" portion.
        run.text = before_text

        # Insert the hyperlink immediately after the original run.
        run._element.addnext(hyperlink_elem)  # type: ignore[attr-defined]

        if after_text:
            after_run = self._build_styled_run(run, after_text)
            hyperlink_elem.addnext(after_run)

    def _build_styled_run(self, template_run, text: str):  # type: ignore[no-untyped-def]
        """Build a new w:r element with text and a copy of the template run's properties."""
        new_run = OxmlElement("w:r")

        # Deep-copy run properties (rPr) so font, color, etc., are preserved.
        rpr = template_run._element.find(qn("w:rPr"))  # type: ignore[attr-defined]
        if rpr is not None:
            new_run.append(deepcopy(rpr))

        t = OxmlElement("w:t")
        t.text = text
        # Preserve leading/trailing whitespace (xml:space="preserve")
        if text != text.strip():
            t.set(qn("xml:space"), "preserve")
        new_run.append(t)
        return new_run

    def _build_hyperlink_element(
        self,
        paragraph,  # type: ignore[no-untyped-def]
        spec: HyperlinkSpec,
        inner_run_element,  # type: ignore[no-untyped-def]
    ):  # type: ignore[no-untyped-def]
        hyperlink = OxmlElement("w:hyperlink")

        if spec.kind is LinkKind.EXTERNAL_URL:
            rel_id = paragraph.part.relate_to(spec.target, RT.HYPERLINK, is_external=True)
            hyperlink.set(qn("r:id"), rel_id)
        elif spec.kind is LinkKind.INTERNAL_BOOKMARK:
            hyperlink.set(qn("w:anchor"), spec.target)
        else:
            # Cross-doc / cross-module: treat target as a file relationship + anchor
            # combination; for Phase 1 we wire the bare URL to keep the API single-path.
            rel_id = paragraph.part.relate_to(spec.target, RT.HYPERLINK, is_external=True)
            hyperlink.set(qn("r:id"), rel_id)

        # Apply the "Hyperlink" character style to the inner run so it renders
        # blue underlined.
        self._apply_hyperlink_style(inner_run_element)
        hyperlink.append(inner_run_element)
        return hyperlink

    @staticmethod
    def _apply_hyperlink_style(run_element) -> None:  # type: ignore[no-untyped-def]
        """Inject `<w:rStyle w:val="Hyperlink"/>` into the run's rPr."""
        rpr = run_element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_element.insert(0, rpr)

        # Don't duplicate if already present
        existing = rpr.find(qn("w:rStyle"))
        if existing is not None:
            existing.set(qn("w:val"), "Hyperlink")
            return

        style_elem = OxmlElement("w:rStyle")
        style_elem.set(qn("w:val"), "Hyperlink")
        rpr.insert(0, style_elem)
