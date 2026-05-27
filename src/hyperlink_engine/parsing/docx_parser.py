"""Layer 2 — Word (.docx) parsing.

Walks paragraphs + table cells in document order, captures run-level styling
(bold/italic/underline/font/color/style), preserves the (paragraph_index,
run_index, char_offset) anchors the injection layer needs, and surfaces
candidate "blue text without hyperlink" runs that the anomaly layer will
review later.

The parser does **not** modify the document — it just observes.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx.document import Document as DocxObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from hyperlink_engine.config.logging_setup import get_logger
from hyperlink_engine.config.settings import get_settings
from hyperlink_engine.ingestion.docx_loader import load_docx
from hyperlink_engine.models import (
    DocxDocument,
    DocxParagraph,
    DocxRun,
    DocumentProvenance,
    RunStyle,
)

_log = get_logger("parsing.docx")


# A "blue" color is loosely anything whose B channel dominates R and G.
# We tolerate ±tolerance points around pure blue (configurable in settings).
_BLUE_HEX_RE = re.compile(r"^[0-9A-Fa-f]{6}$")


def _is_blueish(hex_rgb: str | None, tolerance: int) -> bool:
    if not hex_rgb or not _BLUE_HEX_RE.fullmatch(hex_rgb):
        return False
    r = int(hex_rgb[0:2], 16)
    g = int(hex_rgb[2:4], 16)
    b = int(hex_rgb[4:6], 16)
    return b >= 128 and b - max(r, g) >= tolerance


def _extract_run_style(run: Run) -> RunStyle:
    rgb = None
    try:
        color = run.font.color
        if color is not None and color.rgb is not None:
            rgb = str(color.rgb).upper()
    except (AttributeError, ValueError):
        rgb = None

    size_pt: float | None
    try:
        size_pt = float(run.font.size.pt) if run.font.size is not None else None
    except (AttributeError, ValueError):
        size_pt = None

    style_name = None
    try:
        if run.style is not None:
            style_name = run.style.name
    except AttributeError:
        style_name = None

    return RunStyle(
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        font_name=run.font.name,
        font_size_pt=size_pt,
        color_rgb=rgb,
        style_name=style_name,
        is_hyperlink=False,  # filled in by paragraph walker if the run lives inside a w:hyperlink
    )


def _iter_block_paragraphs(document: DocxObject) -> Iterable[tuple[Paragraph, bool, tuple[int, int, int] | None]]:
    """Yield (paragraph, in_table, (table_idx, row, col)) in document order.

    python-docx exposes top-level paragraphs and tables but does not interleave
    them in document order out of the box. We approximate document order by
    iterating top-level paragraphs first, then table contents — close enough
    for Phase 1 since downstream code uses our `paragraph_index` numbering
    as the canonical identifier.
    """
    for para in document.paragraphs:
        yield para, False, None

    for t_idx, table in enumerate(document.tables):
        assert isinstance(table, Table)
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    yield para, True, (t_idx, r_idx, c_idx)


def _paragraph_hyperlink_run_ids(para: Paragraph) -> set[int]:
    """Return the set of object-ids of run XML elements that live inside a w:hyperlink.

    python-docx flattens runs into `paragraph.runs` and silently drops the
    surrounding `w:hyperlink` wrapper. We re-read the underlying XML to flag
    runs whose parent is a hyperlink so the parser can mark them.
    """
    flagged: set[int] = set()
    for hyperlink in para._p.findall(qn("w:hyperlink")):
        for r in hyperlink.findall(qn("w:r")):
            flagged.add(id(r))
    return flagged


def _collect_existing_hyperlinks(document: DocxObject) -> list[str]:
    """Return every external hyperlink target stored in the package relationships."""
    targets: list[str] = []
    rels = document.part.rels
    for rel in rels.values():
        if "hyperlink" in rel.reltype and rel.target_ref:
            targets.append(rel.target_ref)
    return targets


def parse_docx_document(document: DocxObject, provenance: DocumentProvenance) -> DocxDocument:
    """Project a python-docx Document into our typed DocxDocument shape."""
    settings = get_settings()
    blue_tolerance = settings.blue_text_rgb_tolerance

    core = document.core_properties
    title = core.title or None
    author = core.author or None

    parsed_paragraphs: list[DocxParagraph] = []

    for p_idx, (para, in_table, coords) in enumerate(_iter_block_paragraphs(document)):
        link_run_ids = _paragraph_hyperlink_run_ids(para)
        runs_out: list[DocxRun] = []
        char_cursor = 0
        for r_idx, run in enumerate(para.runs):
            style = _extract_run_style(run)
            if id(run._r) in link_run_ids:
                style = style.model_copy(update={"is_hyperlink": True})
            # Override the blue-ish detection into a denormalized convenience flag
            # via metadata is overkill; downstream callers can re-check via style.
            run_out = DocxRun(
                run_index=r_idx,
                text=run.text,
                style=style,
                char_offset_in_paragraph=char_cursor,
            )
            runs_out.append(run_out)
            char_cursor += len(run.text)

        para_text = "".join(r.text for r in runs_out)
        style_name = para.style.name if para.style is not None else None
        parsed_paragraphs.append(
            DocxParagraph(
                paragraph_index=p_idx,
                style_name=style_name,
                text=para_text,
                runs=runs_out,
                in_table=in_table,
                table_coords=coords,
            )
        )

    existing = _collect_existing_hyperlinks(document)
    parsed = DocxDocument(
        provenance=provenance,
        title=title,
        author=author,
        paragraph_count=len(parsed_paragraphs),
        paragraphs=parsed_paragraphs,
        existing_hyperlinks=existing,
    )

    _log.info(
        "docx_parsed",
        path=str(provenance.source_path),
        paragraphs=parsed.paragraph_count,
        runs=parsed.total_runs,
        chars=parsed.total_chars,
        existing_links=len(existing),
        blue_tolerance=blue_tolerance,
    )
    return parsed


def parse_docx(path: Path) -> DocxDocument:
    """Convenience: load + parse in one call."""
    document, provenance = load_docx(Path(path))
    return parse_docx_document(document, provenance)


def candidate_blue_runs(parsed: DocxDocument) -> list[tuple[int, int, DocxRun]]:
    """Return runs whose color is blue-ish AND which are NOT already a hyperlink.

    These are anomaly candidates: the human author colored the text blue (a
    visual cue that it should be a hyperlink) but no link is actually attached.
    """
    settings = get_settings()
    tolerance = settings.blue_text_rgb_tolerance
    out: list[tuple[int, int, DocxRun]] = []
    for para in parsed.paragraphs:
        for run in para.runs:
            if run.style.is_hyperlink:
                continue
            if _is_blueish(run.style.color_rgb, tolerance):
                out.append((para.paragraph_index, run.run_index, run))
    return out
