"""Per-document **Anchor Index** — PLAN TEN, the core of precise reference targeting.

Problem it solves
-----------------
The injector used to declare a bookmark / named-destination at the *first textual
mention* of a reference (e.g. the first sentence that says "…see Table 14.2.1.1…").
Every link to "Table 14.2.1.1" then jumped to that citation instead of the actual
table. With many tables / sections / appendices, links collapsed onto the wrong
spot ("redirect to a reference, not to the data").

This module builds a map ``canonical_key -> definition_location`` so injection can
anchor each link at the **definition** (the caption / heading), not the citation:

    "table_ref_14_2_1_1"  -> (page 120, bbox)          # PDF
    "section_ref_6_3"     -> (paragraph 842)           # DOCX

Design
------
* **Captions** (Table / Figure / Listing / Appendix) are *already detected* — the
  caption line "Table 14.2.1.1: Demographics" matches the same regex a citation
  does. We distinguish a **definition** from a **citation** by position + shape
  (``_is_caption_definition``), so the index is built straight from the detection
  records with **zero extra parsing**.
* **Sections** are headings, which the citation regex does *not* fire on (no cue
  word), so we read them from the document's own structure: ``doc.get_toc()`` for
  PDF, ``Heading`` paragraph styles for DOCX. Both are cheap.
* **First definition wins** per key (PLAN TEN Step 5) — captions/headings are rare
  and citations are excluded, so the first hit is the real target.

PLAN TEN-bis — structure scan (this revision)
---------------------------------------------
Detection runs **per-span** (PDF) / **per-run** (DOCX), so a caption's ``context`` is
only a fragment, and the citation regexes are case-sensitive — all-caps headings
("APPENDIX A.") are never even detected. So after the detection-driven pass we
additionally **scan the document's own structure** (joined-span *lines* for PDF,
full-paragraph *text* for DOCX) for caption-shaped definitions: case-tolerant,
Table-of-Contents excluded, length-guarded. The scan is *additive* — it only fills
keys not already found, so it can never relocate a working anchor — and PDF/DOCX are
kept in parity. (Captions that live inside DOCX table cells are out of reach until the
injector's paragraph-index addressing learns table cells — a separate effort.)

Everything is **best-effort**: any failure degrades to "no entry for that key", and
the caller falls back to today's behavior (anchor at the citation). It can never
break injection.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from hyperlink_engine.config.logging_setup import get_logger

_log = get_logger("injection.anchor_index")

# Labels whose *definition* is a caption line in the body text.
_CAPTION_LABELS = {"TABLE_REF", "FIGURE_REF", "LISTING_REF", "APPENDIX_REF"}
_LABEL_KEYWORD = {
    "TABLE_REF": "Table",
    "FIGURE_REF": "Figure",
    "LISTING_REF": "Listing",
    "APPENDIX_REF": "Appendix",
}

# Per-label number shape, kept byte-compatible with the detection regexes so a
# scanned caption produces the *same* canonical key as its citations. Appendix
# allows a lone letter ("Appendix A") besides a dotted number ("Appendix 16.1").
_NUMPAT = {
    "TABLE_REF": r"\d+(?:[.\-]\d+){0,4}",
    "FIGURE_REF": r"\d+(?:[.\-]\d+){0,4}",
    "LISTING_REF": r"\d+(?:[.\-]\d+){0,4}",
    "APPENDIX_REF": r"\d+(?:\.\d+){0,3}|[A-Za-z]",
}

# A definition caption starts with "<Keyword> <num>" and is *followed by* a title
# (punctuation, a capitalised word, a tab) or nothing — never a lower-case verb
# ("Table 1 shows ...") which would be a sentence/citation, not the caption.
_TITLE_SEP = set(":.-—\t")
# Captions/headings are short; a long line that merely *starts* with "Table N Word"
# is running prose, not a caption — only trust the Title-Case branch under this
# length so the scanner doesn't anchor at a sentence.
_MAX_CAPTION_LEN = 200

# Table-of-Contents / List-of-Tables entries repeat the caption text but point at
# a page through a run of dot leaders ("Table 5: Demographics …… 42"). They must
# never be chosen as the definition, or links land on the ToC page (PLAN TEN-bis).
_TOC_LEADER = re.compile(r"\.{3,}|(?:\.\s){3,}|…|(?:·\s*){3,}")


def canonical_anchor_key(label: str, num: str) -> str:
    """Canonical bookmark/named-destination key for a numbered reference.

    Must stay byte-identical to ``workers.tasks._resolve_target`` so a citation
    and its definition resolve to the *same* key.  e.g. ("TABLE_REF","14.2.1.1")
    -> "table_ref_14_2_1_1"; ("APPENDIX_REF","B") -> "appendix_ref_B".
    """
    slug = num.replace(".", "_").replace("-", "_")
    return f"{label.lower()}_{slug}"


def canonical_visit_key(unit: str, n: str) -> str:
    """Canonical key for a VISIT_REF ('Week 2 Visit' -> 'visit_ref_week_2').

    Must stay byte-identical to ``workers.tasks._resolve_target`` so a visit
    citation and the visit section heading that defines it resolve to the same
    key. The target section is found by matching the visit phrase against the
    document's section headings (see ``_add_*_visit_anchors``).
    """
    return f"visit_ref_{unit.lower()}_{n}"


# A visit/timepoint phrase inside a section heading: "9.3.3.1 Week 2 (± 3 Days)",
# "9.3.3.6 Month 3 (± 3 Days)". Maps the heading's visit phrase to the canonical
# visit key so a "Week 2 Visit" citation anchors at the visit section. Day-N is
# matched too, though real protocols often title Day 1 as "Study Randomization",
# so a Day reference may not resolve (it then falls back to the citation anchor).
_VISIT_IN_HEADING = re.compile(r"\b(?P<unit>Week|Day|Month)\s+(?P<n>\d+)\b", re.IGNORECASE)


def _is_toc_line(s: str) -> bool:
    """True for a Table-of-Contents / List-of-* entry (dot-leader to a page no.)."""
    return bool(_TOC_LEADER.search(s))


def _caption_def_num(context: str, label: str, *, require_title: bool = False) -> str | None:
    """Return the reference number when *context* is the caption/heading that
    **defines** the reference, else None.

    Case-tolerant (catches all-caps headings like ``APPENDIX A.``), ToC-excluded,
    and length-guarded so prose that merely opens with "Table N …" is not mistaken
    for a caption. The number is normalised to match ``canonical_anchor_key``
    exactly (a lone appendix letter is upper-cased, as the citation regex emits).

    ``require_title=True`` rejects a **bare** "Appendix D" / "Table 5" with no title
    after it. A bare token is ambiguous — PDF span-splitting turns a *citation*
    ("…contemplated in Appendix D of this protocol.") into a lone "Appendix D" span
    that is indistinguishable from a heading by text alone. Callers that only have a
    fragment (the per-span/per-run detection pass) pass ``require_title=True`` and
    leave bare headings to the structure scan, which can font-gate them.
    """
    kw = _LABEL_KEYWORD.get(label)
    numpat = _NUMPAT.get(label)
    if not kw or not numpat:
        return None
    s = (context or "").strip()
    if not s or _is_toc_line(s):
        return None
    m = re.match(rf"{kw}\s+(?P<num>{numpat})\b", s, re.IGNORECASE)
    if not m:
        return None  # doesn't START with the caption → mid-sentence (citation)
    rest = s[m.end():].lstrip()
    if rest == "":
        ok = not require_title  # bare "Appendix D" — only a definition if caller allows it
    elif rest[0] in _TITLE_SEP:
        # A real titled caption has TEXT after the separator ("Table 5: Demographics").
        # A citation that merely ends a sentence ("…refer Appendix 2.") has only the
        # period — reject it when a title is required (the step-1 detection pass, whose
        # context is a lone run/span), so it can't be mistaken for the definition. A
        # genuine bare heading ("APPENDIX A.") is recovered by the structure scan,
        # which calls with require_title=False and sees the whole line.
        ok = bool(rest[1:].strip()) or not require_title
    else:
        # "Table 14.2.1.1 Summary of …" (Title-Case title) is a caption;
        # "Table 1 shows …" (lower-case verb) is a sentence. Long Title-Case lines
        # are prose, not captions.
        ok = rest[:1].isupper() and len(s) <= _MAX_CAPTION_LEN
    if not ok:
        return None
    num = m.group("num")
    if label == "APPENDIX_REF" and len(num) == 1 and num.isalpha():
        num = num.upper()
    return num


def _is_caption_definition(context: str, label: str) -> bool:
    """True when *context* is the caption that **defines** the reference, not a
    sentence that merely **cites** it."""
    return _caption_def_num(context, label) is not None


def _section_num(text: str) -> str | None:
    """Extract the leading dotted section number from a heading / ToC title.

    "6.3 Statistical Methods" -> "6.3"; "Section 6.3 …" -> "6.3"; "6 SAFETY" -> "6".
    """
    m = re.search(r"\b(\d+(?:\.\d+){0,4})\b", text or "")
    return m.group(1) if m else None


# A section *heading* in body text: a bounded dotted number then a Title-Case /
# ALL-CAPS title on a short standalone line — distinct from a numbered footnote or
# inclusion-criterion sentence ("12 Videotaping … optional."). Shape only; the
# font/style gate (larger font for PDF, bold / Heading style for DOCX) is applied
# by the scanners, because shape alone is NOT a safe discriminator in clinical PDFs
# (they are dense with short, bold, numbered footnotes). PLAN TEN-bis.
_SECTION_HEAD = re.compile(r"^(?P<num>\d{1,2}(?:\.\d{1,3}){0,3})\.?\s+(?P<title>[A-Za-z].*)$")
_MAX_SECTION_HEADING_LEN = 80


def _section_heading_num(text: str) -> str | None:
    """Return the section number when *text* is a numbered heading line, else None.

    Shape gate only (ToC-excluded, bounded number, short, Title-Case title, not a
    prose sentence). Callers MUST additionally require a typographic heading signal
    (larger font / bold / Heading style) — see the scanners — or footnotes and
    numbered criteria will be mistaken for headings.
    """
    s = (text or "").strip()
    if not s or _is_toc_line(s) or len(s) > _MAX_SECTION_HEADING_LEN:
        return None
    m = _SECTION_HEAD.match(s)
    if not m:
        return None
    title = m.group("title").strip()
    if not title[:1].isupper():
        return None  # lower-case start → prose, not a heading
    # A trailing sentence terminator marks prose ("… optional."); headings rarely
    # end in punctuation (an ALL-CAPS title like "6 SAFETY." is the exception).
    if s[-1] in ".?!;," and not title.isupper():
        return None
    return m.group("num")


def build_anchor_index(
    detections: list[dict[str, Any]],
    source_path: str | Path,
    *,
    is_pdf: bool,
) -> dict[str, dict[str, Any]]:
    """Build ``canonical_key -> location`` for one document.

    location dicts:
      * PDF  -> ``{"page_index": int, "bbox": [x0,y0,x1,y1]}``
      * DOCX -> ``{"paragraph_index": int}``
    """
    index: dict[str, dict[str, Any]] = {}

    # 1) Captions (Table/Figure/Listing/Appendix) straight from the detections.
    #    ``require_title=True`` — a detection's context is a single span/run, so a
    #    bare "Appendix D" fragment from a *citation* ("…contemplated in Appendix D
    #    of this protocol.") must NOT be trusted as the definition. Title-bearing
    #    captions ("Table 1: Demographics") still qualify; bare headings are left to
    #    the font-gated structure scan below.
    for det in detections:
        label = det.get("label")
        if label not in _CAPTION_LABELS:
            continue
        if _caption_def_num(det.get("context", ""), label, require_title=True) is None:
            continue
        num = (det.get("groups") or {}).get("num") or det.get("text", "")
        if not num:
            continue
        key = canonical_anchor_key(label, num)
        if key in index:
            continue  # first definition wins
        if is_pdf:
            if det.get("page_index") is None:
                continue
            index[key] = {"page_index": det["page_index"], "bbox": det.get("bbox")}
        else:
            if det.get("paragraph_index") is None:
                continue
            index[key] = {"paragraph_index": det["paragraph_index"]}

    # 2) Captions the per-span (PDF) / per-run (DOCX) detector missed — scan the
    #    document's OWN structure (joined-span lines / full-paragraph text) for
    #    caption-shaped definitions. Recovers all-caps headings ("APPENDIX A."),
    #    captions fragmented across spans/runs, and number-alone caption lines that
    #    the case-sensitive citation regex never fired on. ToC lines are excluded.
    #    Best-effort and *additive*: only fills keys not already found above, so it
    #    can never relocate a working anchor. PDF and DOCX are kept in parity here.
    try:
        if is_pdf:
            _scan_pdf_captions(index, source_path)
        else:
            _scan_docx_captions(index, source_path)
    except Exception as exc:  # noqa: BLE001 — never break injection
        _log.warning("anchor_index_caption_scan_failed", source=str(source_path), error=str(exc))

    # 3) Sections from the document's own structure (best-effort, high precision):
    #    embedded ToC / bookmarks (PDF) and Heading paragraph styles (DOCX).
    try:
        if is_pdf:
            _add_pdf_sections(index, source_path)
        else:
            _add_docx_sections(index, source_path)
    except Exception as exc:  # noqa: BLE001 — never break injection
        _log.warning("anchor_index_sections_failed", source=str(source_path), error=str(exc))

    # 4) Section headings the ToC/Heading-style sources missed — scan the body for
    #    numbered headings, gated on a typographic signal (larger font for PDF,
    #    bold / Heading style for DOCX). Recovers sections in PDFs exported without
    #    bookmarks and Word docs with manually-formatted headings. Additive and
    #    fill-only; the font/style gate keeps numbered footnotes & criteria out, so
    #    where headings carry no signal it adds nothing rather than guessing wrong.
    try:
        if is_pdf:
            _scan_pdf_section_headings(index, source_path)
        else:
            _scan_docx_section_headings(index, source_path)
    except Exception as exc:  # noqa: BLE001 — never break injection
        _log.warning("anchor_index_section_scan_failed", source=str(source_path), error=str(exc))

    # 5) Visit/timepoint anchors (clinical protocols): map "Week 2 Visit" → the
    #    "9.3.3.1 Week 2" section heading via the visit phrase in the heading text.
    #    Additive and best-effort; an unresolved visit falls back to the citation
    #    anchor, so it can never relocate a working anchor or break injection.
    try:
        if is_pdf:
            _add_pdf_visit_anchors(index, source_path)
        else:
            _add_docx_visit_anchors(index, source_path)
    except Exception as exc:  # noqa: BLE001 — never break injection
        _log.warning("anchor_index_visit_failed", source=str(source_path), error=str(exc))

    # 6) References-section entries: parse the "REFERENCES" list into a per-entry
    #    bookmark so an in-text citation ("Helget LN, 2024", "[7]") anchors at its
    #    bibliography entry, plus a ``ref_heading`` fallback so a citation whose
    #    specific entry can't be found still lands on the References heading rather
    #    than being dropped (the team's "by not skipping" requirement). Additive
    #    and best-effort.
    try:
        if is_pdf:
            _add_pdf_references(index, source_path)
        else:
            _add_docx_references(index, source_path)
    except Exception as exc:  # noqa: BLE001 — never break injection
        _log.warning("anchor_index_references_failed", source=str(source_path), error=str(exc))

    return index


def _union_bbox(boxes: list[Any]) -> list[float] | None:
    """Union of PyMuPDF span bboxes (x0,y0,x1,y1); the caption line's rectangle."""
    boxes = [b for b in boxes if b]
    if not boxes:
        return None
    return [
        min(b[0] for b in boxes),
        min(b[1] for b in boxes),
        max(b[2] for b in boxes),
        max(b[3] for b in boxes),
    ]


def _scan_pdf_captions(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Fill caption anchors by scanning the PDF's own text lines (joined spans), so
    a caption split across spans reassembles and all-caps headings are seen.

    Title-bearing captions ("Table 5: Demographics", "APPENDIX A.") are trusted on
    text alone. A **bare** "Appendix D" is captured only when it is typographically
    a heading (bold or larger than body) — otherwise it is a citation fragment or a
    List-of-Appendices entry and must not win (the real heading is the bold one).
    """
    import collections

    import fitz  # PyMuPDF

    doc = fitz.open(str(source_path))
    try:
        sizes: collections.Counter = collections.Counter()
        lines: list[tuple[int, str, bool, float, list[float] | None]] = []
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type", 0) != 0:  # keep text blocks only
                    continue
                for line in block.get("lines", []):
                    spans = [s for s in line.get("spans", []) if s.get("text", "").strip()]
                    if not spans:
                        continue
                    text = "".join(s.get("text", "") for s in spans).strip()
                    bold = any(int(s.get("flags", 0)) & 16 for s in spans)
                    max_size = max(round(float(s.get("size", 0.0)), 1) for s in spans)
                    for s in spans:
                        sizes[round(float(s.get("size", 0.0)), 1)] += len(s.get("text", ""))
                    lines.append(
                        (page_index, text, bold, max_size,
                         _union_bbox([s.get("bbox") for s in spans]))
                    )
        body_size = sizes.most_common(1)[0][0] if sizes else 0.0
        for page_index, text, bold, max_size, bbox in lines:
            heading_signal = bold or max_size >= body_size + 1.5
            for label in _CAPTION_LABELS:
                num = _caption_def_num(text, label, require_title=True)
                if num is None and heading_signal:
                    num = _caption_def_num(text, label, require_title=False)  # bare heading
                if not num:
                    continue
                key = canonical_anchor_key(label, num)
                if key in index:
                    continue  # first definition wins
                index[key] = {"page_index": page_index, "bbox": bbox}
    finally:
        doc.close()


def _scan_docx_captions(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Fill caption anchors by scanning the DOCX's own paragraph text (full text, so
    a caption split across runs reassembles). A bare "Appendix D" paragraph counts
    only when it is a Heading style or bold — else it is a citation, not a heading."""
    from docx import Document

    doc = Document(str(source_path))
    for p_idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        style = (getattr(para.style, "name", "") or "").lower()
        heading_signal = (
            style.startswith("heading")
            or style.startswith("title")
            or (bool(para.runs) and any(bool(r.bold) for r in para.runs))
        )
        for label in _CAPTION_LABELS:
            num = _caption_def_num(text, label, require_title=True)
            if num is None and heading_signal:
                num = _caption_def_num(text, label, require_title=False)  # bare heading
            if not num:
                continue
            key = canonical_anchor_key(label, num)
            if key in index:
                continue  # first definition wins
            index[key] = {"paragraph_index": p_idx}


def _add_pdf_sections(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Add section anchors from the PDF's own table-of-contents / bookmarks."""
    import fitz  # PyMuPDF

    doc = fitz.open(str(source_path))
    try:
        toc = doc.get_toc(simple=True)  # [[level, title, page_1based], ...]
    finally:
        doc.close()
    for entry in toc:
        try:
            _level, title, page = entry[0], entry[1], entry[2]
        except (IndexError, TypeError):
            continue
        num = _section_num(title)
        if not num or page is None or page < 1:
            continue
        key = canonical_anchor_key("SECTION_REF", num)
        if key in index:
            continue
        index[key] = {"page_index": int(page) - 1, "bbox": None}


def _add_docx_sections(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Add section anchors from DOCX ``Heading`` paragraphs that start with a number."""
    from docx import Document

    doc = Document(str(source_path))
    for p_idx, para in enumerate(doc.paragraphs):
        style = getattr(para.style, "name", "") or ""
        if not style.lower().startswith("heading"):
            continue
        num = _section_num(para.text)
        if not num:
            continue
        key = canonical_anchor_key("SECTION_REF", num)
        if key in index:
            continue
        index[key] = {"paragraph_index": p_idx}


def _scan_pdf_section_headings(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Fill section anchors from numbered headings in the PDF body, gated on a
    *larger-than-body* font on an isolated line — the only signal that reliably
    separates a heading from the numbered footnotes/criteria clinical PDFs are full
    of. Where headings are not font-distinguished this finds nothing (no guessing).
    """
    import collections

    import fitz  # PyMuPDF

    doc = fitz.open(str(source_path))
    try:
        sizes: collections.Counter = collections.Counter()
        candidates: list[tuple[int, str, float, int, list[float] | None]] = []
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type", 0) != 0:
                    continue
                blk_lines = block.get("lines", [])
                for line in blk_lines:
                    spans = [s for s in line.get("spans", []) if s.get("text", "").strip()]
                    if not spans:
                        continue
                    text = "".join(s.get("text", "") for s in spans).strip()
                    max_size = max(round(float(s.get("size", 0.0)), 1) for s in spans)
                    for s in spans:
                        sizes[round(float(s.get("size", 0.0)), 1)] += len(s.get("text", ""))
                    candidates.append(
                        (page_index, text, max_size, len(blk_lines),
                         _union_bbox([s.get("bbox") for s in spans]))
                    )
        body_size = sizes.most_common(1)[0][0] if sizes else 0.0
        for page_index, text, max_size, n_block_lines, bbox in candidates:
            if n_block_lines != 1 or max_size < body_size + 1.5:
                continue  # heading must stand alone in its block, in a larger font
            num = _section_heading_num(text)
            if not num:
                continue
            key = canonical_anchor_key("SECTION_REF", num)
            if key in index:
                continue  # first definition wins (ToC source already had it)
            index[key] = {"page_index": page_index, "bbox": bbox}
    finally:
        doc.close()


def _scan_docx_section_headings(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Fill section anchors from numbered DOCX paragraphs that are a Heading style
    or **bold** — covering manually-formatted headings that ``_add_docx_sections``
    (Heading-style only) misses, while the style/bold gate keeps body prose out."""
    from docx import Document

    doc = Document(str(source_path))
    for p_idx, para in enumerate(doc.paragraphs):
        num = _section_heading_num(para.text or "")
        if not num:
            continue
        style = (getattr(para.style, "name", "") or "").lower()
        is_heading_style = style.startswith("heading") or style.startswith("title")
        is_bold = bool(para.runs) and any(bool(r.bold) for r in para.runs)
        if not (is_heading_style or is_bold):
            continue  # need a typographic heading signal beyond the numbered shape
        key = canonical_anchor_key("SECTION_REF", num)
        if key in index:
            continue
        index[key] = {"paragraph_index": p_idx}


def _add_pdf_visit_anchors(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Map visit citations to visit sections using the PDF's own ToC / bookmarks.

    A ToC entry "9.3.3.1 Week 2 (± 3 Days)" yields ``visit_ref_week_2`` → that
    section's page, so a "Week 2 Visit" citation anchors at the visit section
    rather than at itself. First heading wins per visit key.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(str(source_path))
    try:
        toc = doc.get_toc(simple=True)  # [[level, title, page_1based], ...]
    finally:
        doc.close()
    for entry in toc:
        try:
            _level, title, page = entry[0], entry[1], entry[2]
        except (IndexError, TypeError):
            continue
        m = _VISIT_IN_HEADING.search(title or "")
        if not m or page is None or page < 1:
            continue
        key = canonical_visit_key(m.group("unit"), m.group("n"))
        if key in index:
            continue
        index[key] = {"page_index": int(page) - 1, "bbox": None}


def _add_docx_visit_anchors(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Map visit citations to visit sections from DOCX ``Heading`` paragraphs whose
    text carries a visit phrase, e.g. a "Week 2" heading → ``visit_ref_week_2``."""
    from docx import Document

    doc = Document(str(source_path))
    for p_idx, para in enumerate(doc.paragraphs):
        style = (getattr(para.style, "name", "") or "").lower()
        if not (style.startswith("heading") or style.startswith("title")):
            continue
        m = _VISIT_IN_HEADING.search(para.text or "")
        if not m:
            continue
        key = canonical_visit_key(m.group("unit"), m.group("n"))
        if key in index:
            continue
        index[key] = {"paragraph_index": p_idx}


def _add_docx_references(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Index the DOCX References section: each entry → its paragraph, plus a
    ``ref_heading`` fallback at the 'REFERENCES' heading itself."""
    from docx import Document

    from hyperlink_engine.core.injection.ref_index import (
        REF_HEADING_KEY,
        is_next_section_after_refs,
        is_references_heading,
        parse_ref_entry_key,
    )

    doc = Document(str(source_path))
    paras = doc.paragraphs
    start = next(
        (i for i, p in enumerate(paras) if is_references_heading(p.text or "")), None
    )
    if start is None:
        return
    index.setdefault(REF_HEADING_KEY, {"paragraph_index": start})
    for i in range(start + 1, len(paras)):
        text = (paras[i].text or "").strip()
        if not text:
            continue
        if is_next_section_after_refs(text):
            break
        key = parse_ref_entry_key(text)
        if key and key not in index:
            index[key] = {"paragraph_index": i}


def _add_pdf_references(index: dict[str, dict[str, Any]], source_path: str | Path) -> None:
    """Index the PDF References section by scanning text lines (joined spans): each
    entry → its page, plus a ``ref_heading`` fallback at the 'REFERENCES' heading."""
    import fitz  # PyMuPDF

    from hyperlink_engine.core.injection.ref_index import (
        REF_HEADING_KEY,
        is_next_section_after_refs,
        is_references_heading,
        parse_ref_entry_key,
    )

    doc = fitz.open(str(source_path))
    try:
        lines: list[tuple[int, str]] = []
        for page_index in range(doc.page_count):
            page = doc.load_page(page_index)
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type", 0) != 0:
                    continue
                for line in block.get("lines", []):
                    spans = [s for s in line.get("spans", []) if s.get("text", "").strip()]
                    if not spans:
                        continue
                    text = "".join(s.get("text", "") for s in spans).strip()
                    lines.append((page_index, text))
    finally:
        doc.close()

    start = next((i for i, (_p, t) in enumerate(lines) if is_references_heading(t)), None)
    if start is None:
        return
    index.setdefault(REF_HEADING_KEY, {"page_index": lines[start][0], "bbox": None})
    for page_index, text in lines[start + 1 :]:
        if is_next_section_after_refs(text):
            break
        key = parse_ref_entry_key(text)
        if key and key not in index:
            index[key] = {"page_index": page_index, "bbox": None}
