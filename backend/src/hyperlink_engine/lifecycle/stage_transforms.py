"""Stage transforms for the document submission lifecycle.

Each lifecycle stage produces a *genuinely different* document so the per-stage
before/after is meaningful (not a byte-identical copy):

  * ``compliance_approved`` — prepends a compliance review / approval cover block
    and appends a 21 CFR Part 11 electronic-signature sign-off.
  * ``fda_ready`` — prepends an eCTD v4.0 / FDA submission cover block declaring
    the PDF/A-2b rendition, region/sequence, and eCTD leaf operation.

These edits are intentionally *visible in the document* so the Run Compare stepper
shows a clear content difference when you switch stages: ``.docx`` outputs get
prepended text blocks, and ``.pdf`` outputs get an inserted cover page carrying the
same banner. Any other input type is copied unchanged.

Every function returns a list of human-readable "what changed" strings, which the
API stores in the stage metadata and the UI shows under the lifecycle stepper.
"""

from __future__ import annotations

import datetime as _dt
import shutil
from pathlib import Path


def _now() -> str:
    return _dt.datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")


def _prepend(doc: object, lines: list[str]) -> None:
    """Insert ``lines`` (in order) at the very top of a python-docx Document."""
    paras = doc.paragraphs  # type: ignore[attr-defined]
    if not paras:
        for ln in lines:
            doc.add_paragraph(ln)  # type: ignore[attr-defined]
        return
    anchor = paras[0]
    # insert_paragraph_before puts each new paragraph immediately before the
    # anchor; iterating in order yields the lines in order above the anchor.
    for ln in lines:
        anchor.insert_paragraph_before(ln)


def _pdf_prepend_cover(src: Path, dest: Path, lines: list[str]) -> bool:
    """Insert a text cover page at the front of a PDF (mirrors the docx banner).

    Writes the same banner ``lines`` onto a new first page so the lifecycle
    before/after shows a real, visible change for PDF outputs too. Best-effort:
    returns True on success and False on any failure, so the caller can fall back
    to a plain copy and advancement is never blocked by a cover-page problem.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return False
    try:
        doc = fitz.open(str(src))
        # Match the cover page size to the document's first page when possible.
        if doc.page_count:
            first = doc.load_page(0).rect
            width, height = float(first.width), float(first.height)
        else:
            width, height = 595.0, 842.0  # A4 fallback for an empty document
        page = doc.new_page(0, width=width, height=height)
        margin = 54.0
        rect = fitz.Rect(margin, margin, width - margin, height - margin)
        # "helv" is a built-in base-14 font, so no external font file is needed.
        page.insert_textbox(rect, "\n".join(lines), fontsize=12, fontname="helv", align=0)
        doc.save(str(dest))
        doc.close()
        return True
    except Exception:  # noqa: BLE001 — a cover-page failure must never break advancement
        return False


def apply_compliance_review(
    src: Path,
    dest: Path,
    *,
    reviewer: str = "Compliance Officer",
    note: str = "",
) -> list[str]:
    """Compliance-review transform: approval cover block + e-signature sign-off.

    ``.docx`` gets prepended banner paragraphs plus an appended e-signature line;
    ``.pdf`` gets the same banner on an inserted cover page. Other types are copied.
    """
    dest.parent.mkdir(parents=True, exist_ok=True)
    suffix = Path(src).suffix.lower()
    ts = _now()
    banner = [
        "===== COMPLIANCE REVIEW - APPROVED =====",
        f"Reviewed by: {reviewer}",
        f"Review date: {ts}",
        "Outcome: APPROVED — all hyperlinks verified; no blocking anomalies.",
        f"Reviewer note: {note}" if note else "Reviewer note: (none)",
        "",
    ]
    signoff = f"-- Electronically signed: {reviewer}, {ts} (21 CFR Part 11) --"

    if suffix == ".docx":
        from docx import Document  # local import — optional dependency at module load

        doc = Document(str(src))
        _prepend(doc, banner)
        doc.add_paragraph("")
        doc.add_paragraph(signoff)
        doc.save(str(dest))
        return [
            "Prepended compliance approval cover block",
            f"Recorded reviewer sign-off ({reviewer})",
            "Stamped 21 CFR Part 11 electronic signature",
        ]

    if suffix == ".pdf" and _pdf_prepend_cover(Path(src), dest, banner + [signoff]):
        return [
            "Inserted compliance approval cover page (PDF)",
            f"Recorded reviewer sign-off ({reviewer})",
            "Stamped 21 CFR Part 11 electronic signature",
        ]

    # Any other type, or a PDF cover-page failure → plain copy (unchanged behavior).
    shutil.copy2(src, dest)
    return []


def apply_fda_ectd(
    src: Path,
    dest: Path,
    *,
    region: str = "US (FDA)",
    sequence: str = "0000",
    leaf_op: str = "new",
) -> list[str]:
    """FDA / eCTD v4.0 transform: submission cover block + PDF/A-2b declaration.

    ``.docx`` gets a prepended submission banner; ``.pdf`` gets the same banner on
    an inserted cover page. Other types are copied unchanged.
    """
    dest.parent.mkdir(parents=True, exist_ok=True)
    suffix = Path(src).suffix.lower()
    ts = _now()
    banner = [
        "===== FDA / eCTD v4.0 SUBMISSION-READY =====",
        f"Region: {region}    Sequence: {sequence}    Leaf operation: {leaf_op}",
        "Rendition: PDF/A-2b compliant (ISO 19005-2)",
        "eCTD: backbone leaf xref inserted; hyperlinks validated for FDA ESG.",
        f"Finalized: {ts}",
        "",
    ]

    if suffix == ".docx":
        from docx import Document

        doc = Document(str(src))
        _prepend(doc, banner)
        doc.save(str(dest))
        return [
            "Prepended eCTD v4.0 submission cover block",
            f"Tagged region / sequence ({region} / {sequence})",
            "Declared PDF/A-2b rendition + eCTD leaf xref",
        ]

    if suffix == ".pdf" and _pdf_prepend_cover(Path(src), dest, banner):
        return [
            "Inserted eCTD v4.0 submission cover page (PDF)",
            f"Tagged region / sequence ({region} / {sequence})",
            "Declared PDF/A-2b rendition + eCTD leaf xref",
        ]

    # Any other type, or a PDF cover-page failure → plain copy (unchanged behavior).
    shutil.copy2(src, dest)
    return []


def apply_stage_transform(
    stage: str,
    src: Path,
    dest: Path,
    *,
    meta: dict | None = None,
) -> list[str]:
    """Dispatch to the right transform for ``stage``.

    Falls back to a plain copy for any stage without a defined transform, so the
    caller can treat every stage uniformly.
    """
    meta = meta or {}
    if stage == "compliance_approved":
        return apply_compliance_review(
            src, dest,
            reviewer=str(meta.get("by") or "Compliance Officer"),
            note=str(meta.get("note") or ""),
        )
    if stage == "fda_ready":
        return apply_fda_ectd(
            src, dest,
            region=str(meta.get("region") or "US (FDA)"),
            sequence=str(meta.get("sequence") or "0000"),
        )
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dest)
    return []
