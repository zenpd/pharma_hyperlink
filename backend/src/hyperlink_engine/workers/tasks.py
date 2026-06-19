"""W7.1 — Pipeline stage tasks.

Each task is a thin wrapper around an existing pure-Python implementation
(from ``detection/``, ``injection/``, ``validation/``…) so that:

  * It can be invoked **synchronously** for tests and the W1.5 spike
    (the underlying function is a normal Python call).
  * It can be invoked **asynchronously** via Celery for the bulk pipeline.

The Celery wiring uses the same factory pattern as Phase 1's Neo4j
adapter — Celery is imported only when ``celery_eager=False`` and a real
worker is running. In eager mode (the default in tests), every Celery
``.delay()`` call returns an ``EagerResult`` containing the synchronous
return value, so callers don't need a branch.

All task inputs and outputs are JSON-serializable (Celery requirement).
Pydantic models flow through ``.model_dump()`` / ``.model_validate()``.
"""

from __future__ import annotations

import hashlib
import re
import time
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from hyperlink_engine.config.logging_setup import get_logger
from hyperlink_engine.config.settings import get_settings
from hyperlink_engine.core.injection.docx_linker import DocxLinker
from hyperlink_engine.core.reporting.csv_exporter import write_link_records
from hyperlink_engine.core.validation.existence_checker import LinkProbe, check_all
from hyperlink_engine.models import (
    LinkKind,
    LinkRecord,
    LinkStatus,
    RunLocation,
)
from hyperlink_engine.workers.cache import ExtractorConfig, get_extractor
from hyperlink_engine.workers.celery_app import (
    PIPELINE_STAGES,
    get_app,
    stage_task_name,
)

_log = get_logger("pipeline.tasks")


# ─────────────────────────────────────────────────────────────────────────
# Stage 1 — ingestion: hash the source doc so downstream stages know it
# ─────────────────────────────────────────────────────────────────────────


def ingest_document(source_path: str) -> dict[str, Any]:
    """Stage 1 (synchronous primitive).

    Computes a streaming SHA-256 of the source document. The result rides
    along through every downstream stage so the audit log can prove the
    pipeline acted on a specific document version (21 CFR Part 11 trail).
    """
    path = Path(source_path)
    if not path.exists():
        raise FileNotFoundError(source_path)
    sha = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            sha.update(chunk)
    size = path.stat().st_size
    return {
        "source_path": str(path),
        "sha256": sha.hexdigest(),
        "file_size_bytes": size,
    }


# ─────────────────────────────────────────────────────────────────────────
# Stage 2 — detection
# ─────────────────────────────────────────────────────────────────────────


def detect_references(
    ingest_record: dict[str, Any],
    *,
    extractor_config: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """Stage 2 — run the cascade against every paragraph/run (docx) or page/span (pdf).

    Returns a flat list of detection records (one per detected span)
    plus the original ingest record for downstream stages.
    """
    source_path = ingest_record["source_path"]
    suffix = Path(source_path).suffix.lower()

    if suffix == ".pdf":
        return _detect_references_pdf(ingest_record, extractor_config=extractor_config)

    from hyperlink_engine.core.injection.ref_index import references_span

    cfg = ExtractorConfig(**(extractor_config or {}))
    extractor = get_extractor(cfg)
    cfg_settings = get_settings()
    doc = Document(ingest_record["source_path"])
    detections: list[dict[str, Any]] = []
    # Where the References section sits, so we don't detect a bibliography entry's
    # own "Surname … Journal 1976" tail as an in-text citation.
    ref_start, ref_end = references_span(
        ["".join(r.text or "" for r in p.runs) for p in doc.paragraphs]
    )
    for p_idx, para in enumerate(doc.paragraphs):
        para_dets: list[dict[str, Any]] = []
        for r_idx, run in enumerate(para.runs):
            if not run.text or not run.text.strip():
                continue
            for ref in extractor.extract(run.text):
                # A CTD module reference with no usable module number — the NER
                # rule-fallback firing on a document's own "MODULE 2.7" title —
                # would resolve to a broken "#m?" self-bookmark (narrative module
                # refs have no anchor-index entry, so it can never navigate). Skip
                # it; real eCTD leaf paths ("m5/53-…") and "Module 2.5.3" (which
                # carry a 'mod'/'subpath' group) are untouched.
                if ref.label == "CTD_LEAF" and not (
                    ref.groups.get("mod") or ref.groups.get("subpath")
                ):
                    continue
                para_dets.append(
                    {
                        "paragraph_index": p_idx,
                        "run_index": r_idx,
                        "char_start": ref.start,
                        "char_end": ref.end,
                        "pattern_id": ref.pattern_id,
                        "label": ref.label,
                        "text": ref.text,
                        "context": run.text,
                        "confidence": ref.confidence,
                        "source_layer": ref.source_layer,
                        "groups": dict(ref.groups),
                        "llm_consulted": ref.llm_consulted,
                        "llm_confidence_before": ref.llm_confidence_before,
                        "llm_confidence_after": ref.llm_confidence_after,
                        "llm_reasoning": ref.llm_reasoning,
                    }
                )
        # "Statistical Analysis Plan (SAP)" → link only the acronym "(SAP)", drop the
        # spelled-out restatement (publishing preference; handles cross-run splits).
        _dedupe_fullname_acronym_docx(para, para_dets)
        # Whole-reference linking: collapse a compound phrase
        # ("Protocol TMX-67_301 Section 6.1") into ONE continuous link to one target.
        _enrich_compound_docx(p_idx, para, para_dets)
        # Literature citations ("Helget LN, 2024", "[7]") must be matched over the
        # WHOLE paragraph: Word splits a citation across runs (e.g. "Helget" +
        # " LN, 2024"), so a per-run pass would miss it. This pass runs the
        # citation matchers on the joined run text. Skipped INSIDE the References
        # section so entries don't self-cite. Collected into ``para_dets`` so the
        # highlight pass below can extend a citation link to its full span.
        if ref_start is None or not (ref_start <= p_idx < ref_end):
            para_dets.extend(_detect_ref_cites_docx(p_idx, para))
        # Highlight safety net: ensure every yellow-highlighted span is fully
        # linked (one continuous link). Additive + a strict no-op on plain docs.
        if cfg_settings.link_highlighted_spans:
            _ensure_highlighted_linked_docx(p_idx, para, para_dets, extractor)
            # Tag detections sitting on a highlighted run, so the resolver may link a
            # highlighted SELF-reference ("SAP" inside SAP.docx) to its own document's
            # top — a doc normally never links to itself, but the author highlighted it,
            # so for parity with the adjacent cross-ref ("CSR") we honor the mark.
            _runs = list(para.runs)
            for _det in para_dets:
                _ri = _det.get("run_index", 0)
                if 0 <= _ri < len(_runs) and _run_is_highlighted(_runs[_ri]):
                    _det["is_highlighted"] = True
        # Re-target the author's pre-existing PLACEHOLDER hyperlinks ("ISE" →
        # about:blank): text inside a <w:hyperlink> is not in para.runs, so the
        # passes above never saw it. Real links (a SharePoint URL) are left alone.
        para_dets.extend(_detect_existing_hyperlinks_docx(p_idx, para, extractor))
        detections.extend(para_dets)
    return {"ingest": ingest_record, "detections": detections}


def _detect_ref_cites_docx(p_idx: int, para: Any) -> list[dict[str, Any]]:
    """Detect in-text literature citations across a DOCX paragraph's joined runs.

    Author-year ("Helget LN, 2024") and numbered ("[7]") citations are found over
    the full paragraph text, then each match is anchored on the FIRST run it
    overlaps with run-relative char offsets — so a citation split across runs still
    yields a valid single-run link span. The resolver keys the link from ``groups``
    (surname+year / num), so the run-clipped ``text`` is only the visible anchor.
    """
    from hyperlink_engine.core.injection.ref_index import author_year_cites, numbered_cites

    runs = list(para.runs)
    full = "".join(r.text or "" for r in runs)
    if not full.strip():
        return []
    bounds: list[tuple[int, int, int]] = []  # (run_index, start_offset, end_offset)
    cum = 0
    for r_idx, r in enumerate(runs):
        rl = len(r.text or "")
        bounds.append((r_idx, cum, cum + rl))
        cum += rl

    out: list[dict[str, Any]] = []

    def _emit(m: Any, groups: dict[str, str], pattern_id: str) -> None:
        for r_idx, rs, re_ in bounds:
            if re_ <= m.start() or rs >= m.end():
                continue  # no overlap with this run
            cs = max(m.start(), rs) - rs
            ce = min(m.end(), re_) - rs
            if ce <= cs:
                continue
            out.append(
                {
                    "paragraph_index": p_idx,
                    "run_index": r_idx,
                    "char_start": cs,
                    "char_end": ce,
                    "pattern_id": pattern_id,
                    "label": "REF_CITE",
                    "text": full[rs + cs : rs + ce],
                    "context": full,
                    "confidence": 0.6,
                    "source_layer": "regex",
                    "groups": groups,
                    "llm_consulted": False,
                    "llm_confidence_before": None,
                    "llm_confidence_after": None,
                    "llm_reasoning": None,
                }
            )
            return

    for m in author_year_cites(full):
        _emit(m, {"surname": m.group("surname"), "year": m.group("year")}, "REF_CITE_AUTHOR_YEAR_V1")
    for m in numbered_cites(full):
        _emit(m, {"num": m.group("num")}, "REF_CITE_NUM_V1")
    return out


# ─────────────────────────────────────────────────────────────────────────────
# Compound cross-reference unification (whole-reference linking)
#
# The publishing team wants the WHOLE phrase — e.g. "Protocol TMX-67_301
# Section 6.1", "SAP Section 5.3", "Protocol TMX-67_301" — to link to ONE
# destination (the section in the named document). Word splits the phrase across
# runs, so the per-run detector emits the pieces (DOC_REF "Protocol", DOC_ID
# "TMX-67_301", SECTION_REF "Section 6.1") as separate detections that resolve
# to different places (or get dropped). We re-scan the JOINED paragraph text for
# the compound phrase and STAMP every per-run detection inside it with a shared
# qualifier doc-type + section anchor, so the constituent links all resolve to
# the same target ("contiguous links, one destination" — no injector changes).
# ─────────────────────────────────────────────────────────────────────────────
_COMPOUND_REF_RE = re.compile(
    r"(?<![A-Za-z])(?P<head>"
    # 'Study ID <id>' identifies the target by its ID ("Study ID TMX-67_301" →
    # the file named TMX-67_301), NOT by a document-type word. The dossier now
    # writes the protocol/study reference this way.
    r"(?i:study\s+id)"
    # 'Protocol' kept for back-compat: docs that still say "Protocol TMX-67_301"
    # route by doc-type (→ the protocol-typed file). Harmless when absent — the new
    # "Study ID …" docs never trigger it. Remove only if a protocol-typed file must
    # never win over the Study-ID file for the same id.
    r"|(?i:protocol)"
    r"|(?i:statistical\s+analysis\s+plan)|SAP"
    r"|(?i:clinical\s+study\s+report)|CSR"
    r"|(?i:integrated\s+summary\s+of\s+safety)|ISS"
    r"|(?i:integrated\s+summary\s+of\s+efficacy)|ISE"
    r")"
    r"(?P<docid>\s+[A-Z]{2,6}-\d{1,3}_\d{1,4})?"
    r"(?P<struct>\s+(?i:section|sect\.?|table|figure|listing|appendix)"
    r"\s+\d+(?:[.\-]\d+){0,4})?"
)

# Compound heads that identify the target by its ID (the trailing <id>) rather
# than by a document type — routed via token-overlap on the id, not doc-type.
_ID_REF_HEADS = {"study id"}

_STRUCT_KW_TO_LABEL = {
    "section": "SECTION_REF",
    "sect": "SECTION_REF",
    "table": "TABLE_REF",
    "figure": "FIGURE_REF",
    "listing": "LISTING_REF",
    "appendix": "APPENDIX_REF",
}


def _compound_doc_type(head: str) -> str:
    """Map a compound phrase's leading doc-type words to a canonical doc type."""
    h = head.strip().lower()
    if "protocol" in h:
        return "protocol"
    if "statistical analysis plan" in h or h == "sap":
        return "sap"
    if "clinical study report" in h or h == "csr":
        return "csr"
    if "integrated summary of safety" in h or h == "iss":
        return "iss"
    if "integrated summary of efficacy" in h or h == "ise":
        return "ise"
    return ""


def _compound_anchor_key(struct: str) -> str:
    """'Section 6.1' -> 'section_ref_6_1' (the sibling document's bookmark key)."""
    from hyperlink_engine.core.injection.anchor_index import canonical_anchor_key

    m = re.match(r"\s*([A-Za-z.]+)\s+(\d+(?:[.\-]\d+){0,4})", struct)
    if not m:
        return ""
    label = _STRUCT_KW_TO_LABEL.get(m.group(1).rstrip(".").lower())
    if not label:
        return ""
    return canonical_anchor_key(label, m.group(2))


_FULLNAME_ACRONYM = (
    ("statistical analysis plan", "sap"),
    ("clinical study report", "csr"),
    ("integrated summary of safety", "iss"),
    ("integrated summary of efficacy", "ise"),
)


def _dedupe_fullname_acronym_docx(para: Any, para_dets: list[dict[str, Any]]) -> None:
    """When a spelled-out doc name is immediately followed by its acronym in parens —
    "Statistical Analysis Plan (SAP)" — keep only the ACRONYM link and drop the
    spelled-out one (publishing preference: link the short tag). Works across runs
    (Word often splits the phrase and the "(SAP)" into different runs), which the
    per-run regex validator cannot see. Mutates ``para_dets`` in place."""
    runs = list(para.runs)
    full = "".join(r.text or "" for r in runs)
    low = full.lower()
    if "(" not in low:
        return
    run_start: list[int] = []
    cum = 0
    for r in runs:
        run_start.append(cum)
        cum += len(r.text or "")

    drop: list[tuple[int, int]] = []  # (lo, hi) spelled-out spans to remove
    for sp, ac in _FULLNAME_ACRONYM:
        start = 0
        while True:
            i = low.find(sp, start)
            if i < 0:
                break
            after = low[i + len(sp) : i + len(sp) + 8].lstrip()
            if after.startswith("(") and after[1:].lstrip().startswith(ac):
                drop.append((i, i + len(sp)))
            start = i + len(sp)
    if not drop:
        return

    kept: list[dict[str, Any]] = []
    for det in para_dets:
        if det.get("label") == "DOC_REF":
            ri = det.get("run_index", 0)
            if ri < len(run_start):
                ds = run_start[ri] + det["char_start"]
                de = run_start[ri] + det["char_end"]
                if any(lo <= ds and de <= hi for lo, hi in drop):
                    continue  # drop the spelled-out form; the acronym det remains
        kept.append(det)
    para_dets[:] = kept


def _enrich_compound_docx(p_idx: int, para: Any, para_dets: list[dict[str, Any]]) -> None:
    """Collapse a compound cross-reference phrase into ONE continuous link.

    A phrase like "Protocol TMX-67_301 Section 6.1" is emitted by the per-run
    detector as separate pieces (DOC_REF "Protocol", DOC_ID "TMX-67_301",
    SECTION_REF "Section 6.1"). Those pieces are REMOVED and replaced by a single
    compound detection that spans the WHOLE phrase — across multiple runs when
    Word split it — so the injector wraps it in one ``w:hyperlink`` (the publishing
    team's "continuous, whole-reference" requirement) and all of it resolves to the
    one destination (the named document's section). Mutates ``para_dets`` in place.
    """
    runs = list(para.runs)
    full = "".join(r.text or "" for r in runs)
    if not full.strip():
        return
    run_start: list[int] = []
    cum = 0
    for r in runs:
        run_start.append(cum)
        cum += len(r.text or "")
    total = cum

    def _locate(off: int) -> tuple[int, int]:
        """Global char offset → (run_index, char-within-run)."""
        for r_idx in range(len(runs)):
            rs = run_start[r_idx]
            re_ = rs + len(runs[r_idx].text or "")
            if rs <= off < re_:
                return r_idx, off - rs
        last = max(len(runs) - 1, 0)
        return last, len(runs[last].text or "") if runs else 0

    compounds: list[dict[str, Any]] = []
    consumed: list[tuple[int, int]] = []  # (lo, hi) spans replaced by a compound
    for m in _COMPOUND_REF_RE.finditer(full):
        # A bare doc-type word (no id, no structural ref) is already handled by
        # the plain DOC_REF pattern — only an id- or section-qualified phrase is a
        # compound that needs unifying.
        if not m.group("docid") and not m.group("struct"):
            continue
        anchor_key = _compound_anchor_key(m.group("struct")) if m.group("struct") else ""
        head_l = m.group("head").strip().lower()
        is_id_ref = head_l in _ID_REF_HEADS
        if is_id_ref:
            # "Study ID TMX-67_301 Section 6.1" — no doc-type word: the target is
            # identified by the trailing id. The link TEXT is the WHOLE phrase (so the
            # injected hyperlink AND the preview box span the whole continuous
            # reference, like the DOC_REF branch), but routing uses the bare id alone
            # via `compound_id_token` so the trailing "Section 6.1"/"Study"/"ID" words
            # never skew the file token-overlap ("Study ID TMX-67_301.pdf"). The
            # section anchor links the whole phrase to that file's section. Requires
            # the id to be present.
            if not m.group("docid"):
                continue
            label = "DOC_ID"
            det_text = None  # → full phrase (continuous link label), set below
            groups: dict[str, Any] = {"compound_id_token": m.group("docid").strip()}
        else:
            dt = _compound_doc_type(m.group("head"))
            if not dt:
                continue
            label = "DOC_REF"
            det_text = None  # → full phrase, set below
            groups = {"compound_qual_dt": dt}
        if anchor_key:
            groups["compound_anchor_key"] = anchor_key
        lo, hi = m.start(), min(m.end(), total)
        if hi <= lo:
            continue
        s_run, cs = _locate(lo)
        e_run, _ = _locate(hi - 1)
        ce = hi - run_start[e_run]
        compounds.append(
            {
                "paragraph_index": p_idx,
                "run_index": s_run,
                "char_start": cs,
                # For a single-run phrase char_end IS the end; for a multi-run phrase
                # it just needs to satisfy RunLocation (>0) — the real end is
                # carried by end_run_index/end_char and used by the injector.
                "char_end": ce if e_run == s_run else max(len(runs[s_run].text or ""), 1),
                "end_run_index": e_run,
                "end_char": ce,
                "pattern_id": "COMPOUND_REF_V1",
                "label": label,
                # Both branches set text = the whole phrase (the continuous link
                # label). DOC_REF routes by `compound_qual_dt`; an id-routed DOC_ID
                # routes by `compound_id_token` (the bare id) — see _resolve_one's
                # token-overlap fallback.
                "text": det_text if det_text is not None else full[lo:hi],
                "context": full,
                "confidence": 0.9,
                "source_layer": "regex",
                "groups": groups,
                "llm_consulted": False,
                "llm_confidence_before": None,
                "llm_confidence_after": None,
                "llm_reasoning": None,
                "is_compound": True,
            }
        )
        consumed.append((lo, hi))

    if not compounds:
        return

    # Drop the per-run pieces that the compound now covers (so they don't also
    # emit their own separate links inside the same phrase).
    kept: list[dict[str, Any]] = []
    for det in para_dets:
        ds = run_start[det["run_index"]] + det["char_start"] if det["run_index"] < len(run_start) else -1
        de = run_start[det["run_index"]] + det["char_end"] if det["run_index"] < len(run_start) else -1
        if any(lo <= ds and de <= hi for lo, hi in consumed):
            continue
        kept.append(det)
    para_dets[:] = kept + compounds


# ─────────────────────────────────────────────────────────────────────────────
# Highlight safety net (whole-reference linking, authored docs)
#
# The publishing team marks every citation that must link with a yellow
# highlight. The detection cascade links by TEXT (regex/NER) — independent of the
# highlight — so on a plain (un-highlighted) document it behaves exactly the same.
# But on an AUTHORED doc we additionally guarantee that no highlighted citation is
# left partially/un-linked: for each contiguous highlighted span that still has an
# un-linked word, we make the WHOLE span one continuous link, routing the target
# through the normal resolver. Additive (never drops a detection link), and a
# strict no-op when there are no highlights.
# ─────────────────────────────────────────────────────────────────────────────
_YELLOWISH_FILLS = {"FFFF00", "FFEE80", "FFFF99", "FFF200", "FFE600", "FFFFCC"}
_DOCTYPE_BEFORE = (
    (r"statistical\s+analysis\s+plan", "sap"),
    (r"clinical\s+study\s+report", "csr"),
    (r"integrated\s+summary\s+of\s+safety", "iss"),
    (r"integrated\s+summary\s+of\s+efficacy", "ise"),
    (r"\bprotocol\b", "protocol"),
    (r"\bSAP\b", "sap"),
    (r"\bCSR\b", "csr"),
    (r"\bISS\b", "iss"),
    (r"\bISE\b", "ise"),
)


def _run_is_highlighted(run: Any) -> bool:
    """A run carries a citation highlight (Word YELLOW highlight or yellow shading)."""
    from docx.enum.text import WD_COLOR_INDEX

    hl = run.font.highlight_color
    if hl is not None and hl != WD_COLOR_INDEX.AUTO:
        return True
    rPr = run._element.rPr
    if rPr is not None:
        shd = rPr.find(qn("w:shd"))
        if shd is not None and (shd.get(qn("w:fill")) or "").upper() in _YELLOWISH_FILLS:
            return True
    return False


def _doctype_before(full: str, lo: int) -> str:
    """Find the doc-type named just before offset ``lo`` (same paragraph), so a bare
    id like 'TMX-67_301' in 'Protocol ID- TMX-67_301' routes to the Protocol doc."""
    prefix = full[:lo]
    best_pos, best_dt = -1, ""
    for pat, dt in _DOCTYPE_BEFORE:
        for m in re.finditer(pat, prefix, re.IGNORECASE):
            if m.start() > best_pos:
                best_pos, best_dt = m.start(), dt
    return best_dt if best_pos >= 0 and (lo - best_pos) <= 45 else ""


def _ensure_highlighted_linked_docx(
    p_idx: int, para: Any, para_dets: list[dict[str, Any]], extractor: Any
) -> None:
    """Guarantee every yellow-highlighted span is one continuous link.

    For each contiguous highlighted span that still has an un-linked *word*
    character, build ONE continuous link over the whole span. The link's target is
    taken from an existing detection inside the span (extended to the span), or —
    when none exists — by classifying the span text with the extractor. A span with
    no resolvable target is left as-is (e.g. a self-reference or a bibliography
    entry), so this never invents a bad link.
    """
    runs = list(para.runs)
    full = "".join(r.text or "" for r in runs)
    if not full.strip():
        return
    run_start: list[int] = []
    cum = 0
    for r in runs:
        run_start.append(cum)
        cum += len(r.text or "")
    total = cum

    hot = [False] * total
    for ri, r in enumerate(runs):
        if _run_is_highlighted(r):
            for k in range(run_start[ri], run_start[ri] + len(r.text or "")):
                hot[k] = True
    if not any(hot):
        return  # strict no-op on plain documents

    def _locate(off: int) -> tuple[int, int]:
        for ri in range(len(runs)):
            rs = run_start[ri]
            re_ = rs + len(runs[ri].text or "")
            if rs <= off < re_:
                return ri, off - rs
        last = max(len(runs) - 1, 0)
        return last, (len(runs[last].text or "") if runs else 0)

    # Which characters are already inside a detection link.
    covered = [False] * total
    det_bounds: list[tuple[int, int, dict[str, Any]]] = []
    for det in para_dets:
        ri = det.get("run_index", 0)
        if ri >= len(run_start):
            continue
        gs = run_start[ri] + det["char_start"]
        er = det.get("end_run_index")
        if er is not None and er < len(run_start):
            ge = run_start[er] + (det.get("end_char") or 0)
        else:
            ge = run_start[ri] + det["char_end"]
        det_bounds.append((gs, ge, det))
        # A bare DOC_ID with no doc-type qualifier may not resolve to a link
        # ("TMX-67_301" alone names no file) — don't let it count as "already linked",
        # so the highlight pass can upgrade it (stamp the doc-type from its context). A
        # DOC_REF ("CSR"/"SAP") already names its own type, so it is left as covered.
        # EXCEPTION: a COMPOUND DOC_ID ("Study ID TMX-67_301") is a deliberate
        # multi-run span — keep it covered so a highlight GAP (the un-highlighted
        # space between "Study ID" and "TMX-67_301") can't split it back into two.
        weak = (
            det.get("label") == "DOC_ID"
            and not (det.get("groups") or {}).get("compound_qual_dt")
            and not det.get("is_compound")
        )
        if not weak:
            for k in range(max(gs, 0), min(ge, total)):
                covered[k] = True

    # Contiguous highlighted spans.
    spans: list[tuple[int, int]] = []
    i = 0
    while i < total:
        if hot[i]:
            j = i
            while j < total and hot[j]:
                j += 1
            spans.append((i, j))
            i = j
        else:
            i += 1

    new_dets: list[dict[str, Any]] = []
    remove_ids: set[int] = set()
    for lo, hi in spans:
        while lo < hi and full[lo].isspace():
            lo += 1
        while hi > lo and full[hi - 1].isspace():
            hi -= 1
        if hi <= lo:
            continue
        # Already fully linked? (ignore non-word chars like '.', '-', '–')
        if all(covered[k] or not full[k].isalnum() for k in range(lo, hi)):
            continue
        overlap = [(gs, ge, d) for gs, ge, d in det_bounds if gs < hi and ge > lo]
        # Don't merge two genuinely different references (e.g. Figure 1 / Figure 2).
        if len({d.get("label") for _gs, _ge, d in overlap}) > 1:
            continue
        span_text = full[lo:hi]
        if overlap:
            carrier = overlap[0][2]
            label = carrier.get("label")
            groups = dict(carrier.get("groups") or {})
            pattern_id = carrier.get("pattern_id") or "HIGHLIGHT_SPAN_V1"
            source_layer = carrier.get("source_layer", "regex")
            confidence = carrier.get("confidence", 0.6)
            if label == "DOC_ID" and not groups.get("compound_qual_dt"):
                dt = _doctype_before(full, lo)
                if dt:
                    groups["compound_qual_dt"] = dt
        else:
            hits = list(extractor.extract(span_text))
            if not hits:
                continue  # unclassifiable highlighted text → leave it (no bad link)
            best = max(hits, key=lambda h: h.end - h.start)
            label = best.label
            groups = dict(best.groups)
            pattern_id = best.pattern_id or "HIGHLIGHT_SPAN_V1"
            source_layer = best.source_layer
            confidence = best.confidence
        for _gs, _ge, d in overlap:
            remove_ids.add(id(d))
        s_run, cs = _locate(lo)
        e_run, _ = _locate(hi - 1)
        ce = hi - run_start[e_run]
        new_dets.append(
            {
                "paragraph_index": p_idx,
                "run_index": s_run,
                "char_start": cs,
                "char_end": ce if e_run == s_run else max(len(runs[s_run].text or ""), 1),
                "end_run_index": e_run,
                "end_char": ce,
                "pattern_id": pattern_id,
                "label": label,
                "text": span_text,
                "context": full,
                "confidence": confidence,
                "source_layer": source_layer,
                "groups": groups,
                "llm_consulted": False,
                "llm_confidence_before": None,
                "llm_confidence_after": None,
                "llm_reasoning": None,
                "is_highlight_span": True,
            }
        )

    if not new_dets and not remove_ids:
        return
    para_dets[:] = [d for d in para_dets if id(d) not in remove_ids] + new_dets


# A pre-existing hyperlink whose target is one of these is an unfilled placeholder
# the author meant to point at a real document — safe to re-target. Anything else
# (a real http URL, a working bookmark) is left untouched.
_PLACEHOLDER_HL_TARGETS = {None, "", "about:blank", "#", "about:blank#"}


def _detect_existing_hyperlinks_docx(
    p_idx: int, para: Any, extractor: Any
) -> list[dict[str, Any]]:
    """Detect references inside the author's pre-existing PLACEHOLDER hyperlinks.

    Text wrapped in a ``w:hyperlink`` is invisible to ``para.runs``, so the normal
    passes never see "ISE"/"ISS"/"SAP" the author already linked to ``about:blank``.
    We read each existing hyperlink's text, and when it is still a placeholder AND
    the text resolves as a reference, emit a detection carrying ``existing_hl_index``
    so injection re-points that exact hyperlink (instead of nesting a new one)."""
    out: list[dict[str, Any]] = []
    hls = para._p.findall(qn("w:hyperlink"))
    for h_idx, hl in enumerate(hls):
        rid = hl.get(qn("r:id"))
        anchor = hl.get(qn("w:anchor"))
        current = None
        if rid is not None:
            rel = para.part.rels.get(rid)
            current = rel.target_ref if rel is not None else None
        elif anchor is not None:
            current = f"#{anchor}"
        if current not in _PLACEHOLDER_HL_TARGETS:
            continue  # a real link (e.g. a SharePoint URL) — never touch it
        htext = "".join(t.text or "" for t in hl.iter(qn("w:t")))
        if not htext.strip():
            continue
        hits = list(extractor.extract(htext))
        if not hits:
            continue
        best = max(hits, key=lambda h: h.end - h.start)
        out.append(
            {
                "paragraph_index": p_idx,
                "run_index": 0,
                "char_start": 0,
                "char_end": 1,  # unused: injection re-targets the existing hyperlink
                "existing_hl_index": h_idx,
                "pattern_id": best.pattern_id or "EXISTING_HL_V1",
                "label": best.label,
                "text": htext,
                "context": htext,
                "confidence": best.confidence,
                "source_layer": best.source_layer,
                "groups": dict(best.groups),
                "llm_consulted": False,
                "llm_confidence_before": None,
                "llm_confidence_after": None,
                "llm_reasoning": None,
            }
        )
    return out


def _detect_references_pdf(
    ingest_record: dict[str, Any],
    *,
    extractor_config: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """PDF-specific detection: extract text from PyMuPDF spans and run the cascade."""
    import fitz

    from hyperlink_engine.core.injection.anchor_index import _union_bbox
    from hyperlink_engine.core.injection.ref_index import numbered_cites as _numbered_cites

    cfg = ExtractorConfig(**(extractor_config or {}))
    extractor = get_extractor(cfg)
    source_path = ingest_record["source_path"]
    doc = fitz.open(str(source_path))
    detections: list[dict[str, Any]] = []

    for page_index in range(doc.page_count):
        page = doc.load_page(page_index)
        page_dict = page.get_text("dict")
        span_idx = 0
        for block in page_dict.get("blocks", []):
            if block.get("type", 0) != 0:
                continue
            for line in block.get("lines", []):
                # Collect the line's spans so a numbered citation marker that PyMuPDF
                # split across spans ("[", "14", "]") can be reassembled — the
                # per-span extractor still runs span-by-span for the other patterns.
                line_spans: list[tuple[int, int, int, tuple[float, float, float, float]]] = []
                line_text = ""
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    bbox = span.get("bbox", (0.0, 0.0, 0.0, 0.0))
                    start_off = len(line_text)
                    line_text += text
                    line_spans.append((span_idx, start_off, len(line_text), bbox))
                    if text and text.strip():
                        for ref in extractor.extract(text):
                            detections.append(
                                {
                                    "is_pdf": True,
                                    "page_index": page_index,
                                    "span_index": span_idx,
                                    "char_start": ref.start,
                                    "char_end": ref.end,
                                    "bbox": list(bbox),
                                    "pattern_id": ref.pattern_id,
                                    "label": ref.label,
                                    "text": ref.text,
                                    "context": text,
                                    "confidence": ref.confidence,
                                    "source_layer": ref.source_layer,
                                    "groups": dict(ref.groups),
                                    "llm_consulted": ref.llm_consulted,
                                    "llm_confidence_before": ref.llm_confidence_before,
                                    "llm_confidence_after": ref.llm_confidence_after,
                                    "llm_reasoning": ref.llm_reasoning,
                                }
                            )
                    span_idx += 1
                # Numbered literature citations ("[7]") over the JOINED line → their
                # References entry. Brackets are required, so a numbered entry
                # ("7. …") is never mistaken for a citation; author-year is NOT run
                # on PDFs (an entry's "Journal 1976" tail would false-match).
                for cm in _numbered_cites(line_text):
                    overlap = [
                        (s_idx, so, eo, bb)
                        for s_idx, so, eo, bb in line_spans
                        if not (eo <= cm.start() or so >= cm.end())
                    ]
                    if not overlap:
                        continue
                    ubox = _union_bbox([bb for _s, _a, _b, bb in overlap])
                    detections.append(
                        {
                            "is_pdf": True,
                            "page_index": page_index,
                            "span_index": overlap[0][0],
                            "char_start": cm.start() - overlap[0][1],
                            "char_end": max(cm.end() - overlap[0][1], 1),
                            "bbox": list(ubox) if ubox else list(overlap[0][3]),
                            "pattern_id": "REF_CITE_NUM_V1",
                            "label": "REF_CITE",
                            "text": cm.group(0),
                            "context": line_text,
                            "confidence": 0.6,
                            "source_layer": "regex",
                            "groups": {"num": cm.group("num")},
                            "llm_consulted": False,
                            "llm_confidence_before": None,
                            "llm_confidence_after": None,
                            "llm_reasoning": None,
                        }
                    )

    doc.close()
    return {"ingest": ingest_record, "detections": detections}


# ─────────────────────────────────────────────────────────────────────────
# Stage 3 — injection
# ─────────────────────────────────────────────────────────────────────────


def inject_links(
    detection_record: dict[str, Any],
    *,
    output_path: str,
    target_anchor_indexes: dict[str, dict[str, Any]] | None = None,
    incoming_anchor_keys: set[str] | None = None,
) -> dict[str, Any]:
    """Stage 3 — turn detected references into actual hyperlinks.

    Mirrors the Phase 1 acceptance script's target-resolution heuristic
    so the gate scoreboard reproduces byte-for-byte from the new pipeline.

    ``target_anchor_indexes`` (PLAN TEN) maps each doc *stem* to its anchor index
    so an internal link lands on the reference's **definition** (the caption /
    heading) rather than the first citation, and a cross-doc link can target the
    exact heading in the sibling document. Optional: when absent the source's own
    index is built inline and cross-doc links stay document-level (today's behavior).
    """
    from hyperlink_engine.core.injection.anchor_index import build_anchor_index
    from hyperlink_engine.core.injection.ref_index import REF_HEADING_KEY

    source = Path(detection_record["ingest"]["source_path"])
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    if source.suffix.lower() == ".pdf":
        return _inject_links_pdf(
            detection_record, output_path=out, target_anchor_indexes=target_anchor_indexes
        )

    linker = DocxLinker(source, out)
    # Where each reference is actually *defined* (caption/heading), so the bookmark
    # anchors the data, not the first mention. Reuse a prebuilt index when the
    # orchestrator passed one (avoids a second build per doc).
    if target_anchor_indexes is not None and source.stem in target_anchor_indexes:
        anchor_index = target_anchor_indexes[source.stem]
    else:
        anchor_index = build_anchor_index(
            detection_record["detections"], str(source), is_pdf=False
        )
    declared_anchors: set[str] = set()
    probes: list[dict[str, Any]] = []

    for det in detection_record["detections"]:
        # Re-target an author's pre-existing PLACEHOLDER hyperlink ("ISE" →
        # about:blank) at its resolved sibling, instead of wrapping a new link.
        if "existing_hl_index" in det:
            rtd = det.get("resolved_target_doc")
            if not rtd:
                continue  # couldn't resolve (e.g. self-reference) → leave placeholder
            rt = Path(rtd)
            if rt.stem == source.stem:
                continue
            sub = _cross_doc_subanchor(det, rt, source, target_anchor_indexes)
            cross = f"{rt.stem}_linked{rt.suffix}" + sub
            linker.retarget_existing_hyperlink(
                det["paragraph_index"], det["existing_hl_index"], url=cross
            )
            probes.append(
                {
                    "source_doc": source.name,
                    "link_text": det["text"],
                    "location_descriptor": f"p{det['paragraph_index']}.hl{det['existing_hl_index']}",
                    "kind": LinkKind.EXTERNAL_URL.value,
                    "target": cross,
                    "target_doc": f"{rt.stem}_linked{rt.suffix}",
                    "detected_by": det.get("source_layer"),
                    "ner_pattern": None,
                    "llm_called": False,
                    "llm_confidence_before": None,
                    "llm_confidence_after": None,
                }
            )
            continue

        location = RunLocation(
            paragraph_index=det["paragraph_index"],
            run_index=det["run_index"],
            char_start=det["char_start"],
            char_end=det["char_end"],
            end_run_index=det.get("end_run_index"),
            end_char=det.get("end_char"),
        )
        kind, target = _resolve_target(det)

        # REF_CITE never-skip: when the exact bibliography entry isn't indexed,
        # fall back to the References heading so a citation still links (team's
        # "by not skipping" requirement) instead of self-anchoring at the citation.
        if (
            det.get("label") == "REF_CITE"
            and target not in anchor_index
            and REF_HEADING_KEY in anchor_index
        ):
            target = REF_HEADING_KEY

        resolved_target_doc = det.get("resolved_target_doc")
        cross_doc_name = ""
        if resolved_target_doc:
            rt = Path(resolved_target_doc)
            if rt.stem != source.stem:
                cross_doc_name = f"{rt.stem}_linked{rt.suffix}"

        # A document-type reference ("the protocol") is only a link when its
        # sibling document is in this batch; otherwise drop it (no broken anchor).
        if det.get("label") in ("DOC_REF", "DOC_ID") and not cross_doc_name:
            # Exception: a highlighted SELF-reference ("SAP" inside SAP.docx) the
            # author marked → link it to THIS document's own top (an internal
            # bookmark at the title paragraph), so it reads consistently next to the
            # adjacent cross-ref. Benign self-navigation, never a broken anchor.
            if det.get("self_ref_top"):
                _SELF = "doc_self_top"
                linker.add_internal_link(location, anchor=_SELF)
                if _SELF not in declared_anchors:
                    declared_anchors.add(_SELF)
                    linker.add_bookmark(
                        RunLocation(paragraph_index=0, run_index=0, char_start=0, char_end=1),
                        _SELF,
                    )
                probes.append(
                    {
                        "source_doc": source.name,
                        "link_text": det["text"],
                        "location_descriptor": (
                            f"p{det['paragraph_index']}.r{det['run_index']}"
                            f":c{det['char_start']}-{det['char_end']}"
                        ),
                        "kind": LinkKind.INTERNAL_BOOKMARK.value,
                        "target": _SELF,
                        "target_doc": str(out),
                        "detected_by": det.get("source_layer"),
                        "ner_pattern": None,
                        "llm_called": False,
                        "llm_confidence_before": None,
                        "llm_confidence_after": None,
                    }
                )
            continue

        if cross_doc_name:
            sub = _cross_doc_subanchor(det, rt, source, target_anchor_indexes)
            linker.add_external_link(location, url=cross_doc_name + sub)
            kind = LinkKind.EXTERNAL_URL
            target = cross_doc_name + sub
        elif kind == LinkKind.EXTERNAL_URL:
            linker.add_external_link(location, url=target)
        else:
            # Issue 5: the actual Table/Section CAPTION (the definition line) must
            # NOT be a clickable link — only citations elsewhere link to it. When
            # this occurrence sits on its own definition paragraph, declare the
            # bookmark there and skip both the hyperlink and the probe.
            loc = anchor_index.get(target)
            def_para = loc.get("paragraph_index") if loc else None
            if def_para is not None and def_para == det["paragraph_index"]:
                if target not in declared_anchors:
                    declared_anchors.add(target)
                    linker.add_bookmark(
                        RunLocation(
                            paragraph_index=def_para, run_index=0, char_start=0, char_end=1
                        ),
                        target,
                    )
                continue  # caption is the link TARGET, never the link source
            linker.add_internal_link(location, anchor=target)
            if target not in declared_anchors:
                declared_anchors.add(target)
                # PLAN TEN: anchor at the reference's DEFINITION (caption/heading)
                # when we found one, else fall back to the citation (legacy).
                if def_para is not None:
                    def_loc = RunLocation(
                        paragraph_index=def_para,
                        run_index=0,
                        char_start=0,
                        # Bookmarks are paragraph-level: _inject_bookmark anchors at
                        # the head of the paragraph and ignores the char span. The
                        # span only has to satisfy RunLocation's invariant (char_end
                        # > 0); 1 is the minimal valid value. (Was 0 → ValidationError
                        # that aborted injection for every docx with a found anchor.)
                        char_end=1,
                    )
                    linker.add_bookmark(def_loc, target)
                else:
                    linker.add_bookmark(location, target)
        probes.append(
            {
                "source_doc": source.name,
                "link_text": det["text"],
                "location_descriptor": (
                    f"p{det['paragraph_index']}.r{det['run_index']}"
                    f":c{det['char_start']}-{det['char_end']}"
                ),
                "kind": kind.value,
                "target": target,
                "target_doc": (
                    cross_doc_name
                    or (str(out) if kind == LinkKind.INTERNAL_BOOKMARK else source.name)
                ),
                "detected_by": det.get("source_layer"),
                "ner_pattern": det.get("pattern_id") if det.get("source_layer") == "ner" else None,
                "llm_called": det.get("llm_consulted", False) or det.get("source_layer") == "llm",
                "llm_confidence_before": det.get("llm_confidence_before"),
                "llm_confidence_after": det.get("llm_confidence_after"),
            }
        )

    # Cross-doc bookmark provisioning: declare a bookmark at the DEFINITION of every
    # anchor a *sibling* document links INTO this one, even when this document never
    # cites it itself. Without this the cross-doc "file.docx#section_ref_6_1" link
    # opens the file but can't scroll — the bookmark would not exist. (Protocol
    # defines §6.1 but never cites it; CSR/ISE/ISS/SAP all link to it.)
    if incoming_anchor_keys:
        for key in incoming_anchor_keys:
            if key in declared_anchors:
                continue
            loc = anchor_index.get(key)
            def_para = loc.get("paragraph_index") if loc else None
            if def_para is None:
                continue
            declared_anchors.add(key)
            linker.add_bookmark(
                RunLocation(
                    paragraph_index=def_para, run_index=0, char_start=0, char_end=1
                ),
                key,
            )

    linker.save()
    return {
        "ingest": detection_record["ingest"],
        "output_path": str(out),
        "probes": probes,
    }


def _inject_links_pdf(
    detection_record: dict[str, Any],
    *,
    output_path: Path,
    target_anchor_indexes: dict[str, dict[str, Any]] | None = None,
) -> dict[str, Any]:
    """PDF-specific injection: use PdfLinker with PdfLocation (page + bbox)."""
    from hyperlink_engine.core.injection.anchor_index import build_anchor_index
    from hyperlink_engine.core.injection.pdf_linker import PdfLinker
    from hyperlink_engine.core.injection.ref_index import REF_HEADING_KEY
    from hyperlink_engine.models import PdfLocation

    source = Path(detection_record["ingest"]["source_path"])
    out = output_path

    linker = PdfLinker(source, out)
    # PLAN TEN: definition (caption/heading) location per reference key, so an
    # internal GOTO lands on the real Table/Section page, not the citation page.
    if target_anchor_indexes is not None and source.stem in target_anchor_indexes:
        anchor_index = target_anchor_indexes[source.stem]
    else:
        anchor_index = build_anchor_index(
            detection_record["detections"], str(source), is_pdf=True
        )
    declared_anchors: set[str] = set()
    probes: list[dict[str, Any]] = []

    for det in detection_record["detections"]:
        bbox = det.get("bbox", [0.0, 0.0, 0.0, 0.0])
        location = PdfLocation(
            page_index=det["page_index"],
            x0=bbox[0],
            y0=bbox[1],
            x1=bbox[2],
            y1=bbox[3],
        )
        kind, target = _resolve_target(det)

        # REF_CITE never-skip: fall back to the References heading when the exact
        # bibliography entry isn't indexed (team's "by not skipping" requirement).
        if (
            det.get("label") == "REF_CITE"
            and target not in anchor_index
            and REF_HEADING_KEY in anchor_index
        ):
            target = REF_HEADING_KEY
        # The reference's canonical anchor key (e.g. "table_ref_14_2_1_1" /
        # "section_ref_6_3"), captured BEFORE a cross-doc target overwrites it, so
        # we can look the reference up in the *target* document's anchor index.
        ref_key = target

        resolved_target_doc = det.get("resolved_target_doc")
        cross_doc_name = ""
        if resolved_target_doc:
            rt = Path(resolved_target_doc)
            if rt.stem != source.stem:
                cross_doc_name = f"{rt.stem}_linked{rt.suffix}"

        # A document-type reference ("the protocol") is only a link when its
        # sibling document is in this batch; otherwise drop it (no broken anchor).
        if det.get("label") in ("DOC_REF", "DOC_ID") and not cross_doc_name:
            continue

        # PLAN TWELVE: 1-based page of the reference's *definition* in the document
        # the link opens, so the UI can open that PDF in a new tab at the right
        # page (#page=N). None ⇒ open at page 1 (whole-document reference).
        target_page: int | None = None

        display = det.get("text") or None
        if cross_doc_name:
            linker.add_external_link(location, url=cross_doc_name, display_text=display)
            kind = LinkKind.EXTERNAL_URL
            target = cross_doc_name
            # The page lives in the *sibling* document's anchor index, keyed by the
            # reference's canonical key. A "the protocol" (DOC_REF) cite has no
            # specific key → stays None → the UI opens that PDF at page 1.
            tloc = (target_anchor_indexes or {}).get(rt.stem, {}).get(ref_key)
            if tloc and tloc.get("page_index") is not None:
                target_page = int(tloc["page_index"]) + 1
        elif kind == LinkKind.EXTERNAL_URL:
            linker.add_external_link(location, url=target, display_text=display)
        else:
            if target not in declared_anchors:
                declared_anchors.add(target)
                loc = anchor_index.get(target)
                if loc and loc.get("page_index") is not None:
                    bbox = loc.get("bbox")
                    if bbox and len(bbox) >= 2:
                        linker.declare_named_destination(
                            target, int(loc["page_index"]), float(bbox[0]), float(bbox[1])
                        )
                    else:
                        linker.declare_named_destination(target, int(loc["page_index"]))
                else:
                    linker.declare_named_destination(target, det["page_index"])
            linker.add_internal_link(location, anchor=target, display_text=display)
            # Internal link → the definition page lives in THIS document's index;
            # fall back to the citation's own page when no definition was found.
            iloc = anchor_index.get(target)
            target_page = (
                int(iloc["page_index"]) + 1
                if iloc and iloc.get("page_index") is not None
                else int(det["page_index"]) + 1
            )

        probes.append(
            {
                "source_doc": source.name,
                "link_text": det["text"],
                "location_descriptor": (
                    f"page{det['page_index']}.span{det.get('span_index', 0)}"
                    f":c{det['char_start']}-{det['char_end']}"
                ),
                "kind": kind.value,
                "target": target,
                "target_doc": (
                    cross_doc_name
                    or (str(out) if kind == LinkKind.INTERNAL_BOOKMARK else source.name)
                ),
                "target_page": target_page,
                "detected_by": det.get("source_layer"),
                "ner_pattern": det.get("pattern_id") if det.get("source_layer") == "ner" else None,
                "llm_called": det.get("llm_consulted", False) or det.get("source_layer") == "llm",
                "llm_confidence_before": det.get("llm_confidence_before"),
                "llm_confidence_after": det.get("llm_confidence_after"),
            }
        )

    linker.save()
    return {
        "ingest": detection_record["ingest"],
        "output_path": str(out),
        "probes": probes,
    }


def _cross_doc_subanchor(
    det: dict[str, Any],
    target_path: Path,
    source_path: Path,
    indexes: dict[str, dict[str, Any]] | None,
) -> str:
    """Sub-anchor ("#key") for a cross-document link when the *target* document
    defines that reference (PLAN TEN Step 4).

    Word resolves ``file.docx#bookmark`` only for a docx->docx pairing; for any
    other format pairing (pdf targets, cross-format) we open the file at document
    level — a viewer limitation, not a code gap — so we return no fragment.
    """
    if not indexes:
        return ""
    if target_path.suffix.lower() != ".docx" or source_path.suffix.lower() != ".docx":
        return ""
    tgt = indexes.get(target_path.stem)
    if not tgt:
        return ""
    _kind, key = _resolve_target(det)
    return f"#{key}" if key in tgt else ""


def _resolve_target(det: dict[str, Any]) -> tuple[LinkKind, str]:
    """Phase-1 heuristic mirrored from scripts/phase1_acceptance.py."""
    label = det["label"]
    groups = det.get("groups", {})
    text = det["text"]
    if label == "URL":
        # Already a full website URL — link straight to it.
        return LinkKind.EXTERNAL_URL, text
    if label == "EXT_REF":
        # External standard / regulatory citation → stable public URL, mirroring
        # the NCT → clinicaltrials.gov mapping below. Unknown EXT_REF subtypes
        # fall back to linking the matched text verbatim.
        pid = det.get("pattern_id", "")
        if pid == "EXT_REF_DOI_V1":
            return LinkKind.EXTERNAL_URL, f"https://doi.org/{text}"
        if pid == "EXT_REF_HELSINKI_V1":
            return (
                LinkKind.EXTERNAL_URL,
                "https://www.wma.net/policies-post/wma-declaration-of-helsinki/",
            )
        if pid == "EXT_REF_ICH_V1":
            # ICH has no deterministic deep-link per guideline code; link to the
            # official guidelines index (a valid, non-broken landing page).
            return LinkKind.EXTERNAL_URL, "https://www.ich.org/page/ich-guidelines"
        if pid == "EXT_REF_CFR_V1":
            title = groups.get("title", "21")
            part = groups.get("part", "")
            return (
                LinkKind.EXTERNAL_URL,
                f"https://www.ecfr.gov/current/title-{title}/part-{part}",
            )
        return LinkKind.EXTERNAL_URL, text
    if label in ("DOC_REF", "DOC_ID"):
        # A document-type reference ("the protocol") or document id ("TMX-67_301").
        # The real target is the sibling document, wired by the resolver via
        # ``resolved_target_doc``; unresolved ones are skipped before injection.
        # When this piece is part of a compound phrase that also names a section
        # ("Protocol TMX-67_301 Section 6.1"), inherit that section's anchor so the
        # whole phrase links to the SAME spot (Protocol_linked.docx#section_ref_6_1).
        compound_key = groups.get("compound_anchor_key")
        if compound_key:
            return LinkKind.INTERNAL_BOOKMARK, compound_key
        return LinkKind.INTERNAL_BOOKMARK, "doc_ref" if label == "DOC_REF" else "doc_id"
    if label == "REF_CITE":
        # A literature citation → its entry in this document's References section.
        # Author-year keys on (surname, year); numbered keys on the bracket number.
        # The anchor index maps the key to the entry paragraph/page; if the exact
        # entry isn't indexed the inject step retargets to the References heading
        # (ref_heading) so the citation is never dropped.
        from hyperlink_engine.core.injection.ref_index import (
            REF_HEADING_KEY,
            canonical_ref_key_author,
            canonical_ref_key_num,
        )

        if groups.get("surname") and groups.get("year"):
            return LinkKind.INTERNAL_BOOKMARK, canonical_ref_key_author(
                groups["surname"], groups["year"]
            )
        if groups.get("num"):
            return LinkKind.INTERNAL_BOOKMARK, canonical_ref_key_num(groups["num"])
        return LinkKind.INTERNAL_BOOKMARK, REF_HEADING_KEY
    if label in {"SECTION_REF", "TABLE_REF", "FIGURE_REF", "LISTING_REF", "APPENDIX_REF"}:
        from hyperlink_engine.core.injection.anchor_index import canonical_anchor_key

        num = groups.get("num") or text
        return LinkKind.INTERNAL_BOOKMARK, canonical_anchor_key(label, num)
    if label == "VISIT_REF":
        # "Week 2 Visit" → visit_ref_week_2; the anchor index maps that to the
        # matching visit section ("9.3.3.1 Week 2"). Unresolved visits (e.g. a
        # "Day 1" with no Day-1 heading) fall back to the citation anchor.
        from hyperlink_engine.core.injection.anchor_index import canonical_visit_key

        return LinkKind.INTERNAL_BOOKMARK, canonical_visit_key(
            groups.get("unit", ""), groups.get("n", "")
        )
    if label == "STUDY_ID" and det["pattern_id"] == "STUDY_ID_NCT_V1":
        return LinkKind.EXTERNAL_URL, f"https://clinicaltrials.gov/study/{text}"
    if label == "STUDY_ID":
        return LinkKind.INTERNAL_BOOKMARK, f"study_{text.replace('-', '_')}"
    if label == "CTD_LEAF":
        mod = groups.get("mod", "?")
        sub = groups.get("sub", "") or groups.get("subpath", "")
        if sub:
            return (
                LinkKind.INTERNAL_BOOKMARK,
                f"m{mod}_" + sub.replace(".", "_").replace("/", "_"),
            )
        return LinkKind.INTERNAL_BOOKMARK, f"m{mod}"
    return LinkKind.INTERNAL_BOOKMARK, text


# ─────────────────────────────────────────────────────────────────────────
# Stage 4 — validation
# ─────────────────────────────────────────────────────────────────────────


def validate_links(injection_record: dict[str, Any]) -> dict[str, Any]:
    """Stage 4 — run existence checks against the linked output."""
    probes: list[LinkProbe] = []
    for p in injection_record["probes"]:
        probes.append(
            LinkProbe(
                source_doc=p["source_doc"],
                link_text=p["link_text"],
                location_descriptor=p["location_descriptor"],
                kind=LinkKind(p["kind"]),
                target=p["target"],
                target_doc=Path(p["target_doc"]) if p.get("target_doc") else None,
                detected_by=p.get("detected_by"),
                ner_pattern=p.get("ner_pattern"),
                llm_called=p.get("llm_called", False),
                llm_confidence_before=p.get("llm_confidence_before"),
                llm_confidence_after=p.get("llm_confidence_after"),
            )
        )
    records = check_all(probes)
    return {
        "ingest": injection_record["ingest"],
        "output_path": injection_record["output_path"],
        "link_records": [r.model_dump(mode="json") for r in records],
    }


# ─────────────────────────────────────────────────────────────────────────
# Stage 5 — reporting (per-document CSV; the batch runner aggregates)
# ─────────────────────────────────────────────────────────────────────────


def write_per_doc_report(
    validation_record: dict[str, Any],
    *,
    output_path: str,
) -> dict[str, Any]:
    """Stage 5 — write a per-document CSV report and return its path."""
    records = [LinkRecord.model_validate(r) for r in validation_record["link_records"]]
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    write_link_records(records, out)
    return {
        "ingest": validation_record["ingest"],
        "output_path": validation_record["output_path"],
        "report_path": str(out),
        "link_records": validation_record["link_records"],
    }


# ─────────────────────────────────────────────────────────────────────────
# Single-doc orchestration (sync) — convenient for unit tests + the W1.5
# spike. The batch runner uses this in a thread pool for parallelism in
# eager mode, and switches to Celery .apply_async() in production.
# ─────────────────────────────────────────────────────────────────────────


@dataclass
class DocPipelineResult:
    source_path: Path
    output_path: Path
    report_path: Path
    detection_count: int = 0
    link_records: list[LinkRecord] = field(default_factory=list)
    duration_seconds: float = 0.0

    @property
    def ok_count(self) -> int:
        return sum(1 for r in self.link_records if r.status == LinkStatus.OK)

    @property
    def broken_count(self) -> int:
        return sum(1 for r in self.link_records if r.status == LinkStatus.BROKEN)

    @property
    def total_links(self) -> int:
        return len(self.link_records)


def process_document(
    source_path: Path,
    *,
    output_path: Path,
    report_path: Path,
    extractor_config: ExtractorConfig | None = None,
) -> DocPipelineResult:
    """Run all five stages synchronously for one document. Always idempotent."""
    started = time.perf_counter()
    ingest = ingest_document(str(source_path))
    detection = detect_references(
        ingest,
        extractor_config=(extractor_config or ExtractorConfig()).__dict__,
    )
    injection = inject_links(detection, output_path=str(output_path))
    validation = validate_links(injection)
    report = write_per_doc_report(validation, output_path=str(report_path))
    elapsed = time.perf_counter() - started
    link_records = [LinkRecord.model_validate(r) for r in report["link_records"]]
    _log.info(
        "process_document_complete",
        source=str(source_path),
        detections=len(detection["detections"]),
        links=len(link_records),
        duration_s=round(elapsed, 3),
    )
    return DocPipelineResult(
        source_path=source_path,
        output_path=output_path,
        report_path=report_path,
        detection_count=len(detection["detections"]),
        link_records=link_records,
        duration_seconds=elapsed,
    )


# ─────────────────────────────────────────────────────────────────────────
# Celery task registration
# ─────────────────────────────────────────────────────────────────────────


_registered_tasks: dict[str, Callable[..., Any]] = {}


def register_celery_tasks() -> dict[str, Any]:
    """Decorate the stage functions as Celery tasks and return them.

    Called once at import time when Celery is available; safe to call
    multiple times — the second call returns the same handles.
    """
    if _registered_tasks:
        return dict(_registered_tasks)
    settings = get_settings()
    app = get_app()

    stage_action_pairs: list[tuple[str, str, Callable[..., Any]]] = [
        (PIPELINE_STAGES[0], "ingest_document", ingest_document),
        (PIPELINE_STAGES[1], "detect_references", detect_references),
        (PIPELINE_STAGES[2], "inject_links", inject_links),
        (PIPELINE_STAGES[3], "validate_links", validate_links),
        (PIPELINE_STAGES[4], "write_per_doc_report", write_per_doc_report),
    ]
    for stage, action, fn in stage_action_pairs:
        task_name = stage_task_name(stage, action)
        task = app.task(
            name=task_name,
            bind=False,
            autoretry_for=(Exception,),
            retry_backoff=settings.pipeline_retry_backoff_seconds,
            retry_kwargs={"max_retries": settings.pipeline_max_retries},
            acks_late=True,
        )(fn)
        _registered_tasks[task_name] = task
    _log.info("celery_tasks_registered", count=len(_registered_tasks))
    return dict(_registered_tasks)


def get_task(stage: str, action: str) -> Callable[..., Any]:
    """Look up a registered task by stage + action. Auto-registers if needed."""
    if not _registered_tasks:
        register_celery_tasks()
    name = stage_task_name(stage, action)
    if name not in _registered_tasks:
        raise KeyError(f"task {name!r} not registered")
    return _registered_tasks[name]
