"""Unit tests for ingestion/docx_loader.py."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from hyperlink_engine.core.ingestion.docx_loader import DocxLoadError, load_docx


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "loader_sample.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("See Section 2.5.3 for details.")
    doc.save(str(path))
    return path


def test_load_returns_provenance(sample_docx: Path) -> None:
    document, provenance = load_docx(sample_docx)
    assert document is not None
    assert provenance.source_path == sample_docx
    assert provenance.file_size_bytes > 0
    assert len(provenance.sha256) == 64
    int(provenance.sha256, 16)  # must be a valid hex digest


def test_sha256_is_deterministic(sample_docx: Path) -> None:
    _, p1 = load_docx(sample_docx)
    _, p2 = load_docx(sample_docx)
    assert p1.sha256 == p2.sha256


def test_load_rejects_missing_file(tmp_path: Path) -> None:
    with pytest.raises(DocxLoadError, match="does not exist"):
        load_docx(tmp_path / "nope.docx")


def test_load_rejects_non_ooxml(tmp_path: Path) -> None:
    bogus = tmp_path / "bogus.docx"
    bogus.write_bytes(b"not a zip at all")
    with pytest.raises(DocxLoadError, match="not a valid .docx"):
        load_docx(bogus)


def test_load_rejects_directory(tmp_path: Path) -> None:
    with pytest.raises(DocxLoadError, match="is not a file"):
        load_docx(tmp_path)


def test_load_does_not_mutate_source(sample_docx: Path) -> None:
    before = sample_docx.read_bytes()
    load_docx(sample_docx)
    after = sample_docx.read_bytes()
    assert before == after
