"""Regression: injecting into a .docx whose anchor index has a *definition*
(a numbered Heading) must NOT raise and must still produce hyperlinks.

Root cause guarded here: the PLAN TEN "anchor at the definition" branch in
``inject_links`` built a ``RunLocation`` for the bookmark with ``char_end=0``.
``RunLocation`` requires ``char_end > 0`` (a link span covers ≥1 char), so the
construction raised ``ValidationError`` — which ``node_inject_links`` swallowed
as ``injection_failed`` and fell back to copying the original. The visible
symptom was a run reporting thousands of *resolved* references but **0 links
injected** across every Word document (e.g. run ``1fa4bf0d`` on
``csr_dossier_large``). Bookmarks are paragraph-level (``_inject_bookmark``
ignores the char span), so the fix is simply a valid span (``char_end=1``).
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from hyperlink_engine.workers.tasks import inject_links


def _make_docx_with_heading(path: Path) -> None:
    """A doc whose Heading 'Section 2.5' is the definition the citation targets."""
    doc = Document()
    doc.add_heading("2.5 Statistical Methods", level=1)  # the DEFINITION (para 0)
    doc.add_paragraph("Refer to Section 2.5 for details.")  # the CITATION (para 1)
    doc.save(str(path))


def test_inject_docx_with_indexed_definition_does_not_raise(tmp_path: Path) -> None:
    src = tmp_path / "csr-body.docx"
    _make_docx_with_heading(src)

    citation = "Refer to Section 2.5 for details."
    start = citation.index("Section 2.5")
    det = {
        "label": "SECTION_REF",
        "text": "Section 2.5",
        "groups": {"num": "2.5"},
        "paragraph_index": 1,
        "run_index": 0,
        "char_start": start,
        "char_end": start + len("Section 2.5"),
        "source_layer": "regex",
        "context": citation,
    }
    rec = {"ingest": {"source_path": str(src)}, "detections": [det]}

    out = tmp_path / "csr-body_linked.docx"
    # Before the fix this raised pydantic ValidationError (char_end=0) → 0 links.
    result = inject_links(rec, output_path=str(out))

    assert len(result["probes"]) == 1
    probe = result["probes"][0]
    assert probe["kind"] == "internal_bookmark"
    assert probe["target"] == "section_ref_2_5"

    # The output really contains the injected hyperlink + the definition bookmark.
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert xml.count("<w:hyperlink") == 1
    assert "<w:bookmarkStart" in xml


def _make_docx_with_table_caption(path: Path) -> None:
    """A doc with a Table caption (the definition) plus a prose citation to it."""
    doc = Document()
    doc.add_paragraph("Table 1.1: Subject Demographics")  # the CAPTION (para 0)
    doc.add_paragraph("Baseline data are shown in Table 1.1.")  # the CITATION (para 1)
    doc.save(str(path))


def test_inject_does_not_hyperlink_the_table_caption(tmp_path: Path) -> None:
    """Issue 5: the actual table caption number is the bookmark TARGET, not a
    clickable link; only the prose citation links to it."""
    src = tmp_path / "csr-body.docx"
    _make_docx_with_table_caption(src)

    caption = "Table 1.1: Subject Demographics"
    citation = "Baseline data are shown in Table 1.1."

    def _det(para: int, line: str) -> dict:
        start = line.index("Table 1.1")
        return {
            "label": "TABLE_REF",
            "text": "Table 1.1",
            "groups": {"num": "1.1"},
            "paragraph_index": para,
            "run_index": 0,
            "char_start": start,
            "char_end": start + len("Table 1.1"),
            "source_layer": "regex",
            "context": line,
        }

    rec = {
        "ingest": {"source_path": str(src)},
        "detections": [_det(0, caption), _det(1, citation)],  # caption first, then citation
    }
    out = tmp_path / "csr-body_linked.docx"
    result = inject_links(rec, output_path=str(out))

    # Only the citation becomes a link/probe — the caption is skipped.
    assert len(result["probes"]) == 1
    assert result["probes"][0]["location_descriptor"].startswith("p1.")

    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8", "ignore")
    assert xml.count("<w:hyperlink") == 1  # citation only, NOT the caption
    assert "<w:bookmarkStart" in xml  # the caption is still a bookmark target
