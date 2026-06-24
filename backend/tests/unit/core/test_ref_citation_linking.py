"""Literature-citation linking — author-year + numbered → References entries.

Grounded in the team's real CSR.docx, which cites three references author-year
("Helget LN, 2024", "Becker MA, 2005", "Schumacher HR, 2010") and lists them in a
"15. REFERENCES" section. The tricky bits this locks in:

* a citation split across Word runs ("Helget" | " LN, 2024") is still detected;
* surname alone is ambiguous — "Schumacher" is the 2010 first author *and* a 2005
  co-author — so matching uses ``(surname, year)``;
* a References entry's "Journal. 2005" tail is NOT mistaken for a citation
  (period guard); month names are rejected (date noise);
* a citation whose entry can't be found falls back to the References heading
  rather than being dropped ("by not skipping").
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from hyperlink_engine.core.injection.anchor_index import build_anchor_index
from hyperlink_engine.core.injection.ref_index import (
    author_year_cites,
    canonical_ref_key_author,
    canonical_ref_key_num,
    is_next_section_after_refs,
    is_references_heading,
    numbered_cites,
    parse_ref_entry_key,
)
from hyperlink_engine.workers.tasks import detect_references, inject_links, _resolve_target


# ── ref_index matchers ───────────────────────────────────────────────────────


def _ay(text: str) -> list[tuple[str, str]]:
    return [(m.group("surname"), m.group("year")) for m in author_year_cites(text)]


def test_author_year_detects_real_citations() -> None:
    assert _ay("active-controlled trial comparing febuxostat see Helget LN, 2024") == [
        ("Helget", "2024")
    ]
    assert _ay("reference Becker MA, 2005 and Schumacher HR, 2010") == [
        ("Becker", "2005"),
        ("Schumacher", "2010"),
    ]


def test_author_year_detects_comma_after_surname_format() -> None:
    """The Afatinib CSR writes citations with the initials AFTER the comma
    ("Ullas, B 2022") rather than before it ("Helget LN, 2024"). Both styles
    must be detected, and resolve on (surname, year)."""
    assert _ay("Primarily used in NSCLC See Ullas, B 2022") == [("Ullas", "2022")]
    assert _ay("Route: Oral (tablet) See Tankere, P 2022.") == [("Tankere", "2022")]
    assert _ay("including see Chiu, CH 2015") == [("Chiu", "2015")]


@pytest.mark.parametrize(
    "text",
    [
        "the study ran from February 2017 to August 2018",  # month-year date noise
        "published in N Engl J Med. 2005;353:2450",  # References-entry journal tail (period)
        "Am J Kidney Dis. 2024",  # period before year, not a citation
    ],
)
def test_author_year_rejects_noise(text: str) -> None:
    assert _ay(text) == [], f"false-positive author-year in: {text!r}"


def test_pdf_citations_skipped_inside_references_section(tmp_path: Path) -> None:
    """A name inside the References list must NOT become a link (it is the link
    TARGET); only a mention elsewhere (the body) links and redirects there. The
    docx path already gated this; this locks the PDF path."""
    import fitz

    from hyperlink_engine.workers.tasks import _detect_references_pdf

    pdf = tmp_path / "cite.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "Refer to Smith J 2020 for the analysis.", fontsize=12)
    page.insert_text((72, 150), "6 References", fontsize=12)
    page.insert_text((72, 180), "Smith J 2020 A study of things. Journal of Tests.", fontsize=12)
    doc.save(str(pdf))
    doc.close()

    rec = _detect_references_pdf({"source_path": str(pdf)})
    cites = [d for d in rec["detections"] if d.get("label") == "REF_CITE"]
    # exactly one — the BODY citation; the References-list "Smith J 2020" is skipped
    assert len(cites) == 1
    assert cites[0]["groups"]["surname"] == "Smith"


def test_pdf_reference_entry_has_bbox_for_pinpoint_scroll(tmp_path: Path) -> None:
    """A PDF citation must land ON the References entry line, not the page top — so
    the entry anchor must carry a bbox (was None → page-level only, the 'doesn't
    scroll to the reference' bug)."""
    import fitz

    from hyperlink_engine.core.injection.anchor_index import build_anchor_index
    from hyperlink_engine.workers.tasks import detect_references

    pdf = tmp_path / "r.pdf"
    doc = fitz.open()
    doc.new_page().insert_text((72, 100), "see Chiu, CH 2015.", fontsize=11)
    p2 = doc.new_page()
    p2.insert_text((72, 90), "6 References", fontsize=12)
    p2.insert_text((72, 300), "Chiu, C.H.; Yang. Treatment. J. Thorac. Oncol. 2015, 10, 793-799.", fontsize=10)
    doc.save(str(pdf))
    doc.close()

    rec = detect_references({"source_path": str(pdf)})
    idx = build_anchor_index(rec["detections"], pdf, is_pdf=True)
    assert idx["ref_chiu_2015"]["page_index"] == 1          # references page
    assert idx["ref_chiu_2015"]["bbox"] is not None         # pinpoint, not page-top
    assert idx["ref_chiu_2015"]["bbox"][1] > 250            # the entry's actual Y


def test_numbered_cites_need_brackets() -> None:
    assert [m.group("num") for m in numbered_cites("as shown in [7] and [12]")] == ["7", "12"]
    # A numbered References *entry* ("7. …") must not be mistaken for a citation.
    assert numbered_cites("7. Marcus R, et al. Biometrika 1976") == []


# ── References-section parsing ───────────────────────────────────────────────


def test_references_heading_and_terminator() -> None:
    assert is_references_heading("15. REFERENCES")
    assert is_references_heading("References")
    assert not is_references_heading("see the references section for details")
    assert is_next_section_after_refs("16. APPENDICES")


def test_parse_entry_key_author_and_numbered() -> None:
    assert parse_ref_entry_key("Helget LN et al. Efficacy … Am J Kidney Dis. 2024") == "ref_helget_2024"
    assert parse_ref_entry_key("7. Marcus R, et al. Biometrika 1976") == "ref_7"
    assert parse_ref_entry_key("   ") is None


def test_citation_and_entry_keys_agree() -> None:
    """The whole scheme hinges on citation key == entry key."""
    assert canonical_ref_key_author("Helget", "2024") == "ref_helget_2024"
    assert canonical_ref_key_num("07") == "ref_7"
    # Disambiguation: Schumacher appears in both entries; the year separates them.
    becker_entry = "Becker MA, Schumacher HR, Wortmann RL, et al. Foo. N Engl J Med. 2005;353"
    schu_entry = "Schumacher HR et al. Confirms trial. Arthritis Res Ther. 2010"
    assert parse_ref_entry_key(becker_entry) == "ref_becker_2005"
    assert parse_ref_entry_key(schu_entry) == "ref_schumacher_2010"


# ── end-to-end docx detect → resolve → inject ────────────────────────────────


def _ref_hyperlinks(docx_path: Path) -> list[tuple[int, str, str]]:
    """(paragraph_index, anchor, visible_text) for each injected ref_* hyperlink."""
    doc = Document(str(docx_path))
    out: list[tuple[int, str, str]] = []
    for p_i, p in enumerate(doc.paragraphs):
        for h in p._p.findall(qn("w:hyperlink")):
            anc = h.get(qn("w:anchor"))
            if anc and anc.startswith("ref_"):
                txt = "".join(t.text or "" for t in h.findall(".//" + qn("w:t")))
                out.append((p_i, anc, txt))
    return out


def _bookmarks(docx_path: Path) -> set[str]:
    doc = Document(str(docx_path))
    return {bm.get(qn("w:name")) for bm in doc.element.body.iter(qn("w:bookmarkStart"))}


def _run(make_doc, tmp_path: Path) -> Path:
    src = tmp_path / "src.docx"
    make_doc(src)
    rec = detect_references({"source_path": str(src)})
    out = tmp_path / "src_linked.docx"
    inject_links(rec, output_path=str(out))
    return out


def test_split_run_author_year_links_to_its_entry(tmp_path: Path) -> None:
    """"Helget" + " LN, 2024" in separate runs still links to the Helget entry."""

    def make(path: Path) -> None:
        d = Document()
        p = d.add_paragraph()
        p.add_run("The trial compared febuxostat and allopurinol see ")
        p.add_run("Helget")  # surname in its own run …
        p.add_run(" LN, 2024")  # … year in the next run
        d.add_paragraph("15. REFERENCES")
        d.add_paragraph("Helget LN et al. Efficacy and safety. Am J Kidney Dis. 2024")
        d.save(str(path))

    out = _run(make, tmp_path)
    links = _ref_hyperlinks(out)
    assert ("ref_helget_2024" in {a for _, a, _ in links})
    # the anchor lands on the surname run, and the entry bookmark exists
    p_i, anc, txt = next(l for l in links if l[1] == "ref_helget_2024")
    assert txt == "Helget"
    assert "ref_helget_2024" in _bookmarks(out)


def test_surname_year_disambiguates_coauthor(tmp_path: Path) -> None:
    """"Schumacher HR, 2010" must hit the 2010 entry, not the 2005 entry it co-authors."""

    def make(path: Path) -> None:
        d = Document()
        d.add_paragraph("Risk was similar see Becker MA, 2005 and Schumacher HR, 2010")
        d.add_paragraph("15. REFERENCES")
        d.add_paragraph("Becker MA, Schumacher HR, et al. Foo. N Engl J Med. 2005;353")  # entry A
        d.add_paragraph("Schumacher HR et al. Bar. Arthritis Res Ther. 2010")  # entry B
        d.save(str(path))

    out = _run(make, tmp_path)
    idx = build_anchor_index(
        detect_references({"source_path": str(tmp_path / "src.docx")})["detections"],
        str(tmp_path / "src.docx"),
        is_pdf=False,
    )
    # entry A is paragraph 2, entry B is paragraph 3
    assert idx["ref_becker_2005"]["paragraph_index"] == 2
    assert idx["ref_schumacher_2010"]["paragraph_index"] == 3
    anchors = {a for _, a, _ in _ref_hyperlinks(out)}
    assert {"ref_becker_2005", "ref_schumacher_2010"} <= anchors


def _make_numbered_doc(path: Path) -> None:
    d = Document()
    d.add_paragraph("Efficacy was confirmed in prior work [1].")
    d.add_paragraph("15. REFERENCES")
    d.add_paragraph("1. Marcus R, et al. Foo bar baz. Biometrika 1976")
    d.save(str(path))


def test_numbered_citation_links_to_numbered_entry(tmp_path: Path, monkeypatch) -> None:
    # Numbered "[n]" citation linking is OFF by default (publishing team found a bare
    # "[1]" → bibliography confusing); this exercises the still-present feature ON.
    import hyperlink_engine.workers.tasks as _tasks

    monkeypatch.setattr(_tasks, "_LINK_NUMBERED_CITATIONS", True)
    out = _run(_make_numbered_doc, tmp_path)
    assert "ref_1" in {a for _, a, _ in _ref_hyperlinks(out)}
    assert "ref_1" in _bookmarks(out)


def test_numbered_citation_disabled_by_default(tmp_path: Path) -> None:
    """DEFAULT: a bare "[1]" must NOT produce a numbered citation link."""
    out = _run(_make_numbered_doc, tmp_path)
    assert "ref_1" not in {a for _, a, _ in _ref_hyperlinks(out)}


def test_pdf_author_year_citation_detected(tmp_path: Path) -> None:
    """Author-year citations are now detected in the PDF path (was numbered-only)."""
    import fitz

    from hyperlink_engine.workers.tasks import _detect_references_pdf

    pdf = tmp_path / "cite.pdf"
    doc = fitz.open()
    doc.new_page().insert_text(
        (72, 100), "Primarily used in NSCLC See Ullas, B 2022 for details.", fontsize=11
    )
    doc.save(str(pdf))
    doc.close()
    rec = _detect_references_pdf({"source_path": str(pdf)})
    cites = [d for d in rec["detections"] if d.get("pattern_id") == "REF_CITE_AUTHOR_YEAR_V1"]
    assert any(c["groups"] == {"surname": "Ullas", "year": "2022"} for c in cites)


def test_unmatched_citation_falls_back_to_heading_not_skipped(tmp_path: Path) -> None:
    """A citation with no matching entry links to the References heading, never dropped."""

    def make(path: Path) -> None:
        d = Document()
        d.add_paragraph("An older finding see Nobody QZ, 1999 was noted.")
        d.add_paragraph("15. REFERENCES")
        d.add_paragraph("Helget LN et al. Something else. Am J Kidney Dis. 2024")
        d.save(str(path))

    out = _run(make, tmp_path)
    anchors = {a for _, a, _ in _ref_hyperlinks(out)}
    assert "ref_heading" in anchors, "unmatched citation must fall back, not be skipped"
    assert "ref_heading" in _bookmarks(out)


def test_resolve_target_reads_groups() -> None:
    ay = {"label": "REF_CITE", "text": "Helget", "groups": {"surname": "Helget", "year": "2024"}}
    assert _resolve_target(ay)[1] == "ref_helget_2024"
    num = {"label": "REF_CITE", "text": "[7]", "groups": {"num": "7"}}
    assert _resolve_target(num)[1] == "ref_7"
