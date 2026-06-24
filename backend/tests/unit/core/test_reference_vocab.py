"""PLAN TWENTY — deterministic closed-vocabulary citation detection.

The References list is the answer key: extract every author surname (any shape,
incl. short "Xu" / apostrophe "O'Brien") and find those exact names in the body
with their year adjacent. Deterministic — same result every run, no model.
"""

from __future__ import annotations

from pathlib import Path

from hyperlink_engine.core.injection.ref_index import (
    build_reference_vocab,
    reference_vocab_cites,
)


def test_build_vocab_handles_diverse_names() -> None:
    lines = [
        "6 References",
        "Xu, H.; Yang, G. EGFR alterations. Front. Pharmacol. 2022, 13, 976731",
        "O'Brien, M.; Kelly, P. Afatinib in NSCLC. Ann Oncol. 2019, 30, 1245-1252.",
        "Chiu, C.H.; Yang, C.T. Treatment Response. J. Thorac. Oncol. 2015, 10, 793-799.",
    ]
    v = build_reference_vocab(lines)
    assert v["Xu"] == ("2022", "ref_xu_2022")          # short surname regex can't see
    assert v["O'Brien"] == ("2019", "ref_obrien_2019")  # apostrophe kept for search; key stripped
    assert v["Chiu"] == ("2015", "ref_chiu_2015")


def test_build_vocab_empty_without_references() -> None:
    assert build_reference_vocab(["1 Introduction", "Some body text without a References section."]) == {}


def test_vocab_cites_catches_short_and_compound() -> None:
    vocab = {"Xu": ("2022", "ref_xu_2022"), "O'Brien": ("2019", "ref_obrien_2019")}
    body = "See Xu, H 2022 and also O'Brien, M 2019 for details."
    keys = {(c.group("surname"), c.group("year")) for c in reference_vocab_cites(body, vocab)}
    assert ("Xu", "2022") in keys
    assert ("O'Brien", "2019") in keys


def test_vocab_cites_requires_adjacent_year() -> None:
    vocab = {"Li": ("2020", "ref_li_2020")}
    # surname present but NO nearby year → not this citation (avoids stray short-name FPs)
    assert reference_vocab_cites("The Li cohort was studied carefully.", vocab) == []
    # year adjacent → matched
    assert len(reference_vocab_cites("see Li, Q 2020 here", vocab)) == 1


def test_vocab_word_boundary_no_substring_match() -> None:
    vocab = {"Xu": ("2022", "ref_xu_2022")}
    # "Xu" inside "Xuanwu" must NOT match
    assert reference_vocab_cites("the Xuanwu 2022 district hospital", vocab) == []


def test_vocab_mode_links_diverse_name_end_to_end(tmp_path: Path, monkeypatch) -> None:
    """Full docx flow in vocab mode: a body 'Xu, H 2022' (regex can't see it) links
    to its References entry with a destination bookmark."""
    monkeypatch.setattr("hyperlink_engine.core.injection.ref_index._REFERENCE_DETECTOR", "vocab")
    from docx import Document
    from docx.oxml.ns import qn

    from hyperlink_engine.workers.tasks import detect_references, inject_links

    src = tmp_path / "c.docx"
    d = Document()
    d.add_paragraph("Findings see Xu, H 2022 for the cohort.")
    d.add_paragraph("6 References")
    d.add_paragraph("Xu, H.; Yang, G. EGFR uncommon alterations. Front. Pharmacol. 2022, 13, 976731.")
    d.save(str(src))

    rec = detect_references({"source_path": str(src)})
    cites = [x for x in rec["detections"] if x.get("label") == "REF_CITE"]
    assert any(c["groups"].get("surname") == "Xu" for c in cites)

    out = tmp_path / "c_linked.docx"
    res = inject_links(rec, output_path=str(out))
    assert any(p.get("target") == "ref_xu_2022" for p in res.get("probes", []))
    bms = {b.get(qn("w:name")) for b in Document(str(out)).element.iter(qn("w:bookmarkStart"))}
    assert "ref_xu_2022" in bms  # destination bookmark exists → scroll target
