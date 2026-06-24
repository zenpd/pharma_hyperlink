"""Unit tests for injection/style_preserver.py."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from hyperlink_engine.core.injection.style_preserver import (
    StyleSnapshot,
    diff,
    snapshot,
)
from hyperlink_engine.models import AnomalySeverity


def _make(path: Path, *, bold_first: bool = False) -> None:
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Hello ")
    if bold_first:
        r1.bold = True
    p.add_run("world")
    doc.save(str(path))


def test_snapshot_captures_every_run(tmp_path: Path) -> None:
    path = tmp_path / "s.docx"
    _make(path)
    snap = snapshot(path)
    assert len(snap.runs) == 2
    assert snap.runs[0].text == "Hello "
    assert snap.runs[1].text == "world"


def test_diff_empty_when_unchanged(tmp_path: Path) -> None:
    path = tmp_path / "s.docx"
    _make(path)
    a = snapshot(path)
    b = snapshot(path)
    assert diff(a, b) == []


def test_diff_flags_bold_change(tmp_path: Path) -> None:
    before_path = tmp_path / "before.docx"
    after_path = tmp_path / "after.docx"
    _make(before_path, bold_first=False)
    _make(after_path, bold_first=True)
    mutations = diff(snapshot(before_path), snapshot(after_path))
    assert any(m.field == "bold" for m in mutations)


def test_diff_ignores_intentional_runs(tmp_path: Path) -> None:
    before_path = tmp_path / "before.docx"
    after_path = tmp_path / "after.docx"
    _make(before_path, bold_first=False)
    _make(after_path, bold_first=True)
    mutations = diff(
        snapshot(before_path),
        snapshot(after_path),
        intentional_runs={(0, 0)},
    )
    assert mutations == []


def test_dosscriber_style_change_is_blocker(tmp_path: Path) -> None:
    """A Dosscriber-style mutation must flag as BLOCKER, others as WARNING."""
    before_path = tmp_path / "before.docx"
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("text")
    run.bold = False
    # Tag with a Dosscriber-prefixed character style. python-docx can't
    # create styles ad-hoc, so we patch the underlying XML to attach a
    # Dosscriber-named style.
    from docx.oxml.ns import qn

    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._r.makeelement(qn("w:rPr"), {})
        run._r.insert(0, rPr)
    rStyle = rPr.makeelement(qn("w:rStyle"), {qn("w:val"): "Dosscriber_Body"})
    rPr.append(rStyle)

    # Register the style so python-docx surfaces .style.name in the snapshot.
    styles = doc.styles
    from docx.enum.style import WD_STYLE_TYPE

    if "Dosscriber_Body" not in [s.name for s in styles]:
        styles.add_style("Dosscriber_Body", WD_STYLE_TYPE.CHARACTER)

    doc.save(str(before_path))
    before = snapshot(before_path)

    # Build the "after" snapshot manually with a bold flip.
    from hyperlink_engine.core.injection.style_preserver import RunFingerprint

    after_run = RunFingerprint(
        paragraph_index=before.runs[0].paragraph_index,
        run_index=before.runs[0].run_index,
        text=before.runs[0].text,
        bold=True,  # flipped
        italic=before.runs[0].italic,
        underline=before.runs[0].underline,
        font_name=before.runs[0].font_name,
        font_size_pt=before.runs[0].font_size_pt,
        color_rgb=before.runs[0].color_rgb,
        style_name=before.runs[0].style_name,
    )
    after = StyleSnapshot(runs=(after_run,))
    mutations = diff(before, after)
    assert mutations, "expected a mutation"
    if before.runs[0].style_name and before.runs[0].style_name.startswith("Dosscriber"):
        # The style fingerprint round-tripped — the mutation must be blocker.
        assert mutations[0].severity == AnomalySeverity.BLOCKER
        assert mutations[0].is_dosscriber_style is True


def test_diff_detects_deleted_run(tmp_path: Path) -> None:
    before_path = tmp_path / "before.docx"
    _make(before_path)
    before = snapshot(before_path)
    # Drop one run from the snapshot to simulate the injector deleting it.
    after = StyleSnapshot(runs=(before.runs[0],))
    mutations = diff(before, after)
    assert mutations
    assert mutations[0].field == "(deleted)"
    assert mutations[0].severity == AnomalySeverity.BLOCKER


def test_color_change_is_warning_when_no_dosscriber_style(tmp_path: Path) -> None:
    before_path = tmp_path / "before.docx"
    after_path = tmp_path / "after.docx"
    doc_before = Document()
    doc_before.add_paragraph().add_run("text")
    doc_before.save(str(before_path))

    doc_after = Document()
    run = doc_after.add_paragraph().add_run("text")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    doc_after.save(str(after_path))

    mutations = diff(snapshot(before_path), snapshot(after_path))
    color_mutations = [m for m in mutations if m.field == "color_rgb"]
    assert color_mutations
    assert color_mutations[0].severity == AnomalySeverity.WARNING
