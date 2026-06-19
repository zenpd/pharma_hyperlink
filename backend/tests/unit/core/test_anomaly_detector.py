"""Unit tests for validation/anomaly_detector.py (W8.1)."""

from __future__ import annotations

from pathlib import Path

import yaml
from docx import Document

from hyperlink_engine.core.validation.anomaly_detector import (
    AnomalyReport,
    _load_deprecated_registry,
    aggregate_anomaly_reports,
    detect_blue_text_no_link,
    detect_deprecated_study_ids,
    detect_orphaned_references,
    detect_suspicious_targets,
    run_anomaly_detection,
)
from hyperlink_engine.models import (
    AnomalyKind,
    AnomalySeverity,
    LinkRecord,
    LinkStatus,
)

# ── Helpers ───────────────────────────────────────────────────────────────────


def _make_link_record(
    link_text: str = "Section 2.5.3",
    status: LinkStatus = LinkStatus.OK,
    target_anchor: str | None = "section_ref_2_5_3",
) -> LinkRecord:
    return LinkRecord(
        source_doc="doc.docx",
        link_text=link_text,
        link_location_descriptor="p0.r0:c0-10",
        target_anchor=target_anchor,
        status=status,
        confidence=0.9,
    )


def _make_registry_yaml(tmp_path: Path, entries: list[dict]) -> Path:
    path = tmp_path / "deprecated_ids.yaml"
    path.write_text(yaml.dump({"deprecated": entries}), encoding="utf-8")
    return path


# ── detect_blue_text_no_link ───────────────────────────────────────────────


def test_blue_text_no_link_returns_empty_for_no_blue_runs(tmp_path: Path) -> None:
    """A doc with no blue runs produces no anomalies."""
    from hyperlink_engine.core.parsing.docx_parser import parse_docx

    src = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("Section 2.5.3 is here")
    doc.save(str(src))

    parsed = parse_docx(src)
    anomalies = detect_blue_text_no_link(parsed, document_path=src)
    assert anomalies == []


def test_blue_text_no_link_detects_blue_unlinked_run(tmp_path: Path) -> None:
    """A blue-coloured run without a hyperlink triggers a WARNING anomaly."""
    from docx.shared import RGBColor

    from hyperlink_engine.core.parsing.docx_parser import parse_docx

    src = tmp_path / "blue.docx"
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("Blue text without link")
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # pure blue
    doc.save(str(src))

    parsed = parse_docx(src)
    anomalies = detect_blue_text_no_link(parsed, document_path=str(src))

    assert any(a.kind == AnomalyKind.BLUE_TEXT_NO_LINK for a in anomalies)
    assert all(a.severity == AnomalySeverity.WARNING for a in anomalies)
    assert all(a.document == str(src) for a in anomalies)


# ── detect_orphaned_references ─────────────────────────────────────────────


def test_orphan_detection_no_orphans() -> None:
    texts = ["Section 2.5.3"]
    records = [_make_link_record("Section 2.5.3")]
    anomalies = detect_orphaned_references(texts, records, document_path="doc.docx")
    assert anomalies == []


def test_orphan_detection_finds_unlinked_ref() -> None:
    texts = ["Section 2.5.3", "Section 5.1"]
    records = [_make_link_record("Section 2.5.3")]
    anomalies = detect_orphaned_references(texts, records, document_path="doc.docx")
    assert len(anomalies) == 1
    assert anomalies[0].kind == AnomalyKind.ORPHANED_REFERENCE
    assert "Section 5.1" in anomalies[0].text
    assert anomalies[0].severity == AnomalySeverity.WARNING


def test_orphan_detection_all_orphaned() -> None:
    texts = ["Section 1", "Table 2"]
    anomalies = detect_orphaned_references(texts, [], document_path="doc.docx")
    assert len(anomalies) == 2


def test_orphan_detection_empty_inputs() -> None:
    assert detect_orphaned_references([], [], document_path="doc.docx") == []


# ── detect_deprecated_study_ids ────────────────────────────────────────────


def test_deprecated_registry_loads_from_yaml(tmp_path: Path) -> None:
    path = _make_registry_yaml(
        tmp_path,
        [{"id": "SP-2019-001", "reason": "Terminated", "replaced_by": None}],
    )
    entries = _load_deprecated_registry(path)
    assert len(entries) == 1
    assert entries[0].id == "SP-2019-001"


def test_deprecated_registry_missing_file_returns_empty(tmp_path: Path) -> None:
    entries = _load_deprecated_registry(tmp_path / "missing.yaml")
    assert entries == []


def test_deprecated_ids_hit(tmp_path: Path) -> None:
    path = _make_registry_yaml(
        tmp_path,
        [{"id": "SP-2019-001", "reason": "Terminated", "replaced_by": "SP-2022-001"}],
    )
    text = "See Study SP-2019-001 for background data."
    anomalies = detect_deprecated_study_ids(
        text, document_path="doc.docx", registry_path=path
    )
    assert len(anomalies) == 1
    assert anomalies[0].kind == AnomalyKind.DEPRECATED_STUDY_ID
    assert anomalies[0].severity == AnomalySeverity.BLOCKER
    assert "SP-2022-001" in anomalies[0].suggested_fix


def test_deprecated_ids_no_match(tmp_path: Path) -> None:
    path = _make_registry_yaml(
        tmp_path,
        [{"id": "SP-2019-001", "reason": "Old"}],
    )
    text = "See Study SP-2024-999 for data."
    anomalies = detect_deprecated_study_ids(
        text, document_path="doc.docx", registry_path=path
    )
    assert anomalies == []


def test_deprecated_ids_empty_registry(tmp_path: Path) -> None:
    path = _make_registry_yaml(tmp_path, [])
    text = "Any text mentioning SP-2019-001"
    anomalies = detect_deprecated_study_ids(
        text, document_path="doc.docx", registry_path=path
    )
    assert anomalies == []


def test_deprecated_ids_no_partial_match(tmp_path: Path) -> None:
    """'SP-2019-001B' should NOT match the 'SP-2019-001' registry entry."""
    path = _make_registry_yaml(
        tmp_path,
        [{"id": "SP-2019-001", "reason": "Old"}],
    )
    text = "Study SP-2019-001B is different."
    anomalies = detect_deprecated_study_ids(
        text, document_path="doc.docx", registry_path=path
    )
    assert anomalies == []


# ── detect_suspicious_targets ──────────────────────────────────────────────


def test_suspicious_target_matching_numbers_no_anomaly() -> None:
    record = _make_link_record("Section 2.5.3", target_anchor="section_ref_2_5_3")
    anomalies = detect_suspicious_targets([record], document_path="doc.docx")
    assert anomalies == []


def test_suspicious_target_mismatched_numbers() -> None:
    record = _make_link_record("Section 5.3.2", target_anchor="section_ref_4_1_0")
    anomalies = detect_suspicious_targets(
        [record], document_path="doc.docx", similarity_threshold=0.2
    )
    assert len(anomalies) == 1
    assert anomalies[0].kind == AnomalyKind.SUSPICIOUS_TARGET
    assert anomalies[0].severity == AnomalySeverity.WARNING


def test_suspicious_target_broken_links_skipped() -> None:
    """Already-broken links shouldn't produce suspicious-target anomalies."""
    record = _make_link_record(
        "Section 5.3.2",
        status=LinkStatus.BROKEN,
        target_anchor="section_ref_4_1_0",
    )
    anomalies = detect_suspicious_targets([record], document_path="doc.docx")
    assert anomalies == []


def test_suspicious_target_no_anchor_skipped() -> None:
    record = _make_link_record("Section 2.5.3", target_anchor=None)
    anomalies = detect_suspicious_targets([record], document_path="doc.docx")
    assert anomalies == []


def test_suspicious_target_no_numbers_in_text_skipped() -> None:
    record = LinkRecord(
        source_doc="doc.docx",
        link_text="Appendix",
        link_location_descriptor="p0.r0:c0-8",
        target_anchor="appendix_a",
        status=LinkStatus.OK,
        confidence=0.9,
    )
    anomalies = detect_suspicious_targets([record], document_path="doc.docx")
    assert anomalies == []


# ── run_anomaly_detection (integration) ───────────────────────────────────


def test_run_anomaly_detection_smoke(tmp_path: Path) -> None:
    path = _make_registry_yaml(
        tmp_path,
        [{"id": "SP-2019-001", "reason": "Old"}],
    )
    report = run_anomaly_detection(
        document_path="doc.docx",
        detection_texts=["Section 2.5.3", "Orphaned ref"],
        link_records=[_make_link_record("Section 2.5.3")],
        full_text="See SP-2019-001 for background.",
        deprecated_registry_path=path,
    )
    assert isinstance(report, AnomalyReport)
    # Should find orphan + deprecated ID
    kinds = {a.kind for a in report.anomalies}
    assert AnomalyKind.ORPHANED_REFERENCE in kinds
    assert AnomalyKind.DEPRECATED_STUDY_ID in kinds


def test_run_anomaly_detection_all_checks_disabled() -> None:
    report = run_anomaly_detection(
        document_path="doc.docx",
        check_blue_text=False,
        check_orphans=False,
        check_deprecated=False,
        check_suspicious=False,
    )
    assert report.total == 0


# ── AnomalyReport helpers ─────────────────────────────────────────────────


def test_anomaly_report_counts() -> None:
    from hyperlink_engine.models import Anomaly, AnomalyKind, AnomalySeverity

    report = AnomalyReport(
        document="doc.docx",
        anomalies=[
            Anomaly(
                kind=AnomalyKind.BLUE_TEXT_NO_LINK,
                severity=AnomalySeverity.BLOCKER,
                document="doc.docx",
                text="x",
                confidence=0.9,
            ),
            Anomaly(
                kind=AnomalyKind.ORPHANED_REFERENCE,
                severity=AnomalySeverity.WARNING,
                document="doc.docx",
                text="y",
                confidence=0.8,
            ),
        ],
    )
    assert report.blocker_count == 1
    assert report.warning_count == 1
    assert report.total == 2
    assert len(report.by_kind(AnomalyKind.BLUE_TEXT_NO_LINK)) == 1


# ── DossierAnomalySummary ────────────────────────────────────────────────


def test_aggregate_anomaly_reports() -> None:
    from hyperlink_engine.models import Anomaly, AnomalyKind, AnomalySeverity

    def _make_report(doc: str, kind: AnomalyKind, sev: AnomalySeverity) -> AnomalyReport:
        return AnomalyReport(
            document=doc,
            anomalies=[
                Anomaly(kind=kind, severity=sev, document=doc, text="x", confidence=0.9)
            ],
        )

    reports = [
        _make_report("a.docx", AnomalyKind.DEPRECATED_STUDY_ID, AnomalySeverity.BLOCKER),
        _make_report("b.docx", AnomalyKind.ORPHANED_REFERENCE, AnomalySeverity.WARNING),
    ]
    summary = aggregate_anomaly_reports(reports)
    assert summary.total_anomalies == 2
    assert summary.total_blockers == 1
    assert summary.total_warnings == 1
    assert len(summary.by_kind(AnomalyKind.DEPRECATED_STUDY_ID)) == 1
    assert "a.docx" in summary.documents_with_blockers()
