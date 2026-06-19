"""PLAN THIRTEEN — root-reference correctness.

Fix 1 (this plan): a study's protocol/sap/csr/listings siblings must share ONE
study key so a "the protocol" / "the SAP" reference routes to the right sibling.
Before the fix, the doc-type word prefixing the filename corrupted the key
(``protocol-sp-2026-001`` keyed on ``colsp2026001``), so siblings got different
keys and every DOC_REF silently dropped.

Fix 2 (already provided by the PLAN TEN-bis caption scan, locked here): an appendix
*definition* line — numbered ("Appendix 16.1 …") or letter ("Appendix A: …") — is
indexed under ``appendix_ref_<n>`` so the appendix citation lands on it, not on the
first mention or the document top.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from hyperlink_engine.core.injection.anchor_index import build_anchor_index
from hyperlink_engine.orchestration.nodes import (
    _build_file_index,
    _norm,
    _resolve_one,
    _robust_study_key,
)
from hyperlink_engine.workers.tasks import detect_references


# ── Fix 1 — study key is doc-type-agnostic ───────────────────────────────────

def test_study_key_strips_doctype_prefix() -> None:
    """All four siblings of one study collapse to the same key."""
    for stem in (
        "csr-sp-2026-001-body",
        "listings-sp-2026-001",
        "protocol-sp-2026-001",
        "sap-sp-2026-001",
    ):
        assert _robust_study_key(_norm(stem)) == "sp2026001", stem


def test_study_key_nct_unchanged() -> None:
    """The NCT path is matched first and untouched by the prefix stripping."""
    assert _robust_study_key(_norm("NCT04089566_Protocol")) == "nct04089566"
    assert _robust_study_key(_norm("NCT04089566_SAP")) == "nct04089566"


def test_doc_ref_routes_to_same_study_sibling() -> None:
    """'the protocol' / 'the SAP' resolve to that study's protocol / sap doc and
    never cross studies."""
    files = [
        "csr-sp-2026-001-body.docx",
        "protocol-sp-2026-001.docx",
        "sap-sp-2026-001.docx",
        "listings-sp-2026-001.docx",
        "csr-sp-2026-002-body.docx",
        "protocol-sp-2026-002.docx",
    ]
    fidx = _build_file_index([{"source_path": f} for f in files])

    src1 = "csr-sp-2026-001-body.docx"
    assert Path(_resolve_one({"label": "DOC_REF", "text": "the protocol"}, fidx, src1)).name == "protocol-sp-2026-001.docx"
    assert Path(_resolve_one({"label": "DOC_REF", "text": "the SAP"}, fidx, src1)).name == "sap-sp-2026-001.docx"

    # study 002's protocol ref must route to 002's protocol, not 001's.
    src2 = "csr-sp-2026-002-body.docx"
    assert Path(_resolve_one({"label": "DOC_REF", "text": "the protocol"}, fidx, src2)).name == "protocol-sp-2026-002.docx"


# ── Fix 2 — appendix definitions are anchored (lock PLAN TEN-bis behavior) ────

def test_appendix_definition_is_anchored(tmp_path: Path) -> None:
    p = tmp_path / "with_appendix.docx"
    d = Document()
    d.add_paragraph("See Appendix 16.1 and Appendix A for details.")  # citations (para 0)
    d.add_paragraph("Body text.")                                     # para 1
    d.add_paragraph("Appendix 16.1 Schedule of Assessments")          # numbered def (para 2)
    d.add_paragraph("Appendix A: Subject Data Listings")              # letter def   (para 3)
    d.save(str(p))

    drec = detect_references({"source_path": str(p), "filename": p.name, "suffix": ".docx"})
    idx = build_anchor_index(drec["detections"], str(p), is_pdf=False)

    # citations land on the DEFINITION paragraphs, not the para-0 mention.
    assert idx.get("appendix_ref_16_1") == {"paragraph_index": 2}
    assert idx.get("appendix_ref_A") == {"paragraph_index": 3}
