"""PLAN SEVENTEEN — whole-reference linking, ISS/ISE, NCT-always-external.

Locks in the behaviour the publishing team asked for on the real NCT01101035
dossier (CSR / Clinical Summary / ISE / ISS / Protocol / SAP):

* The WHOLE compound phrase ("Protocol TMX-67_301 Section 6.1", "SAP Section 5.3")
  links to ONE destination — the section in the named document — so the pieces
  (DOC_REF + DOC_ID + SECTION_REF) all resolve to the same target.
* ISS / ISE are detected and route to the ISS / ISE documents.
* An NCT id is ALWAYS an external clinicaltrials.gov link, never a cross-doc link.
* The target document declares a bookmark for every anchor a sibling links INTO
  it, so the cross-doc citation actually scrolls (not just opens the file).
"""

from __future__ import annotations

import pytest
from docx import Document as Docx
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

from hyperlink_engine.core.detection.regex_patterns import default_registry
from hyperlink_engine.core.injection.anchor_index import _caption_def_num
from hyperlink_engine.workers.cache import ExtractorConfig, get_extractor
from hyperlink_engine.models import LinkKind
from hyperlink_engine.orchestration.nodes import _build_file_index, _resolve_one
from hyperlink_engine.workers.tasks import (
    _compound_anchor_key,
    _compound_doc_type,
    _dedupe_fullname_acronym_docx,
    _enrich_compound_docx,
    _ensure_highlighted_linked_docx,
    _resolve_target,
    inject_links,
)


def _doc_ref_hits(text: str) -> list[str]:
    reg = default_registry()
    return [m.text for m in reg.find_all(text) if reg.get(m.pattern_id).label == "DOC_REF"]


def _index(*names: str):
    return _build_file_index([{"source_path": rf"C:\u\{n}.docx"} for n in names])


# ── ISS / ISE detection ──────────────────────────────────────────────────────


@pytest.mark.parametrize(
    "text,needle",
    [
        ("safety analyses are described in the ISS", "ISS"),
        ("results are interpreted in the ISE", "ISE"),
        ("see the Integrated Summary of Safety for details", "Integrated Summary of Safety"),
        ("per the Integrated Summary of Efficacy", "Integrated Summary of Efficacy"),
    ],
)
def test_iss_ise_detected(text: str, needle: str) -> None:
    assert any(needle in h for h in _doc_ref_hits(text)), f"missing {needle!r} in {text!r}"


def test_iss_parenthetical_restatement_deduped() -> None:
    # "Integrated Summary of Safety (ISS)" → ONE link, on the ACRONYM "(ISS)"; the
    # spelled-out form is dropped (same rule as "(SAP)"/"(CSR)").
    hits = _doc_ref_hits("This Integrated Summary of Safety (ISS) provides an evaluation")
    assert "ISS" in hits
    assert "Integrated Summary of Safety" not in hits


# ── Compound phrase enrichment ───────────────────────────────────────────────


def test_compound_doc_type_and_anchor_key() -> None:
    assert _compound_doc_type("Protocol") == "protocol"
    assert _compound_doc_type("SAP") == "sap"
    assert _compound_doc_type("Integrated Summary of Safety") == "iss"
    assert _compound_anchor_key(" Section 6.1") == "section_ref_6_1"
    assert _compound_anchor_key(" Table 14.2.1") == "table_ref_14_2_1"


def test_enrich_compound_merges_into_one_continuous_link() -> None:
    """'Protocol TMX-67_301 Section 6.1' split across runs collapses to ONE
    compound detection spanning all the runs — so the injector wraps the whole
    phrase in a single continuous hyperlink (not three adjacent boxes). The
    fragment detections are removed; the compound carries the shared doc-type +
    section anchor and a multi-run span (run_index → end_run_index)."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("The design is in ")  # run 0
    p.add_run("Protocol ")  # run 1
    p.add_run("TMX-67_301")  # run 2
    p.add_run(" Section 6.1")  # run 3
    p.add_run(", and more text")  # run 4
    para_dets = [
        {"label": "DOC_REF", "text": "Protocol", "run_index": 1, "char_start": 0, "char_end": 8, "groups": {}},
        {"label": "DOC_ID", "text": "TMX-67_301", "run_index": 2, "char_start": 0, "char_end": 10, "groups": {}},
        {"label": "SECTION_REF", "text": "Section 6.1", "run_index": 3, "char_start": 1, "char_end": 12, "groups": {}},
    ]
    _enrich_compound_docx(0, p, para_dets)
    comp = [det for det in para_dets if det.get("is_compound")]
    assert len(comp) == 1, "exactly one compound detection should replace the pieces"
    c = comp[0]
    assert c["text"] == "Protocol TMX-67_301 Section 6.1"
    assert c["groups"]["compound_qual_dt"] == "protocol"
    assert c["groups"]["compound_anchor_key"] == "section_ref_6_1"
    # multi-run span: the link runs from run 1 through run 3 as ONE hyperlink.
    assert c["run_index"] == 1 and c["end_run_index"] == 3
    # the fragment detections are gone (they would have been separate links).
    assert not [det for det in para_dets if det["label"] in ("DOC_ID", "SECTION_REF")]


def test_enrich_single_run_compound_is_one_span() -> None:
    """When Word kept the whole phrase in ONE run, the compound is a single-run
    span (end_run_index == run_index) covering the entire phrase."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("The design is in ")  # run 0
    p.add_run("Protocol TMX-67_301 Section 6.1")  # run 1 (whole phrase)
    p.add_run(", and more.")  # run 2
    para_dets = [
        {"label": "DOC_REF", "text": "Protocol", "run_index": 1, "char_start": 0, "char_end": 8, "groups": {}},
        {"label": "SECTION_REF", "text": "Section 6.1", "run_index": 1, "char_start": 20, "char_end": 31, "groups": {}},
    ]
    _enrich_compound_docx(0, p, para_dets)
    comp = [det for det in para_dets if det.get("is_compound")]
    assert len(comp) == 1
    c = comp[0]
    assert c["text"] == "Protocol TMX-67_301 Section 6.1"
    assert c["run_index"] == 1 and c["end_run_index"] == 1
    assert c["char_start"] == 0 and c["char_end"] == 31
    assert c["groups"]["compound_anchor_key"] == "section_ref_6_1"


def test_enrich_bare_doctype_not_merged() -> None:
    """A bare 'SAP' with no id and no section is NOT a compound (handled by the
    plain DOC_REF pattern) — it must stay as-is, never merged or stamped."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("as defined in the SAP and reported elsewhere")
    para_dets = [
        {"label": "DOC_REF", "text": "SAP", "run_index": 0, "char_start": 18, "char_end": 21, "groups": {}},
    ]
    _enrich_compound_docx(0, p, para_dets)
    assert len(para_dets) == 1
    assert not para_dets[0].get("is_compound")
    assert "compound_qual_dt" not in para_dets[0]["groups"]


def test_injector_wraps_multi_run_span_as_one_hyperlink(tmp_path) -> None:
    """The docx injector wraps a span that crosses runs in ONE w:hyperlink, so a
    compound phrase Word split into runs renders as a single continuous link."""
    from docx.oxml.ns import qn

    from hyperlink_engine.core.injection.docx_linker import DocxLinker
    from hyperlink_engine.models import RunLocation

    src = tmp_path / "src.docx"
    d = Docx()
    p = d.add_paragraph()
    p.add_run("see ")  # run 0
    p.add_run("Protocol ")  # run 1
    p.add_run("TMX-67_301")  # run 2
    p.add_run(" Section 6.1")  # run 3
    p.add_run(" now.")  # run 4
    d.save(src)

    linker = DocxLinker(src, tmp_path / "out.docx")
    linker.add_external_link(
        RunLocation(
            paragraph_index=0, run_index=1, char_start=0, char_end=9,
            end_run_index=3, end_char=12,
        ),
        url="Protocol_linked.docx#section_ref_6_1",
    )
    out = linker.save()

    doc = Docx(str(out))
    hls = doc.paragraphs[0]._p.findall(qn("w:hyperlink"))
    assert len(hls) == 1, "the multi-run span must be ONE hyperlink element"
    htext = "".join(n.text or "" for n in hls[0].iter(qn("w:t")))
    assert htext == "Protocol TMX-67_301 Section 6.1"
    # surrounding text is preserved outside the link.
    assert "see " in doc.paragraphs[0].text
    assert " now." in doc.paragraphs[0].text


# ── Resolution: compound pieces → same target; NCT external ───────────────────


def test_compound_doc_id_routes_to_protocol() -> None:
    idx = _index("CSR", "Protocol", "SAP")
    det = {
        "label": "DOC_ID", "text": "TMX-67_301", "pattern_id": "DOC_ID_V1",
        "groups": {"compound_qual_dt": "protocol"}, "context": "",
    }
    target = _resolve_one(det, idx, r"C:\u\CSR.docx")
    assert target is not None and target.endswith("Protocol.docx")


def test_compound_section_routes_to_qualified_sibling() -> None:
    idx = _index("CSR", "Protocol", "SAP")
    det = {
        "label": "SECTION_REF", "text": "Section 6.1", "pattern_id": "SECTION_REF_LABELED_V1",
        "groups": {"compound_qual_dt": "protocol"}, "context": "Protocol TMX-67_301 Section 6.1",
    }
    target = _resolve_one(det, idx, r"C:\u\CSR.docx")
    assert target is not None and target.endswith("Protocol.docx")


def test_iss_doc_ref_routes_to_iss_doc() -> None:
    idx = _index("CSR", "ISS", "ISE")
    det = {"label": "DOC_REF", "text": "ISS", "pattern_id": "DOC_REF_ISS_V1", "groups": {}, "context": ""}
    target = _resolve_one(det, idx, r"C:\u\CSR.docx")
    assert target is not None and target.endswith("ISS.docx")


def test_sap_ref_prefers_same_format_sibling_over_clinicaltrials_pdf() -> None:
    """Adding a clinicaltrials NCT…_SAP.pdf must NOT break the Word 'SAP Section 5.3'
    link: with two sap-typed docs, a .docx source resolves to the .docx SAP, not the
    PDF copy in the sub-folder. (Regression: the PDF made the ref ambiguous → dropped.)"""
    idx = _build_file_index(
        [
            {"source_path": r"C:\u\CSR.docx"},
            {"source_path": r"C:\u\SAP.docx"},
            {"source_path": r"C:\u\Protocol.docx"},
            {"source_path": r"C:\u\clinicaltrials\NCT01101035_SAP.pdf"},
        ]
    )
    compound = {
        "label": "DOC_REF", "text": "SAP Section 5.3", "pattern_id": "COMPOUND_REF_V1",
        "groups": {"compound_qual_dt": "sap", "compound_anchor_key": "section_ref_5_3"},
        "context": "as defined in SAP Section 5.3",
    }
    target = _resolve_one(compound, idx, r"C:\u\CSR.docx")
    assert target is not None and target.endswith("SAP.docx")
    # Bare "SAP" from a different Word source resolves the same way.
    bare = {"label": "DOC_REF", "text": "SAP", "pattern_id": "DOC_REF_SAP_V1", "groups": {},
            "context": "described in detail in the SAP"}
    assert _resolve_one(bare, idx, r"C:\u\Protocol.docx").endswith("SAP.docx")


def test_ambiguous_same_type_same_format_stays_unresolved() -> None:
    """The 'don't guess among unrelated studies' guard is preserved: two same-format
    SAPs in the same folder with no study-key match resolve to None, not a coin-flip."""
    idx = _build_file_index(
        [{"source_path": r"C:\u\CSR.docx"},
         {"source_path": r"C:\u\SAP_a.docx"},
         {"source_path": r"C:\u\SAP_b.docx"}]
    )
    det = {"label": "DOC_REF", "text": "SAP", "pattern_id": "DOC_REF_SAP_V1", "groups": {}, "context": ""}
    assert _resolve_one(det, idx, r"C:\u\CSR.docx") is None


def test_nct_never_resolves_cross_doc() -> None:
    idx = _index("CSR", "Protocol", "SAP")
    det = {
        "label": "STUDY_ID", "text": "NCT01101035", "pattern_id": "STUDY_ID_NCT_V1",
        "groups": {}, "context": "described in the Protocol for Study NCT01101035",
    }
    assert _resolve_one(det, idx, r"C:\u\CSR.docx") is None


def test_nct_resolves_to_external_url() -> None:
    det = {"label": "STUDY_ID", "text": "NCT01101035", "pattern_id": "STUDY_ID_NCT_V1", "groups": {}}
    kind, target = _resolve_target(det)
    assert kind == LinkKind.EXTERNAL_URL
    assert target == "https://clinicaltrials.gov/study/NCT01101035"


# ── Injection: cross-doc target declares the incoming bookmark ────────────────


def test_score_reflects_real_broken_links(tmp_path) -> None:
    """The readiness score must move off 100 when a link is genuinely broken — it
    used to be hard-coded 'ok' (always 100). node_validate now does a real
    bookmark/file/URL check and node_score_and_report counts the true broken set."""
    from hyperlink_engine.core.injection.docx_linker import DocxLinker
    from hyperlink_engine.models import RunLocation
    from hyperlink_engine.orchestration.nodes import node_score_and_report, node_validate

    # A target doc that declares ONLY section_ref_6_1 (so 9_9 is genuinely broken).
    src = tmp_path / "Protocol.docx"
    d = Docx()
    d.add_paragraph("Intro")
    d.add_paragraph("6.1 Study Design")
    d.save(src)
    linker = DocxLinker(src, tmp_path / "Protocol_linked.docx")
    linker.add_bookmark(RunLocation(paragraph_index=1, run_index=0, char_start=0, char_end=1), "section_ref_6_1")
    linker.save()

    state = {
        "run_id": "t",
        "output_dir": tmp_path,
        "injection_records": [
            {
                "ingest": {"source_path": str(tmp_path / "CSR.docx")},
                "probes": [
                    {"source_doc": "CSR.docx", "link_text": "Section 6.1",
                     "target": "Protocol_linked.docx#section_ref_6_1",
                     "target_doc": "Protocol_linked.docx", "kind": "external_url"},
                    {"source_doc": "CSR.docx", "link_text": "Section 9.9",
                     "target": "Protocol_linked.docx#section_ref_9_9",
                     "target_doc": "Protocol_linked.docx", "kind": "external_url"},
                    {"source_doc": "CSR.docx", "link_text": "NCT01101035",
                     "target": "https://clinicaltrials.gov/study/NCT01101035",
                     "target_doc": "CSR.docx", "kind": "external_url"},
                ],
            }
        ],
    }
    node_validate(state)  # type: ignore[arg-type]
    status = {l["link_text"]: l["status"] for l in state["links"]}
    assert status["Section 6.1"] == "ok"          # bookmark exists → resolvable
    assert status["Section 9.9"] == "broken"      # bookmark missing → broken
    assert status["NCT01101035"] == "ok"          # valid external URL
    assert status["Section 6.1"] == "ok" and status["Section 9.9"] == "broken"

    node_score_and_report(state)  # type: ignore[arg-type]
    assert state["score"] < 100.0                 # the score now REACTS (was always 100)


def test_incoming_anchor_bookmark_declared(tmp_path) -> None:
    """A document declares a bookmark for an anchor a sibling links INTO it, even
    when it never cites that section itself — so the cross-doc link can scroll."""
    src = tmp_path / "Protocol.docx"
    d = Docx()
    d.add_paragraph("Introduction")
    d.add_paragraph("Methods overview")
    d.add_paragraph("6.1 Study Design and Randomization")  # definition, para index 2
    d.save(src)

    drec = {"ingest": {"source_path": str(src)}, "detections": []}
    out = tmp_path / "Protocol_linked.docx"
    idx = {"Protocol": {"section_ref_6_1": {"paragraph_index": 2}}}
    inject_links(
        drec,
        output_path=str(out),
        target_anchor_indexes=idx,
        incoming_anchor_keys={"section_ref_6_1"},
    )
    names = {
        b.get(qn("w:name"))
        for b in Docx(out).element.body.iter(qn("w:bookmarkStart"))
    }
    assert "section_ref_6_1" in names


# ── Highlight safety net: every highlighted span gets linked; plain docs untouched ──


def _hl(run):
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


def test_ensure_pass_is_noop_without_highlights() -> None:
    """On a plain (un-highlighted) paragraph the highlight pass changes nothing — so
    documents uploaded without yellow marking link purely by detection."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("see Table 5.1 for details")
    para_dets = [
        {"label": "TABLE_REF", "text": "Table 5.1", "run_index": 0, "char_start": 4,
         "char_end": 13, "groups": {}, "pattern_id": "T", "source_layer": "regex",
         "confidence": 0.9},
    ]
    before = [dict(x) for x in para_dets]
    _ensure_highlighted_linked_docx(0, p, para_dets, get_extractor(ExtractorConfig()))
    assert para_dets == before  # strict no-op


def test_ensure_pass_links_highlighted_bare_id_via_context() -> None:
    """A highlighted bare id ('TMX-67_301' in 'Protocol ID- TMX-67_301') is upgraded:
    the highlight pass stamps the doc-type from the preceding 'Protocol' so it routes
    to the Protocol document (was dropped as an un-resolvable DOC_ID)."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("Protocol ID- ")  # run 0 (not highlighted)
    _hl(p.add_run("TMX-67_301"))  # run 1 (highlighted)
    para_dets = [
        {"label": "DOC_ID", "text": "TMX-67_301", "run_index": 1, "char_start": 0,
         "char_end": 10, "groups": {}, "pattern_id": "DOC_ID_V1",
         "source_layer": "regex", "confidence": 0.7},
    ]
    _ensure_highlighted_linked_docx(0, p, para_dets, get_extractor(ExtractorConfig()))
    tmx = [x for x in para_dets if "TMX" in x["text"]]
    assert tmx and tmx[0]["groups"].get("compound_qual_dt") == "protocol"


def test_ensure_pass_does_not_relabel_doc_ref_from_context() -> None:
    """A highlighted 'CSR' must NOT be re-stamped as 'sap' just because 'SAP' sits
    earlier in the sentence ('between SAP and CSR') — DOC_REF keeps its own type."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("traceability between SAP and ")  # run 0
    _hl(p.add_run("CSR"))  # run 1 (highlighted)
    p.add_run(".")  # run 2
    para_dets = [
        {"label": "DOC_REF", "text": "CSR", "run_index": 1, "char_start": 0,
         "char_end": 3, "groups": {}, "pattern_id": "DOC_REF_CSR_V1",
         "source_layer": "regex", "confidence": 0.7},
    ]
    _ensure_highlighted_linked_docx(0, p, para_dets, get_extractor(ExtractorConfig()))
    csr = [x for x in para_dets if x["text"].strip() == "CSR"]
    assert csr and csr[0]["groups"].get("compound_qual_dt") != "sap"


def test_appendix_citation_tail_not_a_definition() -> None:
    """A citation that merely ends a sentence ('refer Appendix 2.') is NOT the
    appendix definition (so the anchor lands on the real 'Appendix 2 Clinical Safety'
    heading), while a titled heading still is."""
    # require_title=True (the detection-context pass): the period-terminated citation
    # must be rejected.
    assert _caption_def_num("Appendix 2.", "APPENDIX_REF", require_title=True) is None
    # a real titled heading is the definition.
    assert _caption_def_num("Appendix 2 Clinical Safety", "APPENDIX_REF", require_title=True) == "2"
    # a bare all-caps heading is still recovered by the structure scan (no title req).
    assert _caption_def_num("APPENDIX A.", "APPENDIX_REF", require_title=False) == "A"


# ── Re-target the author's pre-existing placeholder hyperlinks ("ISE" → about:blank) ──


def _add_existing_hyperlink(para, text: str, target: str) -> None:
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement

    rid = para.part.relate_to(target, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)


def test_existing_placeholder_hyperlink_detected_and_retargeted(tmp_path) -> None:
    """An author's 'ISE' → about:blank placeholder (invisible to para.runs) is
    detected and the existing hyperlink is re-pointed at the ISE document."""
    from hyperlink_engine.core.injection.docx_linker import DocxLinker
    from hyperlink_engine.workers.tasks import _detect_existing_hyperlinks_docx

    src = tmp_path / "CSR.docx"
    d = Docx()
    p = d.add_paragraph("These findings are presented in the ")
    _add_existing_hyperlink(p, "ISE", "about:blank")
    d.save(src)

    d2 = Docx(str(src))
    dets = _detect_existing_hyperlinks_docx(
        0, d2.paragraphs[0], get_extractor(ExtractorConfig())
    )
    assert dets and dets[0]["existing_hl_index"] == 0
    assert "ISE" in dets[0]["text"] and dets[0]["label"] == "DOC_REF"

    linker = DocxLinker(src, tmp_path / "out.docx")
    linker.retarget_existing_hyperlink(0, 0, url="ISE_linked.docx")
    out = linker.save()
    od = Docx(str(out))
    hl = od.paragraphs[0]._p.findall(qn("w:hyperlink"))[0]
    assert od.part.rels[hl.get(qn("r:id"))].target_ref == "ISE_linked.docx"


def test_existing_real_url_left_untouched(tmp_path) -> None:
    """A real external link (e.g. a SharePoint URL) is NOT a placeholder, so it is
    never detected for re-targeting."""
    from hyperlink_engine.workers.tasks import _detect_existing_hyperlinks_docx

    d = Docx()
    p = d.add_paragraph("see ")
    _add_existing_hyperlink(p, "ISE", "https://example.sharepoint.com/real.docx")
    dets = _detect_existing_hyperlinks_docx(0, p, get_extractor(ExtractorConfig()))
    assert dets == []


# ── "Study ID <id>" id-routed compound (2026-06-19 demo) ─────────────────────


def test_study_id_compound_is_one_id_routed_span() -> None:
    """'Study ID TMX-67_301 Section 6.1' → ONE compound DOC_ID whose TEXT is the
    WHOLE phrase (so the injected link AND the preview box span the continuous
    reference), routed by the bare id via `compound_id_token` and carrying the
    section anchor — NOT a doc-type compound."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("The study design is described in Study ID TMX-67_301 Section 6.1, while")
    dets: list[dict] = []
    _enrich_compound_docx(0, p, dets)
    comp = [x for x in dets if x.get("is_compound")]
    assert len(comp) == 1
    c = comp[0]
    assert c["label"] == "DOC_ID"
    assert c["text"] == "Study ID TMX-67_301 Section 6.1"  # whole continuous phrase
    assert c["groups"].get("compound_id_token") == "TMX-67_301"  # routing id only
    assert c["groups"].get("compound_anchor_key") == "section_ref_6_1"
    assert "compound_qual_dt" not in c["groups"]  # routes by id, not doc-type


def test_study_id_bare_compound_no_anchor() -> None:
    """Bare 'Study ID TMX-67_301' (no section) is still ONE DOC_ID span (whole-doc),
    text = the whole phrase, routed by the id token."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("administration defined in the Study ID TMX-67_301, with safety")
    dets: list[dict] = []
    _enrich_compound_docx(0, p, dets)
    comp = [x for x in dets if x.get("is_compound")]
    assert len(comp) == 1 and comp[0]["label"] == "DOC_ID"
    assert comp[0]["text"] == "Study ID TMX-67_301"
    assert comp[0]["groups"].get("compound_id_token") == "TMX-67_301"
    assert "compound_anchor_key" not in comp[0]["groups"]


def test_study_id_doc_id_resolves_to_named_file_by_id() -> None:
    """A 'Study ID' compound resolves to the file whose name carries the id, even
    when that file is a PDF (the doc-type is 'other') — and even though the link
    TEXT is the whole phrase, routing uses only `compound_id_token` so the trailing
    'Section 6.1' words never skew the match."""
    idx = _build_file_index(
        [{"source_path": r"C:\u\CSR.docx"}, {"source_path": r"C:\u\Study ID TMX-67_301.pdf"}]
    )
    det = {
        "label": "DOC_ID", "text": "Study ID TMX-67_301 Section 6.1",
        "pattern_id": "COMPOUND_REF_V1",
        "groups": {"compound_id_token": "TMX-67_301", "compound_anchor_key": "section_ref_6_1"},
        "context": "",
    }
    target = _resolve_one(det, idx, r"C:\u\CSR.docx")
    assert target is not None and target.endswith("Study ID TMX-67_301.pdf")


# ── Highlighted self-reference links to the doc's own top ────────────────────


def test_highlighted_self_reference_flags_self_top() -> None:
    """A HIGHLIGHTED self-reference ('SAP' inside SAP.docx) is flagged for a
    link to its own top; it is never routed to a different document."""
    idx = _index("SAP", "CSR")
    det = {
        "label": "DOC_REF", "text": "SAP", "pattern_id": "DOC_REF_SAP_V1",
        "groups": {}, "context": "between SAP and CSR", "is_highlighted": True,
    }
    assert _resolve_one(det, idx, r"C:\u\SAP.docx") is None
    assert det.get("self_ref_top") is True


def test_unhighlighted_self_reference_not_flagged() -> None:
    """An un-highlighted self-reference stays a plain self-reference (no link)."""
    idx = _index("SAP", "CSR")
    det = {
        "label": "DOC_REF", "text": "SAP", "pattern_id": "DOC_REF_SAP_V1",
        "groups": {}, "context": "this SAP describes",
    }
    assert _resolve_one(det, idx, r"C:\u\SAP.docx") is None
    assert det.get("self_ref_top") is not True


# ── Cross-run "FullName (ACR)" de-dupe keeps the acronym ─────────────────────


def test_dedupe_fullname_acronym_cross_run_keeps_acronym() -> None:
    """When Word splits 'Statistical Analysis Plan (SAP)' across runs, the
    paragraph-level pass drops the spelled-out DOC_REF and keeps the acronym."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("detailed in the Statistical Analysis Plan (")  # run 0
    p.add_run("SAP")  # run 1
    p.add_run("), and more")  # run 2
    full = p.runs[0].text
    sp = full.index("Statistical Analysis Plan")
    dets = [
        {"label": "DOC_REF", "text": "Statistical Analysis Plan", "run_index": 0,
         "char_start": sp, "char_end": sp + len("Statistical Analysis Plan"), "groups": {}},
        {"label": "DOC_REF", "text": "SAP", "run_index": 1,
         "char_start": 0, "char_end": 3, "groups": {}},
    ]
    _dedupe_fullname_acronym_docx(p, dets)
    texts = [x["text"] for x in dets]
    assert "SAP" in texts
    assert "Statistical Analysis Plan" not in texts


# ── Compound DOC_ID survives a highlight GAP (not split into two) ────────────


def test_compound_doc_id_not_split_by_highlight_gap() -> None:
    """'Study ID' + un-highlighted space + 'TMX-67_301' (two highlighted runs):
    the compound DOC_ID span must stay ONE link — the ensure-net must not split
    it back into 'Study ID' and 'TMX-67_301'."""
    d = Docx()
    p = d.add_paragraph()
    p.add_run("defined in the ")  # run 0 — plain
    r1 = p.add_run("Study ID")  # run 1 — highlighted
    r1.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run(" ")  # run 2 — the un-highlighted GAP
    r3 = p.add_run("TMX-67_301")  # run 3 — highlighted
    r3.font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run(", with safety")  # run 4 — plain

    dets: list[dict] = []
    _enrich_compound_docx(0, p, dets)
    # tag highlighted-run dets like detect_references does (so the ensure-net runs)
    runs = list(p.runs)
    for det in dets:
        ri = det.get("run_index", 0)
        if ri < len(runs) and runs[ri].font.highlight_color == WD_COLOR_INDEX.YELLOW:
            det["is_highlighted"] = True
    _ensure_highlighted_linked_docx(0, p, dets, get_extractor(ExtractorConfig()))

    spans = [x for x in dets if x.get("label") == "DOC_ID"]
    assert len(spans) == 1, f"expected ONE continuous span, got {[s['text'] for s in spans]}"
    assert spans[0].get("end_run_index", spans[0]["run_index"]) > spans[0]["run_index"]
