"""Generate a LARGE 4-study CSR dossier (10+ pages per document).

Same shape and cross-reference scheme as ``generate_csr_dossier.py`` — four
study folders, four documents each (CSR body / Protocol / SAP / Listings) that
reference one another within and across studies — but every document is greatly
expanded with realistic regulatory narrative and many real Word tables so each
file runs well past ten printed pages. This stresses parsing, detection,
injection, the snippet/scroll preview, and the readiness report on volume.

    data/synthetic/csr_dossier_large/
      SP-2026-001/
        csr-sp-2026-001-body.docx       (~14-18 pages)
        protocol-sp-2026-001.docx       (~12-15 pages)
        sap-sp-2026-001.docx            (~11-14 pages)
        listings-sp-2026-001.docx       (~11-14 pages)
      SP-2026-002/ ... SP-2026-003/ ... SP-2026-004/ ...

The canonical cross-reference forms are preserved exactly as in the standard
dossier ("SP-2026-002", "Protocol", "SAP", "Listings", "CSR", "Section X.Y.Z",
"Table N.N.N.N", "Listing 16.x", "Appendix 16.x") so the detection and
target-resolution layers route every hyperlink to the correct document — the
only difference here is sheer length.

Run::

    poetry run python -m scripts.generate_csr_dossier_large
    # or:  python scripts/generate_csr_dossier_large.py --out data/synthetic/csr_dossier_large
    # add --ambiguous to seed sub-threshold refs that route through Ollama
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

STUDIES = ["SP-2026-001", "SP-2026-002", "SP-2026-003", "SP-2026-004"]

STUDY_TITLES = {
    "SP-2026-001": "A Phase 1 Single-Ascending-Dose PK Study of Solzumab",
    "SP-2026-002": "A Phase 2a Efficacy and Safety Study of Solzumab",
    "SP-2026-003": "A Phase 2b Dose-Finding Study of Solzumab",
    "SP-2026-004": "A Phase 3 Pivotal Efficacy Study of Solzumab",
}

# Per-study enrolment used to keep numbers internally consistent across docs.
ENROLLED = {"SP-2026-001": 64, "SP-2026-002": 240, "SP-2026-003": 360, "SP-2026-004": 612}


# ─────────────────────────────────────────────────────────────────────────────
# Primitives
# ─────────────────────────────────────────────────────────────────────────────


def _slug(study: str) -> str:
    return study.lower()


def _others(study: str) -> list[str]:
    return [s for s in STUDIES if s != study]


def _nct(study: str) -> str:
    n = study.split("-")[-1]
    return f"NCT0482799{int(n)}"


def _table(caption: str, rows: list[list[str]]) -> dict:
    """A table block: a caption paragraph followed by a real Word table."""
    return {"caption": caption, "rows": rows}


def _para(*sentences: str) -> str:
    """Join sentences into one paragraph (keeps content builders readable)."""
    return " ".join(s.strip() for s in sentences if s.strip())


# Canonical, resolvable cross-reference sentences distributed through the bulk
# content so hyperlinks appear across the FULL length of every document rather
# than clustering in the first half. Every token uses a form that the detection
# and target-resolution layers route to a real document / section / table.
_XREF_TEMPLATES = [
    "Full results for this topic are reported in CSR {s} Section 2.7 and summarized in Table 14.3.1.2, "
    "with the analysis populations defined in SAP {s} Section 4.2.",
    "The corresponding study design is described in Protocol {s} Section 6.1, and the primary efficacy "
    "output appears in Table 14.2.1.1 of CSR {s}.",
    "These data should be read together with CSR {o0} Section 2.7.4 and CSR {o1} Section 5.3.5, "
    "consistent with the pooling methodology in CSR {o2} Section 2.5.3.",
    "Subject-level source data supporting this section are provided in Listings {s} Section 16.2 and in "
    "Listing 16.2.5; aggregate demographics appear in Table 14.1.1.1.",
    "Statistical handling is specified in SAP {s} Section 5.1, with the testing hierarchy in Table 5.1.1 "
    "and sensitivity analyses cross-referenced to Protocol {s} Section 10.2.",
    "Safety monitoring is governed by Protocol {s} Section 9.5 and reconciled in CSR {s} Section 2.7; the "
    "deviation inventory is documented in Appendix 16.1.",
    "Exposure and compliance reconcile with Listings {s} Section 16.2, while laboratory shifts are "
    "summarized in Table 14.3.4.1 of CSR {s}.",
    "The estimand and analysis populations trace to SAP {s} Section 4.2 and Table 4.2.1, and subject "
    "disposition is given in Table 14.1.2.1 of CSR {s}.",
]


def _xref_para(study: str, i: int) -> str:
    """One realistic cross-reference paragraph (rotates through the templates)."""
    o = _others(study)
    return _para(_XREF_TEMPLATES[i % len(_XREF_TEMPLATES)].format(s=study, o0=o[0], o1=o[1], o2=o[2]))


# Topical narrative passages used to bulk every document past ten printed pages.
# Each entry is (heading, [paragraphs]); a cross-reference paragraph is appended
# to every topic so the hyperlinks span the whole document, not just the front.
def _expansion(study: str, section_prefix: str, topics: list[tuple[str, list[str]]]) -> list:
    out: list = []
    for i, (heading, bodies) in enumerate(topics):
        out.append("")
        out.append(f"## {heading}")
        for b in bodies:
            out.append(_para(b.format(study=study)))
        out.append(_xref_para(study, i))
    return out


_CSR_EXPANSION = [
    ("3.1 Investigational Product and Administration", [
        "Solzumab drug product was supplied as a sterile liquid in single-use prefilled syringes at the "
        "labeled strength. Study drug was stored at 2–8 °C, protected from light, and administered "
        "subcutaneously by trained site personnel. Accountability records were maintained for every unit "
        "dispensed and returned, and reconciliation was performed at each monitoring visit. Temperature "
        "excursions were documented and adjudicated for impact on product quality before further use.",
        "Matching placebo was identical in appearance, volume, and packaging to maintain the double blind. "
        "Randomization codes were held by the interactive response technology vendor and were not accessible "
        "to site staff, sponsor study team members, or subjects. Emergency unblinding was available through "
        "the system on a per-subject basis and was not invoked during the conduct of the study.",
    ]),
    ("3.2 Compliance and Treatment Exposure", [
        "Treatment compliance was calculated as the number of administered doses divided by the number of "
        "expected doses over the treatment period. Mean compliance exceeded ninety-five percent in all "
        "treatment groups, and the proportion of subjects with at least eighty percent compliance was "
        "comparable across arms. Exposure was summarized as total duration on study drug and cumulative "
        "dose, with no meaningful between-group differences observed.",
        "Dose interruptions and reductions were infrequent and were driven primarily by transient, "
        "mild-to-moderate adverse events that resolved without sequelae. The exposure achieved in this "
        "study supports the interpretation of both the efficacy and the safety findings as representative "
        "of the intended therapeutic regimen.",
    ]),
    ("4.1 Efficacy Evaluation Methods", [
        "The primary efficacy variable was assessed by trained, blinded raters using a validated, "
        "disease-specific instrument administered at each scheduled visit. Rater training and "
        "certification were completed before enrollment, and inter-rater reliability was monitored "
        "throughout the study to minimize measurement variability. Assessments performed outside the "
        "protocol-defined window were flagged and handled according to pre-specified rules.",
        "Responder status was derived programmatically from the recorded assessment scores using the "
        "algorithm fixed before database lock. Continuous endpoints were analyzed using mixed models for "
        "repeated measures with an unstructured covariance matrix, providing valid inference under the "
        "missing-at-random assumption while remaining robust to monotone dropout patterns.",
    ]),
    ("4.3 Pharmacokinetic and Immunogenicity Results", [
        "Serum concentrations of Solzumab increased with dose and reached steady state within the "
        "treatment period. Exposure metrics, including area under the concentration-time curve and maximum "
        "observed concentration, were approximately dose-proportional across the studied range. "
        "Inter-subject variability was moderate and consistent with expectations for a monoclonal antibody.",
        "The incidence of treatment-emergent anti-drug antibodies was low, and titers were generally "
        "transient and of low magnitude. No clear association was observed between anti-drug antibody "
        "status and either efficacy or the incidence of hypersensitivity or injection-site reactions, "
        "supporting a favorable immunogenicity profile for the investigational product.",
    ]),
    ("4.4 Health-Related Quality of Life", [
        "Patient-reported outcomes were collected using validated questionnaires completed by subjects "
        "before any other study procedure to avoid influence from clinical assessments. Completion rates "
        "were high across all visits, and missing items were handled according to each instrument's "
        "scoring manual. Improvements favored the active treatment arms and were consistent with the "
        "clinician-rated efficacy findings.",
        "Domain-level analyses indicated that improvements were broad-based rather than driven by a single "
        "component, encompassing physical functioning, role limitations, and overall well-being. These "
        "findings support the clinical relevance of the observed treatment effect from the patient "
        "perspective.",
    ]),
    ("6.3 Benefit-Risk Assessment", [
        "The totality of evidence from this study demonstrates a clinically meaningful treatment benefit "
        "with an acceptable and manageable safety profile. The magnitude of the responder-rate difference "
        "versus placebo, the dose-ordered response, and the consistency of secondary and patient-reported "
        "endpoints together establish a favorable benefit-risk balance in the studied population.",
        "Identified risks were predominantly mild-to-moderate, transient, and manageable with standard "
        "clinical care. No unexpected risks emerged, and serious events were infrequent and balanced "
        "between active treatment and placebo. Routine pharmacovigilance is considered sufficient to "
        "characterize and manage the risks identified in this study.",
    ]),
]

_PROTOCOL_EXPANSION = [
    ("4. Subject Selection", [
        "Eligible subjects were adults who met the diagnostic criteria for the target indication, had "
        "disease severity above the protocol-defined threshold at screening, and were able to provide "
        "written informed consent. Key exclusion criteria included clinically significant comorbidities, "
        "recent use of prohibited therapies, and any condition that, in the investigator's judgment, would "
        "compromise subject safety or the interpretation of study data.",
        "Screening procedures confirmed eligibility within the protocol-defined window before "
        "randomization. Re-screening was permitted once for subjects who failed a reversible eligibility "
        "criterion, subject to medical monitor approval and documentation in the source record.",
    ]),
    ("5. Treatment Plan and Prior/Concomitant Therapy", [
        "Study drug was administered according to the schedule of activities for the full treatment "
        "duration. Permitted concomitant medications were recorded with indication, dose, and dates. "
        "Prohibited medications, including other investigational agents and disease-modifying therapies "
        "not specified by the protocol, were not allowed during the treatment period.",
        "Rescue medication was permitted for subjects meeting pre-defined criteria; use of rescue "
        "medication was recorded and accounted for in the efficacy analyses according to the estimand "
        "framework. Management of overdose and special situations followed documented procedures.",
    ]),
    ("8. Assessment of Efficacy", [
        "Efficacy assessments were performed by qualified, blinded raters at each scheduled visit using "
        "validated instruments. The primary and secondary endpoints, their timing, and the derivation of "
        "responder status were specified to ensure consistent, reproducible measurement across sites. "
        "Assessment windows and procedures for handling out-of-window visits were defined to preserve the "
        "integrity of the analysis.",
        "Quality-control procedures included centralized review of assessment data, automated range and "
        "consistency checks, and ongoing rater performance monitoring. These measures minimized "
        "measurement error and supported the reliability of the efficacy conclusions.",
    ]),
    ("11. Ethics and Regulatory Considerations", [
        "The study was conducted under the oversight of the responsible Institutional Review Boards or "
        "Independent Ethics Committees and in compliance with applicable national and regional "
        "regulations. Substantial amendments were submitted for review and approval before implementation, "
        "except where immediate action was required to protect subject safety.",
        "Subject confidentiality was protected throughout the study. Personal data were handled in "
        "accordance with applicable data-protection regulations, and records were retained for the period "
        "required by regulation and sponsor policy. Insurance and indemnification arrangements were in "
        "place for all participating sites.",
    ]),
    ("12. Data Handling and Quality Assurance", [
        "Data were recorded in a validated electronic data capture system with role-based access control "
        "and a complete audit trail. Edit checks, manual review, and external data reconciliation "
        "supported data quality. Independent quality-assurance audits of selected sites and of the trial "
        "master file were conducted according to a risk-based plan.",
        "Monitoring followed a risk-based strategy combining centralized statistical surveillance with "
        "targeted on-site source-data verification. Identified issues were escalated, documented, and "
        "resolved through a formal corrective and preventive action process to maintain the reliability "
        "of the trial results.",
    ]),
]

_SAP_EXPANSION = [
    ("3. General Analysis Conventions", [
        "All analyses were performed using validated statistical software under version control. Output "
        "was generated by qualified statistical programmers and independently verified through a "
        "double-programming process for all key tables, figures, and listings. Decimal precision, "
        "rounding conventions, and the presentation of visit windows were standardized across outputs.",
        "Baseline was defined as the last non-missing assessment before the first dose of study drug. "
        "Visit windowing rules mapped actual assessment dates to nominal study visits, and procedures for "
        "handling unscheduled and repeated assessments were specified to ensure a single value per "
        "subject per nominal visit in the analyses.",
    ]),
    ("7. Safety Analysis Details", [
        "Safety analyses were descriptive and performed on the safety population. Treatment-emergent "
        "adverse events were summarized by system organ class and preferred term, by severity, and by "
        "relationship to study drug. Events of special interest, serious adverse events, and events "
        "leading to discontinuation were summarized separately to support a comprehensive safety review.",
        "Laboratory data were summarized using descriptive statistics and shift tables relative to the "
        "reference ranges. Vital signs and electrocardiogram parameters were summarized as observed values "
        "and changes from baseline, with potentially clinically significant values flagged according to "
        "pre-defined criteria.",
    ]),
    ("8. Handling of Missing Data", [
        "The primary analysis treated subjects with missing Week 12 response who discontinued for lack of "
        "efficacy as non-responders, consistent with the estimand. Sensitivity analyses, including "
        "multiple imputation under both missing-at-random and missing-not-at-random assumptions, assessed "
        "the robustness of the primary conclusion to alternative missing-data mechanisms.",
        "A tipping-point analysis identified the degree of departure from the missing-at-random assumption "
        "that would be required to overturn the primary result. The pattern and extent of missing data "
        "were summarized by treatment group and visit to support interpretation of the sensitivity "
        "analyses.",
    ]),
    ("9. Interim Analyses and Data Monitoring", [
        "Pre-specified interim reviews of blinded safety data were conducted by the independent Data "
        "Safety Monitoring Board according to its charter. No formal interim efficacy analysis with "
        "alpha-spending was planned for this study, and no actions affecting the conduct or analysis of "
        "the study resulted from the safety reviews.",
        "The firewall between the unblinded statistical support for the monitoring board and the sponsor "
        "study team was maintained throughout. Documentation of board meetings, recommendations, and "
        "sponsor responses was retained in the trial master file.",
    ]),
    ("10. Changes from the Protocol-Planned Analyses", [
        "Any differences between the analyses described in this plan and those outlined in the protocol "
        "were documented with rationale before database lock and unblinding. Such changes were limited to "
        "clarifications and methodological refinements that did not alter the primary estimand or the "
        "interpretation of the confirmatory analyses.",
        "Post hoc analyses, clearly labeled as exploratory, were conducted to further characterize the "
        "treatment effect and were interpreted with appropriate caution given the absence of multiplicity "
        "control.",
    ]),
]

_LISTINGS_EXPANSION = [
    ("16.5 Protocol Deviation Listings", [
        "All important protocol deviations were identified before unblinding and classified by category, "
        "including eligibility, treatment, assessment, and prohibited-medication deviations. Each entry "
        "records the subject identifier, the deviation category, the date, and a brief description "
        "sufficient to assess potential impact on the efficacy and safety analyses.",
        "Deviations were reviewed at a blinded data review meeting to confirm population assignments and "
        "to determine the handling of affected data points. The complete deviation inventory supports the "
        "per-protocol sensitivity analysis and the overall assessment of study conduct quality.",
    ]),
    ("16.6 Exposure and Compliance Listings", [
        "Per-subject exposure listings present the dates of first and last dose, the number of doses "
        "administered, total duration on treatment, and computed compliance. Interruptions and reductions "
        "are flagged with associated reasons to allow reviewers to relate exposure to efficacy and safety "
        "outcomes at the individual level.",
        "These listings reconcile with the drug accountability records maintained at each site. Any "
        "discrepancies identified during reconciliation were investigated and resolved before database "
        "lock, with the resolution documented in the source record.",
    ]),
    ("16.7 Laboratory Data Listings", [
        "Individual laboratory results are listed by subject, parameter, and visit, with values flagged "
        "relative to the laboratory reference ranges and against potentially clinically significant "
        "thresholds. Both observed values and changes from baseline are presented to support the shift "
        "analyses and the overall laboratory safety evaluation.",
        "Repeat and unscheduled laboratory assessments are included with their actual collection dates. "
        "Hemolyzed, clotted, or otherwise non-evaluable samples are annotated so that reviewers can "
        "distinguish true findings from sample-handling artifacts.",
    ]),
    ("16.8 Vital Signs and ECG Listings", [
        "Vital sign listings present systolic and diastolic blood pressure, heart rate, respiratory rate, "
        "and temperature by subject and visit, including changes from baseline and any values meeting "
        "pre-defined alert criteria. Electrocardiogram listings include interval measurements and the "
        "overall interpretation recorded at each scheduled assessment.",
        "Clinically significant findings, where present, are cross-annotated with the corresponding "
        "adverse event entries to provide a coherent view of each subject's safety experience across data "
        "domains.",
    ]),
]


# Long reusable narrative blocks. These add bulk (page count) without disturbing
# the cross-reference sentences that drive the hyperlink detection.
_GCP_BOILERPLATE = (
    "The study was conducted in accordance with the protocol, the principles of "
    "the Declaration of Helsinki, and the International Council for Harmonisation "
    "(ICH) E6(R2) Good Clinical Practice guideline. All subjects provided written "
    "informed consent before any study-specific procedure was performed. The "
    "protocol, informed consent form, and all amendments were reviewed and "
    "approved by the Institutional Review Board or Independent Ethics Committee at "
    "each participating site. Source data were verified against case report forms "
    "by clinical monitors throughout the study, and all data management activities "
    "followed pre-specified data validation plans with documented query resolution."
)
_QUALITY_BOILERPLATE = (
    "Data were captured in a validated 21 CFR Part 11-compliant electronic data "
    "capture system. Audit trails recorded every change to a data field with the "
    "user identity, timestamp, and reason for change. Database lock occurred only "
    "after resolution of all outstanding queries, completion of medical coding to "
    "MedDRA version 27.0 and WHO Drug Global B3 dictionary, and sign-off by the "
    "study statistician, data manager, and medical monitor."
)
_STATS_BOILERPLATE = (
    "Continuous variables are summarized using the number of non-missing "
    "observations, mean, standard deviation, median, minimum, and maximum. "
    "Categorical variables are summarized using frequency counts and percentages. "
    "Confidence intervals are two-sided at the 95% level unless otherwise stated. "
    "No imputation was performed for the primary analysis; sensitivity analyses "
    "using multiple imputation and tipping-point approaches assessed the impact of "
    "missing data on the primary efficacy conclusion."
)


# ─────────────────────────────────────────────────────────────────────────────
# Writer
# ─────────────────────────────────────────────────────────────────────────────


def _write(path: Path, title: str, paragraphs: list) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for para in paragraphs:
        if isinstance(para, dict) and "rows" in para:
            if para.get("caption"):
                doc.add_paragraph(para["caption"])
            rows = para["rows"]
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                try:
                    tbl.style = "Light Grid Accent 1"
                except Exception:  # noqa: BLE001 — style is cosmetic only
                    pass
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        tbl.cell(ri, ci).text = str(val)
            continue
        if para.startswith("# "):
            doc.add_heading(para[2:], level=1)
        elif para.startswith("## "):
            doc.add_heading(para[3:], level=2)
        elif para.startswith("### "):
            doc.add_heading(para[4:], level=3)
        elif para == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(para)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


# ─────────────────────────────────────────────────────────────────────────────
# CSR body (~14-18 pages)
# ─────────────────────────────────────────────────────────────────────────────


def _csr_body(study: str) -> list[str]:
    o = _others(study)
    n = ENROLLED[study]
    return [
        f"Clinical Study Report — {STUDY_TITLES[study]} ({study}).",
        f"Sponsor: Solzumab Therapeutics. Report status: Final. ICH E3-compliant.",
        "",
        "# 1. Synopsis",
        _para(
            f"This Clinical Study Report (CSR {study}) presents the complete results of study {study},",
            f"registered as {_nct(study)} on ClinicalTrials.gov. The full study design is described in",
            f"Protocol {study} Section 6.1, and the pre-specified analyses are defined in the SAP {study}.",
        ),
        _para(_GCP_BOILERPLATE),
        _para(
            "The synopsis below summarizes objectives, design, population, treatments, endpoints, and",
            f"principal conclusions. Detailed results appear in Sections 2.3 through 5.3 of this CSR {study}.",
        ),
        _table(
            "Table 1.1 - Study Synopsis",
            [
                ["Item", "Description"],
                ["Title", STUDY_TITLES[study]],
                ["Registration", _nct(study)],
                ["Phase", study.split('-')[-1] + " (program staging)"],
                ["Design", "Randomized, double-blind, placebo-controlled, parallel-group"],
                ["Population", "Adults with the target indication"],
                ["Planned N", str(n)],
                ["Primary endpoint", "Proportion of responders at Week 12"],
                ["Sponsor", "Solzumab Therapeutics"],
            ],
        ),
        "",
        "# 2. Introduction and Study Background",
        _para(
            "Solzumab is a humanized monoclonal antibody under development for the target indication.",
            f"The scientific rationale, prior clinical experience, and benefit-risk assessment that",
            f"motivated study {study} are summarized here and detailed in the Investigator's Brochure.",
        ),
        _para(
            "This trial was conducted in accordance with ICH E6(R2) GCP and reported per ICH E3 guidance.",
            "Sponsor reporting follows FDA guidance available at",
            "https://www.fda.gov/regulatory-information/search-fda-guidance-documents and the corresponding",
            "EMA scientific guidelines. No interim database was unblinded outside the charter of the",
            "independent Data Safety Monitoring Board.",
        ),
        _para(
            f"The development program comprises four studies. CSR {study} should be interpreted alongside",
            f"CSR {o[0]}, CSR {o[1]}, and CSR {o[2]}, which together establish the dose-response, efficacy,",
            "and safety profile of Solzumab across the intended population.",
        ),
        "",
        "# 2.3 Demographic and Baseline Characteristics",
        _para(
            f"A total of {n} subjects were enrolled across multiple sites and randomized to treatment.",
            f"Refer to Listings {study} for individual subject data. Key baseline characteristics are",
            "presented in Table 14.1.1.1, and the disposition of subjects is summarized in Table 14.1.2.1.",
        ),
        _table(
            "Table 14.1.1.1 - Subject Demographics (Intent-to-Treat Population)",
            [
                ["Characteristic", "Solzumab 200 mg", "Solzumab 100 mg", "Placebo", "Total"],
                ["N randomized", str(n // 3), str(n // 3), str(n - 2 * (n // 3)), str(n)],
                ["Age (years), mean (SD)", "52.3 (12.1)", "53.1 (11.8)", "52.7 (12.4)", "52.7 (12.1)"],
                ["Age ≥ 65, n (%)", "31 (20.7)", "30 (20.1)", "29 (19.6)", "90 (20.1)"],
                ["Male, n (%)", "89 (59.3)", "88 (59.1)", "88 (59.5)", "265 (59.3)"],
                ["Race: White, n (%)", "125 (83.3)", "124 (83.2)", "123 (83.1)", "372 (83.2)"],
                ["BMI (kg/m²), mean (SD)", "27.9 (4.2)", "28.1 (4.4)", "27.8 (4.0)", "27.9 (4.2)"],
                ["Baseline severity score, mean", "41.8", "42.1", "41.9", "41.9"],
            ],
        ),
        _table(
            "Table 14.1.2.1 - Subject Disposition",
            [
                ["Disposition", "Solzumab 200 mg", "Solzumab 100 mg", "Placebo"],
                ["Randomized", str(n // 3), str(n // 3), str(n - 2 * (n // 3))],
                ["Completed Week 12", "138", "136", "134"],
                ["Discontinued — adverse event", "5", "6", "4"],
                ["Discontinued — withdrawal of consent", "4", "5", "6"],
                ["Discontinued — lost to follow-up", "3", "2", "4"],
            ],
        ),
        _para(
            "Treatment groups were well balanced for all demographic and baseline disease characteristics.",
            "No clinically meaningful differences were observed between arms, supporting the validity of the",
            f"randomized comparison defined in SAP {study} Section 4.2.",
        ),
        "",
        "# 2.5 Clinical Overview and Efficacy",
        _para(
            f"Efficacy results are summarized in Section 2.5.3 and Table 14.2.1.1. Subject-level source data",
            f"are provided in Listings {study}. The primary endpoint analysis compares all treatment arms",
            "against placebo using logistic regression with treatment, baseline severity, and region as",
            "covariates, consistent with the model pre-specified in SAP {0} Section 5.1.".format(study),
        ),
        _para(_STATS_BOILERPLATE),
        _table(
            "Table 14.2.1.1 - Primary Efficacy Summary (ITT Population)",
            [
                ["Treatment Arm", "N", "Responders (%)", "95% CI", "p-value"],
                ["Solzumab 200 mg", str(n // 3), "92 (61.3%)", "(53.0–69.6%)", "0.0003"],
                ["Solzumab 100 mg", str(n // 3), "77 (51.7%)", "(43.3–60.0%)", "0.0156"],
                ["Placebo", str(n - 2 * (n // 3)), "56 (37.8%)", "(29.8–46.2%)", "—"],
            ],
        ),
        "## 2.5.3 Secondary and Exploratory Efficacy",
        _para(
            "Secondary endpoints included time to first response, change from baseline in the continuous",
            "severity score at Week 12, and patient-reported quality of life. The hierarchy of testing and",
            f"multiplicity control are specified in SAP {study} Section 5.1 and were followed without",
            "deviation. Results were directionally consistent with the primary endpoint across all arms.",
        ),
        _table(
            "Table 14.2.2.1 - Secondary Efficacy Endpoints (ITT Population)",
            [
                ["Endpoint", "Solzumab 200 mg", "Solzumab 100 mg", "Placebo"],
                ["Median time to response (days)", "29", "38", "61"],
                ["Δ severity score, LS mean", "-21.4", "-16.9", "-9.2"],
                ["QoL change, LS mean", "+14.1", "+10.8", "+5.3"],
                ["Responders sustained to Wk 16, %", "57.8", "47.2", "33.1"],
            ],
        ),
        _para(
            f"For the pooled safety population, cross-reference CSR {o[0]} Section 2.7 and CSR {o[1]}",
            "Section 2.7.4. Subgroup analyses by age, sex, and baseline severity appear in Appendix 16.1 and",
            "did not reveal subgroups without benefit.",
        ),
        "",
        "# 2.7 Clinical Summary and Safety Analysis",
        _para(
            f"Integrated efficacy follows the methodology of CSR {o[2]} Section 5.3.5. Adverse events are",
            f"tabulated in Table 14.3.1.2 and detailed in Listings {study}. The safety population consists",
            "of all subjects who received at least one dose of study drug.",
        ),
        _para(
            "The overall incidence of treatment-emergent adverse events (TEAEs) was higher in the active",
            "arms than placebo but was driven predominantly by mild-to-moderate events. No new safety signal",
            f"was identified relative to CSR {o[0]} and CSR {o[1]}. Serious adverse events were infrequent",
            "and balanced across arms.",
        ),
        _table(
            "Table 14.3.1.2 - Treatment-Emergent Adverse Events by System Organ Class",
            [
                ["System Organ Class", "Solzumab (n=299)", "Placebo (n=148)", "Total (n=447)"],
                ["Any TEAE", "223 (74.6%)", "98 (66.2%)", "321 (71.8%)"],
                ["Gastrointestinal disorders", "72 (24.1%)", "31 (20.9%)", "103 (23.1%)"],
                ["Infections and infestations", "51 (17.1%)", "22 (14.9%)", "73 (16.3%)"],
                ["Nervous system disorders", "38 (12.7%)", "14 (9.5%)", "52 (11.6%)"],
                ["Skin and subcutaneous tissue", "29 (9.7%)", "11 (7.4%)", "40 (8.9%)"],
                ["Serious AEs", "12 (4.0%)", "7 (4.7%)", "19 (4.3%)"],
            ],
        ),
        _table(
            "Table 14.3.1.3 - Treatment-Related Adverse Events (Safety Population)",
            [
                ["Event", "Solzumab (n=299)", "Placebo (n=148)"],
                ["Related TEAE", "156 (52.2%)", "31 (20.9%)"],
                ["Mild–Moderate", "142 (47.5%)", "28 (18.9%)"],
                ["Severe", "14 (4.7%)", "3 (2.0%)"],
                ["Leading to discontinuation", "8 (2.7%)", "2 (1.4%)"],
            ],
        ),
        _table(
            "Table 14.3.2.1 - Most Frequent TEAEs by Preferred Term (≥5% any arm)",
            [
                ["Preferred Term", "Solzumab 200 mg", "Solzumab 100 mg", "Placebo"],
                ["Headache", "27 (18.0%)", "24 (16.1%)", "15 (10.1%)"],
                ["Nausea", "21 (14.0%)", "19 (12.8%)", "9 (6.1%)"],
                ["Nasopharyngitis", "18 (12.0%)", "17 (11.4%)", "14 (9.5%)"],
                ["Injection-site reaction", "16 (10.7%)", "13 (8.7%)", "3 (2.0%)"],
                ["Fatigue", "12 (8.0%)", "11 (7.4%)", "8 (5.4%)"],
            ],
        ),
        _para(_QUALITY_BOILERPLATE),
        _para(
            f"Deaths and serious adverse events are reconciled against Protocol {study} Section 9.5 and the",
            f"SAP {study} Section 4.2. No deaths were considered related to study drug by the investigator or",
            "the sponsor's medical monitor.",
        ),
        "## 2.7.4 Laboratory, Vital Signs, and ECG Evaluations",
        _para(
            "Clinical laboratory parameters, vital signs, and 12-lead ECG findings were assessed at the",
            f"timepoints defined in Protocol {study} Section 9.5. Mean changes from baseline were small and",
            "not clinically meaningful, and shift-table analyses revealed no treatment-related trends.",
        ),
        _table(
            "Table 14.3.4.1 - Selected Laboratory Shifts (Normal → High), Safety Population",
            [
                ["Parameter", "Solzumab (n=299)", "Placebo (n=148)"],
                ["ALT", "9 (3.0%)", "4 (2.7%)"],
                ["AST", "7 (2.3%)", "3 (2.0%)"],
                ["Creatinine", "4 (1.3%)", "2 (1.4%)"],
                ["Total bilirubin", "2 (0.7%)", "1 (0.7%)"],
            ],
        ),
        "",
        "# 5.3 Clinical Study Reports — Cross-Program Integration",
        _para(
            f"This CSR {study} should be read together with CSR {o[0]}, CSR {o[1]}, and CSR {o[2]} for the",
            "complete clinical picture of the Solzumab program. Cross-references within this report are",
            "denoted as 'Section X.Y.Z' or 'Table N.N.N.N' and redirect to the specific location.",
        ),
        _para(
            "The integrated summary of efficacy and integrated summary of safety draw on the pooled data",
            f"from all four studies. Methodology for pooling is harmonized with CSR {o[2]} Section 5.3.5 and",
            f"the analysis populations defined in SAP {study} Section 4.2.",
        ),
        _table(
            "Table 5.3.1 - Cross-Study Reference Map",
            [
                ["This study references", "Target document", "Location"],
                ["Study design", f"Protocol {study}", "Section 6.1"],
                ["Analysis populations", f"SAP {study}", "Section 4.2"],
                ["Subject-level data", f"Listings {study}", "Section 16.2"],
                ["Pooled safety", f"CSR {o[0]}", "Section 2.7"],
                ["Pooled safety (alt)", f"CSR {o[1]}", "Section 2.7.4"],
                ["Integrated efficacy", f"CSR {o[2]}", "Section 5.3.5"],
            ],
        ),
        "",
        "# 6. Discussion and Overall Conclusions",
        _para(
            f"Study {study} met its primary objective. Solzumab demonstrated a statistically significant and",
            "clinically meaningful improvement in the responder rate versus placebo, with a dose-ordered",
            "response and an acceptable safety profile. These findings are consistent across the development",
            f"program and support the conclusions integrated in CSR {o[2]}.",
        ),
        _para(
            "Limitations include the controlled trial setting and the 12-week primary assessment window.",
            "Long-term durability is being characterized in the open-label extension. Taken together, the",
            "benefit-risk balance of Solzumab in the studied population is favorable.",
        ),
    ] + _expansion(study, "", _CSR_EXPANSION)


# ─────────────────────────────────────────────────────────────────────────────
# Protocol (~12-15 pages)
# ─────────────────────────────────────────────────────────────────────────────


def _protocol(study: str) -> list[str]:
    return [
        f"Clinical Trial Protocol — {study}.",
        f"Protocol title: {STUDY_TITLES[study]}.",
        "",
        "# 1. Protocol Summary",
        _para(
            f"Protocol {study} is a randomized, double-blind, placebo-controlled trial of Solzumab in",
            "patients with the target condition. The study design, safety monitoring, and statistical",
            "approach are detailed in the sections below and operationalized through the study manual.",
        ),
        _para(_GCP_BOILERPLATE),
        _table(
            "Table 1.2 - Schedule of Activities (abbreviated)",
            [
                ["Procedure", "Screening", "Baseline", "Wk 2", "Wk 4", "Wk 8", "Wk 12", "Follow-up"],
                ["Informed consent", "X", "", "", "", "", "", ""],
                ["Randomization", "", "X", "", "", "", "", ""],
                ["Study drug administration", "", "X", "X", "X", "X", "X", ""],
                ["Efficacy assessment", "", "X", "X", "X", "X", "X", "X"],
                ["Safety labs", "X", "X", "", "X", "X", "X", "X"],
                ["PK sampling", "", "X", "X", "X", "", "X", ""],
            ],
        ),
        "",
        "# 2. Background and Rationale",
        _para(
            "The unmet medical need, mechanism of action of Solzumab, and prior nonclinical and clinical",
            f"experience justifying study {study} are summarized here. Dose selection is informed by the",
            f"exposure-response relationship characterized in the earlier program studies.",
        ),
        _para(
            "The benefit-risk assessment supporting initiation of this study considered all available safety",
            "data. Risk-mitigation measures, including stopping rules and DSMB oversight, are described in",
            "Section 9.5.",
        ),
        "",
        "# 3. Objectives and Endpoints",
        _table(
            "Table 3.1 - Objectives and Corresponding Endpoints",
            [
                ["Objective", "Endpoint", "Timepoint"],
                ["Primary efficacy", "Proportion of responders", "Week 12"],
                ["Key secondary", "Time to first response", "Through Week 12"],
                ["Secondary", "Change in severity score", "Week 12"],
                ["Safety", "Incidence of TEAEs and SAEs", "Throughout"],
                ["PK", "AUC, Cmax, trough concentration", "Per PK schedule"],
            ],
        ),
        _para(
            f"Endpoints align with the analysis defined in SAP {study}. Results are reported in CSR {study}",
            "Section 2.5 (efficacy) and Section 2.7 (safety).",
        ),
        "",
        "# 6.1 Study Design",
        _para(
            f"This protocol governs study {study}. Subjects are randomized 1:1:1 to Solzumab 200 mg,",
            "Solzumab 100 mg, or matching placebo, administered subcutaneously per the schedule of",
            f"activities. Endpoints align with the analysis defined in SAP {study}, and results are reported",
            f"in CSR {study} Section 2.5.",
        ),
        _table(
            "Table 6.1.1 - Study Design Parameters",
            [
                ["Parameter", "Value"],
                ["Type", "Randomized, double-blind, placebo-controlled"],
                ["Allocation", "1:1:1 (200 mg : 100 mg : placebo)"],
                ["Duration", "12 weeks treatment + 4 weeks follow-up"],
                ["Primary Endpoint", "Proportion of responders at Week 12"],
                ["Secondary Endpoints", "Time to response, quality of life, safety"],
                ["Target N", f"~{ENROLLED[study]} (power analysis per SAP {study})"],
            ],
        ),
        "## 6.2 Randomization and Blinding",
        _para(
            "Randomization is performed via a validated interactive response technology system using",
            "permuted blocks stratified by region and baseline severity. The sponsor, investigators, and",
            f"subjects remain blinded until database lock as specified in SAP {study} Section 6.",
        ),
        "",
        "# 7. Pharmacokinetics",
        _para(
            f"PK sampling and analysis are specified in Section 7.2 of this protocol and cross-referenced in",
            f"SAP {study} Section 5.1. All subjects provide PK samples at protocol-defined timepoints. Bioanalysis",
            "uses a validated immunoassay with a documented lower limit of quantification.",
        ),
        "## 7.2 PK Sampling Schedule and Bioanalysis",
        _table(
            "Table 7.2.1 - Pharmacokinetic Sampling Schedule",
            [
                ["Visit", "Nominal Time", "Sample Type"],
                ["Baseline", "Pre-dose", "Serum"],
                ["Week 2", "Pre-dose + 4 h", "Serum"],
                ["Week 4", "Pre-dose", "Serum"],
                ["Week 12", "Pre-dose + 4 h", "Serum"],
                ["Follow-up", "Single", "Serum"],
            ],
        ),
        "",
        "# 9. Safety Assessments and Monitoring",
        _para(
            f"Safety data feed Table 14.3.1.2 of CSR {study} and the subject Listings {study}. Adverse",
            "events are graded by severity and assessed for causality by the investigator. Pregnancies and",
            "events of special interest follow expedited reporting procedures.",
        ),
        "## 9.5 Data Safety Monitoring and Stopping Rules",
        _para(
            f"An independent Data Safety Monitoring Board meets every 8 weeks to review blinded safety data",
            "and pre-specified interim efficacy signals. Stopping rules for individual subjects and for the",
            f"study as a whole are defined here and reconciled in CSR {study} Section 2.7.",
        ),
        _table(
            "Table 9.5.1 - Safety Assessment Schedule",
            [
                ["Assessment", "Screening", "Baseline", "Weekly", "End of Study"],
                ["Physical Exam", "X", "X", "", "X"],
                ["Clinical Lab", "X", "X", "X", "X"],
                ["AE Monitoring", "", "X", "X", "X"],
                ["ECG", "X", "", "", "X"],
                ["Vital signs", "X", "X", "X", "X"],
            ],
        ),
        "",
        "# 10. Statistical Considerations",
        _para(
            f"The sample size, analysis populations, and primary analysis model are specified in SAP {study}",
            "Section 6.1 and Section 5.1. The study is powered at 90% to detect the pre-specified treatment",
            "difference in the responder rate at a two-sided alpha of 0.05.",
        ),
        "## 10.2 Handling of Protocol Deviations and Amendments",
        _para(
            "Important protocol deviations are identified before unblinding and summarized by category.",
            f"Sensitivity analyses excluding deviating subjects are pre-specified in SAP {study} Section 6.2.",
            f"Cross-references to CSR {study} Appendix 16.1 document the final deviation inventory.",
        ),
    ] + _expansion(study, "", _PROTOCOL_EXPANSION)


# ─────────────────────────────────────────────────────────────────────────────
# SAP (~11-14 pages)
# ─────────────────────────────────────────────────────────────────────────────


def _sap(study: str) -> list[str]:
    return [
        f"Statistical Analysis Plan — {study}.",
        f"Associated protocol: Protocol {study}. Associated report: CSR {study}.",
        "",
        "# 1. Overview and Objectives",
        _para(
            f"This Statistical Analysis Plan (SAP {study}) defines all statistical analyses for protocol",
            f"{study}. It was finalized before database lock and approved by the independent statistician.",
            "Key analysis populations, methods, estimands, and outputs are defined below.",
        ),
        _para(_STATS_BOILERPLATE),
        "",
        "# 2. Estimands and Analysis Framework",
        _para(
            "The primary estimand targets the difference between Solzumab and placebo in the proportion of",
            "responders at Week 12 in the ITT population, treating discontinuation due to lack of efficacy as",
            "non-response. Intercurrent events and their handling strategies are tabulated below.",
        ),
        _table(
            "Table 2.1 - Estimand Attributes",
            [
                ["Attribute", "Specification"],
                ["Population", "Intention-to-treat"],
                ["Variable", "Responder at Week 12 (yes/no)"],
                ["Intercurrent events", "Treatment discontinuation, rescue medication"],
                ["Strategy", "Composite (treatment-policy for safety)"],
                ["Summary measure", "Risk difference vs placebo, 95% CI"],
            ],
        ),
        "",
        "# 4.2 Analysis Populations",
        _para(
            f"Analyses defined here are reported in CSR {study} Section 2.7. The design assumptions trace to",
            f"Protocol {study} Section 6.1. Subjects are assigned to populations before unblinding.",
        ),
        _table(
            "Table 4.2.1 - Definition of Analysis Populations",
            [
                ["Population", "Description", "N (est.)"],
                ["Intention-to-Treat (ITT)", "All randomized subjects", f"~{ENROLLED[study]}"],
                ["Per-Protocol (PP)", "Completers with ≥80% compliance", "~400"],
                ["Safety (SAF)", "All subjects who received ≥1 dose", f"~{ENROLLED[study]}"],
                ["PK-evaluable", "All with valid PK samples", "~440"],
            ],
        ),
        "",
        "# 5.1 Statistical Analysis Methods",
        _para(
            f"Outputs map to Table 14.2.1.1 and Table 14.3.1.2 of CSR {study}, with supporting subject-level",
            f"data in Listings {study}. The primary analysis uses logistic regression with treatment,",
            "baseline severity, and region as covariates.",
        ),
        _table(
            "Table 5.1.1 - Primary and Secondary Endpoint Analyses",
            [
                ["Endpoint", "Population", "Method", "Significance Level"],
                ["Primary (responder rate)", "ITT", "Logistic regression vs placebo", "α=0.05"],
                ["Secondary (time to response)", "ITT", "Kaplan-Meier, log-rank test", "Exploratory"],
                ["Safety (TEAE incidence)", "SAF", "Descriptive (% and 95% CI)", "Exploratory"],
                ["PK (AUC, Cmax)", "PK-evaluable", "ANOVA, dose-linearity", "Exploratory"],
            ],
        ),
        "## 5.2 Multiplicity and Testing Hierarchy",
        _para(
            "A fixed-sequence hierarchical procedure controls the family-wise error rate across the primary",
            "and key secondary endpoints. Testing proceeds only while each preceding null hypothesis is",
            "rejected at the two-sided 0.05 level.",
        ),
        _table(
            "Table 5.2.1 - Testing Hierarchy",
            [
                ["Step", "Hypothesis", "Condition to proceed"],
                ["1", "200 mg vs placebo (primary)", "Reject at 0.05"],
                ["2", "100 mg vs placebo (primary)", "Step 1 rejected"],
                ["3", "200 mg time-to-response", "Step 2 rejected"],
            ],
        ),
        "",
        "# 6. Sample Size and Sensitivity Analyses",
        _para(
            f"The sample size derivation, sensitivity analyses per protocol amendment, and baseline-risk",
            f"stratification subgroups are detailed in Section 6.2. Cross-references to Protocol {study}",
            f"Section 10.2 and CSR {study} Appendix 16.1 document the final analysis decisions.",
        ),
        "## 6.1 Sample Size Justification",
        _para(
            f"With {ENROLLED[study]} subjects randomized 1:1:1, the study provides 90% power to detect a 20",
            "percentage-point difference in responder rate at a two-sided alpha of 0.05, allowing for 10%",
            "non-evaluable subjects.",
        ),
        "## 6.2 Sensitivity and Subgroup Analyses",
        _table(
            "Table 6.2.1 - Pre-specified Sensitivity Analyses",
            [
                ["Analysis", "Purpose"],
                ["Per-protocol re-analysis", "Robustness to compliance"],
                ["Multiple imputation", "Missing-data robustness"],
                ["Tipping-point", "Worst-case missingness"],
                ["Region subgroup", "Consistency across geography"],
            ],
        ),
    ] + _expansion(study, "", _SAP_EXPANSION)


# ─────────────────────────────────────────────────────────────────────────────
# Listings (~11-14 pages)
# ─────────────────────────────────────────────────────────────────────────────


def _listings(study: str) -> list[str]:
    rows_eff = [
        ["Subject ID", "Treatment", "Baseline", "Week 4", "Week 12", "Responder"],
    ]
    # Build a long efficacy listing so the document spans many pages.
    base = int(study[-3:])
    for i in range(1, 31):
        arm = ["Solzumab 200 mg", "Solzumab 100 mg", "Placebo"][i % 3]
        b = 40 + (i % 7)
        w4 = b - (12 if "200" in arm else 8 if "100" in arm else 3)
        w12 = b - (26 if "200" in arm else 19 if "100" in arm else 6)
        resp = "Yes" if (b - w12) >= 15 else "No"
        rows_eff.append([f"S{base:03d}-{i:03d}", arm, str(b), str(w4), str(w12), resp])

    rows_ae = [["Subject ID", "Preferred Term", "Severity", "Related", "Outcome"]]
    terms = ["Headache", "Nausea", "Nasopharyngitis", "Injection-site reaction", "Fatigue", "Dizziness"]
    for i in range(1, 25):
        rows_ae.append([
            f"S{base:03d}-{i:03d}",
            terms[i % len(terms)],
            ["Mild", "Moderate", "Severe"][i % 3],
            ["Yes", "No"][i % 2],
            ["Recovered", "Recovering", "Ongoing"][i % 3],
        ])

    return [
        f"Subject Data Listings — {study}.",
        f"Supporting report: CSR {study}. Analysis plan: SAP {study}.",
        "",
        "# 16.1 Baseline and Demographics Listings",
        _para(
            f"Individual subject demographics and baseline characteristics supporting Table 14.1.1.1 in CSR",
            f"{study}. Data are sorted by randomization order and treatment group within CSR {study} Section",
            "2.3. These listings constitute the source for the demographic summaries.",
        ),
        _table(
            "Listing 16.1.1 - Subject Demographics (excerpt)",
            [
                ["Subject ID", "Age", "Sex", "Race", "Region", "Arm"],
                [f"S{base:03d}-001", "54", "M", "White", "NA", "Solzumab 200 mg"],
                [f"S{base:03d}-002", "61", "F", "Asian", "EU", "Placebo"],
                [f"S{base:03d}-003", "47", "M", "Black", "NA", "Solzumab 100 mg"],
                [f"S{base:03d}-004", "59", "F", "White", "ROW", "Solzumab 200 mg"],
                [f"S{base:03d}-005", "52", "M", "White", "EU", "Placebo"],
            ],
        ),
        "## 16.1.2 Disposition Listing",
        _para(
            "Per-subject disposition including randomization date, last dose, completion status, and reason",
            f"for discontinuation. Aggregate disposition is summarized in CSR {study} Table 14.1.2.1.",
        ),
        "",
        "# 16.2 Subject Efficacy Listings",
        _para(
            f"These listings support CSR {study} Section 2.5 and the efficacy tables (Table 14.2.1.1) defined",
            f"in SAP {study} Section 5.1. Individual subject efficacy assessments by visit and responder",
            "status are included for the full intent-to-treat population.",
        ),
        _table("Listing 16.2.1 - Subject Efficacy Assessment Summary", rows_eff),
        _para(
            "Responder status is derived per the algorithm in SAP {0} Section 5.1. Subjects discontinuing for".format(study),
            "lack of efficacy are counted as non-responders in the primary analysis.",
        ),
        "",
        "# 16.3 Subject Safety Listings",
        _para(
            f"Adverse events by subject, with MedDRA coding to version 27.0 as specified in SAP {study}",
            f"Section 5.1. Complete SAE narratives and follow-up outcomes are included. Table 14.3.1.2 in CSR",
            f"{study} summarizes incidence by system organ class.",
        ),
        _table("Listing 16.2.5 - Adverse Event Listing (excerpt)", rows_ae),
        "## 16.3.2 Serious Adverse Event Narratives",
        _para(
            "Each serious adverse event is accompanied by a narrative describing onset, course, treatment,",
            f"causality assessment, and outcome. Narratives reconcile with Protocol {study} Section 9.5 and",
            f"the safety summary in CSR {study} Section 2.7.",
        ),
        "",
        "# 16.4 Concomitant Medication Listings",
        _table(
            "Listing 16.4.1 - Concomitant Medications (excerpt)",
            [
                ["Subject ID", "Medication", "Indication", "Ongoing"],
                [f"S{base:03d}-001", "Acetaminophen", "Headache", "No"],
                [f"S{base:03d}-002", "Lisinopril", "Hypertension", "Yes"],
                [f"S{base:03d}-003", "Ibuprofen", "Myalgia", "No"],
                [f"S{base:03d}-004", "Atorvastatin", "Dyslipidemia", "Yes"],
            ],
        ),
    ] + _expansion(study, "", _LISTINGS_EXPANSION)


# ─────────────────────────────────────────────────────────────────────────────
# Ambiguous-reference block (opt-in via --ambiguous)
# ─────────────────────────────────────────────────────────────────────────────


def _ambiguous_block(study: str) -> list[str]:
    o = _others(study)
    return [
        "",
        "# 9.9 Supplementary Cross-References",
        _para(
            f"For the dosing rationale, see 2.5.3 and the methods described in 5.3.5 of the broader Solzumab",
            f"program ({study}).",
        ),
        _para(
            f"Integrated safety for the pooled population is summarized per 14.2.1 and should be read",
            f"together with the {o[0]} findings noted earlier.",
        ),
        _para(
            "Additional exploratory analyses referenced as 6.2 and 16.1 provide supportive context but are",
            "not part of the confirmatory testing hierarchy.",
        ),
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Shared regulatory appendix — appended to EVERY document to guarantee that each
# file runs well past ten printed pages. The prose is generic enough to read
# naturally at the end of any clinical/regulatory document and contains no
# cross-reference tokens, so it does not affect hyperlink detection counts.
# ─────────────────────────────────────────────────────────────────────────────

_COMMON_APPENDIX_TOPICS = [
    ("Appendix A. Study Administrative Structure", [
        "The study was sponsored by Solzumab Therapeutics and conducted with the support of a contract "
        "research organization responsible for site management, monitoring, data management, and "
        "pharmacovigilance services. Roles and responsibilities were defined in a detailed responsibility "
        "matrix agreed before study start and governed by quality agreements between the sponsor and each "
        "vendor. Oversight of vendor performance was maintained through periodic governance meetings, "
        "key-performance-indicator review, and for-cause and routine audits conducted under a risk-based "
        "quality-management plan.",
        "A central project team coordinated cross-functional activities, including clinical operations, "
        "biostatistics, statistical programming, medical writing, regulatory affairs, and drug supply. "
        "Decision authority, escalation pathways, and communication plans were documented to ensure timely "
        "resolution of operational issues. The trial master file was maintained contemporaneously and was "
        "subject to periodic completeness review to support inspection readiness throughout the conduct of "
        "the study.",
    ]),
    ("Appendix B. Investigators, Sites, and Committees", [
        "The study was conducted at multiple investigative sites selected on the basis of access to the "
        "target population, qualified and experienced investigators, adequate facilities, and a "
        "demonstrated record of regulatory compliance. Each principal investigator was responsible for the "
        "conduct of the study at the site, the supervision of delegated study staff, and the integrity of "
        "the data generated. A delegation log documented the tasks assigned to each qualified team member.",
        "An independent Data Safety Monitoring Board provided ongoing oversight of subject safety according "
        "to a charter agreed before enrollment. The board comprised clinicians and a statistician without "
        "other involvement in the study. Where applicable, an independent adjudication committee reviewed "
        "pre-specified endpoints and events using charter-defined criteria, blinded to treatment assignment, "
        "to ensure consistent and unbiased classification across sites.",
    ]),
    ("Appendix C. Informed Consent and Ethical Conduct", [
        "Before any study-specific procedure, the nature, objectives, potential risks, and anticipated "
        "benefits of the study were explained to each prospective subject in language they could "
        "understand, and adequate time was provided to consider participation and to ask questions. Written "
        "informed consent was obtained and documented in the source record. Subjects were informed of their "
        "right to withdraw at any time without penalty or loss of benefits to which they were otherwise "
        "entitled.",
        "The study adhered to the ethical principles originating in the Declaration of Helsinki and to all "
        "applicable regulatory requirements governing the protection of human subjects. Revisions to the "
        "informed consent form were reviewed and approved by the responsible ethics committees, and "
        "affected subjects were re-consented as required. Vulnerable-population protections, where relevant, "
        "were applied in accordance with the protocol and local regulation.",
    ]),
    ("Appendix D. Data Management and Quality Control", [
        "Clinical data were collected using a validated electronic data capture system configured with "
        "programmed edit checks to identify out-of-range, inconsistent, or missing values at the point of "
        "entry. A data management plan defined the data flow, query-management process, external-data "
        "reconciliation, and coding conventions. Medical history, adverse events, and concomitant "
        "medications were coded using standardized dictionaries to support consistent aggregation and "
        "analysis.",
        "Quality control encompassed source-data verification proportionate to risk, central monitoring of "
        "data trends and site performance, and independent review of critical data. Database lock followed "
        "completion of data cleaning, resolution of outstanding queries, finalization of coding, and formal "
        "sign-off by the responsible functions. Post-lock changes, if any, were governed by a documented "
        "unlock-and-relock procedure with full justification and audit trail.",
    ]),
    ("Appendix E. Statistical Software and Reproducibility", [
        "All statistical analyses were produced using validated, version-controlled software environments. "
        "Analysis datasets and outputs conformed to recognized data standards to facilitate review and "
        "reproducibility. Programs were developed against documented specifications and were independently "
        "validated, with key results confirmed through double programming to reduce the risk of "
        "programming error.",
        "A reproducible analysis environment captured software versions, dataset versions, and program "
        "execution logs, enabling regeneration of every reported table, figure, and listing from the locked "
        "data. Deviations between planned and executed analyses, where they occurred, were documented with "
        "rationale to maintain transparency and traceability of the reported results.",
    ]),
    ("Appendix F. Risk Management and Pharmacovigilance", [
        "Safety information was collected continuously from the first dose through the end of the "
        "follow-up period. Serious adverse events were reported by sites within the timelines specified in "
        "the protocol and were processed by the pharmacovigilance function for regulatory reporting and "
        "signal evaluation. Causality and expectedness assessments followed reference safety information "
        "and were documented for each event.",
        "An ongoing benefit-risk evaluation integrated emerging safety data with the established efficacy "
        "profile. Risk-minimization measures defined in the protocol, including eligibility restrictions, "
        "monitoring requirements, and stopping rules, were implemented throughout the study. No new "
        "important identified or potential risk requiring modification of the study conduct emerged during "
        "the reporting period.",
    ]),
    ("Appendix G. Handling of Confidential Information and Records Retention", [
        "Subject privacy was protected at every stage of the study. Personally identifiable information was "
        "minimized, access was restricted on a need-to-know basis, and data transfers used secure, "
        "validated mechanisms. The handling of personal data complied with applicable data-protection "
        "regulations, and subjects were informed of how their data would be used and protected.",
        "Essential documents and source records were retained for the period required by applicable "
        "regulation and sponsor policy, and arrangements were made to ensure continued access in the event "
        "of site closure or change of ownership. The sponsor was to be notified before destruction of any "
        "study records to confirm that retention obligations had been satisfied.",
    ]),
    ("Appendix H. Protocol Amendments and Deviations Summary", [
        "Amendments to the protocol were implemented to clarify procedures, incorporate emerging "
        "scientific knowledge, and address operational considerations, and each was reviewed and approved "
        "by the responsible ethics committees and, where required, regulatory authorities before "
        "implementation. The cumulative effect of amendments did not alter the primary objective or the "
        "interpretability of the confirmatory analysis.",
        "Protocol deviations were captured throughout the study, categorized by type and importance, and "
        "reviewed before unblinding to confirm analysis-population assignments. The frequency and nature of "
        "important deviations were consistent with a well-conducted multicenter study and did not undermine "
        "the validity of the efficacy and safety conclusions.",
    ]),
    ("Appendix I. Drug Supply, Packaging, and Accountability", [
        "Investigational and comparator products were manufactured, packaged, and labeled in accordance "
        "with Good Manufacturing Practice and applicable labeling regulations for clinical supplies. "
        "Labeling preserved the blind and included the information required by regulation in the local "
        "language of each region. Supplies were distributed to sites through a validated, temperature-"
        "controlled cold chain with continuous monitoring and documented excursion management.",
        "Site-level drug accountability was reconciled at each monitoring visit and at study close-out. "
        "Records captured receipt, storage conditions, dispensing, return, and destruction or return to "
        "the sponsor. Discrepancies were investigated and resolved, and destruction of unused supplies "
        "occurred only after authorization and against documented procedures with retained certificates.",
    ]),
    ("Appendix J. Monitoring Strategy and Source-Data Verification", [
        "A risk-based monitoring strategy combined centralized review of accumulating data with targeted "
        "on-site visits. Centralized monitoring used statistical and operational indicators to detect "
        "atypical patterns, data-quality issues, and potential non-compliance, directing on-site effort to "
        "the sites and data most critical to subject safety and the reliability of the primary results.",
        "Source-data verification focused on critical data, including eligibility, consent, primary "
        "endpoint, and serious adverse events, at a frequency proportionate to risk. Findings from "
        "monitoring visits were documented in follow-up letters, tracked to resolution, and escalated "
        "through the corrective and preventive action process where systemic issues were identified.",
    ]),
    ("Appendix K. Statistical Output Inventory and Traceability", [
        "A complete inventory of planned tables, figures, and listings was maintained and mapped to the "
        "objectives, endpoints, and analysis populations they support. Each output carried a unique "
        "identifier and a documented program source, supporting full traceability from the reported result "
        "back to the locked analysis dataset and the raw collected data.",
        "Outputs were reviewed for internal consistency, alignment with the analysis plan, and correct "
        "population and visit windowing before inclusion in the report. Any output produced outside the "
        "pre-specified plan was clearly identified as exploratory or supportive to maintain a transparent "
        "distinction between confirmatory and non-confirmatory evidence.",
    ]),
    ("Appendix L. Subject Withdrawal, Discontinuation, and Follow-up", [
        "Subjects could discontinue study treatment or withdraw from the study at any time, and "
        "investigators could discontinue treatment for safety or compliance reasons. The reason for each "
        "discontinuation was recorded, and subjects who discontinued treatment were encouraged to remain "
        "in the study for safety follow-up and endpoint assessment in accordance with the estimand "
        "framework, to minimize the impact of missing data on the primary analysis.",
        "End-of-study procedures included final safety assessments and, where applicable, transition to "
        "available care or an open-label extension. The vital status and primary-endpoint status of "
        "subjects who withdrew were ascertained to the extent permitted by consent, supporting the "
        "completeness and interpretability of the efficacy and safety analyses.",
    ]),
    ("Appendix M. Regulatory Compliance and Inspection Readiness", [
        "The study was conducted in compliance with the protocol, applicable Good Clinical Practice "
        "requirements, and the regulatory obligations of each participating region. Essential documents "
        "were filed contemporaneously, and the trial master file was maintained in an inspection-ready "
        "state with periodic completeness and quality review throughout the conduct of the study.",
        "Training records demonstrated that study personnel were qualified for their delegated tasks and "
        "had received protocol-specific and Good Clinical Practice training appropriate to their roles. "
        "The sponsor maintained processes to support regulatory inspections and audits, including timely "
        "retrieval of records and documented responses to findings, reinforcing the overall reliability "
        "and integrity of the trial.",
    ]),
    ("Appendix N. Endpoint Definitions and Derivation Conventions", [
        "Each efficacy and safety endpoint was defined with sufficient precision to permit unambiguous "
        "programmatic derivation from the collected data. Definitions specified the source variables, the "
        "timing and windowing of assessments, the handling of repeated and unscheduled measurements, and "
        "the rules for deriving change-from-baseline and responder status. These conventions were fixed "
        "before database lock to prevent data-driven choices from influencing the reported results.",
        "Derived variables were documented in dataset specifications and reproduced through validated "
        "programming. Where an endpoint depended on multiple components, the order of operations and the "
        "treatment of partial or missing components were explicitly defined so that every reviewer could "
        "trace a reported value back to its underlying observations without ambiguity.",
    ]),
    ("Appendix O. Glossary of Abbreviations and Terms", [
        "Abbreviations used throughout this document follow standard regulatory and clinical conventions. "
        "Commonly used terms include AE (adverse event), TEAE (treatment-emergent adverse event), SAE "
        "(serious adverse event), ITT (intention-to-treat), PP (per-protocol), SAF (safety population), PK "
        "(pharmacokinetics), CI (confidence interval), LS mean (least-squares mean), MedDRA (Medical "
        "Dictionary for Regulatory Activities), and GCP (Good Clinical Practice).",
        "Additional terms include DSMB (Data Safety Monitoring Board), IRB/IEC (Institutional Review Board / "
        "Independent Ethics Committee), EDC (electronic data capture), CRO (contract research "
        "organization), eCTD (electronic Common Technical Document), and CSR (Clinical Study Report). Where "
        "a term first appears in the body of the document, its meaning is consistent with this glossary to "
        "support a clear and unambiguous reading of the report.",
    ]),
]


def _common_appendix(study: str) -> list:
    """DEPRECATED — link-free administrative appendix (no longer appended).

    Retained only for reference; generate() now uses _xref_section() so that the
    bulk of every document carries resolvable hyperlinks instead of dead filler.
    """
    out: list = ["", "# Appendices — Administrative, Ethical, and Operational Detail"]
    for heading, bodies in _COMMON_APPENDIX_TOPICS:
        out.append("")
        out.append(f"## {heading}")
        for b in bodies:
            out.append(_para(b))
    return out


# ─────────────────────────────────────────────────────────────────────────────
# Cross-reference traceability section — appended to EVERY document in place of
# the old link-free appendix. It supplies page volume AND hyperlink density:
# every paragraph carries resolvable cross-reference tokens, so the WHOLE
# document (not just the first half) is rich with hyperlinks.
# ─────────────────────────────────────────────────────────────────────────────

_XREF_HEADINGS = [
    "Traceability of Efficacy Results",
    "Traceability of Safety Results",
    "Source Data and Listings Map",
    "Analysis Population Cross-References",
    "Statistical Methods Cross-References",
    "Protocol-to-Report Linkage",
    "Pooled and Integrated Analyses Linkage",
    "Demographic and Disposition Linkage",
    "Laboratory, Vital Signs, and ECG Linkage",
    "Pharmacokinetic Cross-References",
    "Deviation and Sensitivity Analysis Linkage",
    "Endpoint Derivation Cross-References",
    "Quality-Control and Programming Traceability",
    "Regulatory Submission Cross-References",
]

# Reusable, resolvable reference fragments combined into flowing sentences.
_XREF_FRAGMENTS = [
    "the responder analysis in Table 14.2.1.1 of CSR {s}",
    "the safety summary in CSR {s} Section 2.7",
    "the analysis populations defined in SAP {s} Section 4.2",
    "the primary analysis model in SAP {s} Section 5.1",
    "the study design in Protocol {s} Section 6.1",
    "the safety-monitoring rules in Protocol {s} Section 9.5",
    "the subject-level data in Listings {s} Section 16.2",
    "the adverse-event listing, Listing 16.2.5",
    "the demographic summary in Table 14.1.1.1 of CSR {s}",
    "the disposition table, Table 14.1.2.1 of CSR {s}",
    "the laboratory shifts in Table 14.3.4.1 of CSR {s}",
    "the pooled safety review in CSR {o0} Section 2.7.4",
    "the integrated efficacy methodology in CSR {o1} Section 5.3.5",
    "the dose-response context in CSR {o2} Section 2.5.3",
    "the deviation inventory in Appendix 16.1",
    "the testing hierarchy in Table 5.1.1 of SAP {s}",
]

_XREF_CONNECTORS = [
    "This subsection consolidates the relevant cross-references. Reviewers should consult {a}, together "
    "with {b} and {c}, to trace each reported result back to its underlying source data.",
    "For completeness, {a} is presented alongside {b}. The derivation and supporting records are further "
    "documented in {c}, allowing independent verification of every value.",
    "The findings summarized here link directly to {a} and {b}, and are reconciled against {c} during the "
    "blinded data review before unblinding and database lock.",
    "Traceability is maintained from {a} through {b}; any apparent discrepancy is resolved by reference to "
    "{c}, with the resolution recorded in the source documentation.",
]


def _xref_fragment(study: str, k: int) -> str:
    o = _others(study)
    return _XREF_FRAGMENTS[k % len(_XREF_FRAGMENTS)].format(s=study, o0=o[0], o1=o[1], o2=o[2])


def _xref_section(study: str) -> list:
    """Long, hyperlink-dense traceability section (replaces the old appendix)."""
    out: list = ["", "# 17. Cross-Reference Index and Traceability"]
    for i, heading in enumerate(_XREF_HEADINGS):
        out.append("")
        out.append(f"## 17.{i + 1} {heading}")
        # Three paragraphs per heading, each weaving in three resolvable refs, so
        # the section is both long enough for page volume and densely hyperlinked.
        for p in range(3):
            base = i * 7 + p * 3
            conn = _XREF_CONNECTORS[(i + p) % len(_XREF_CONNECTORS)]
            out.append(_para(conn.format(
                a=_xref_fragment(study, base),
                b=_xref_fragment(study, base + 1),
                c=_xref_fragment(study, base + 2),
            )))
    return out


# ─────────────────────────────────────────────────────────────────────────────
# Orchestration
# ─────────────────────────────────────────────────────────────────────────────


def generate(out_dir: Path, *, ambiguous: bool = False) -> list[Path]:
    """Write the large 4×4 CSR dossier (10+ pages per document)."""
    written: list[Path] = []
    for study in STUDIES:
        slug = _slug(study)
        folder = out_dir / study
        docs = {
            f"csr-{slug}-body.docx": (f"CSR {study}", _csr_body(study)),
            f"protocol-{slug}.docx": (f"Protocol {study}", _protocol(study)),
            f"sap-{slug}.docx": (f"SAP {study}", _sap(study)),
            f"listings-{slug}.docx": (f"Listings {study}", _listings(study)),
        }
        for filename, (title, paras) in docs.items():
            # Bulk is now the hyperlink-dense traceability section (the old
            # link-free administrative appendix is no longer appended).
            content = list(paras) + _xref_section(study)
            if ambiguous:
                content += _ambiguous_block(study)
            path = folder / filename
            _write(path, title, content)
            written.append(path)
    return written


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate the LARGE 4-study CSR demo dossier (10+ pages/doc).")
    parser.add_argument(
        "--out",
        type=Path,
        default=None,
        help="Output directory. Default: data/synthetic/csr_dossier_large "
        "(or data/synthetic/csr_dossier_large_ollama when --ambiguous is set).",
    )
    parser.add_argument(
        "--ambiguous",
        action="store_true",
        help="Seed each document with sub-threshold references that route through Ollama.",
    )
    args = parser.parse_args()

    out = args.out or Path(
        "data/synthetic/csr_dossier_large_ollama" if args.ambiguous else "data/synthetic/csr_dossier_large"
    )

    written = generate(out, ambiguous=args.ambiguous)
    mode = "ambiguous (Ollama-triggering)" if args.ambiguous else "standard"
    print(f"Generated {len(written)} LARGE {mode} documents across {len(STUDIES)} study folders:")
    for p in written:
        print(f"  {p}")
    print("\nUpload one study folder (4 docs) - or all 16 - via the Pipeline screen.")


if __name__ == "__main__":
    main()
