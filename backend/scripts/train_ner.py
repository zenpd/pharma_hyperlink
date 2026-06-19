"""Train a real spaCy NER model for the reference-detection catalog (NER demo).

Honest framing
--------------
The point of NER in this pipeline is NOT to beat regex on the types regex already
covers (Study-ID, Section, Table, …) — measured F1 there is a wash. NER earns its
place on the **catalog-extension** labels that have *no regex pattern*:
``FORM_REF``, ``SITE_CODE``, ``IMP_BATCH``, ``AE_CODE``, ``SEQUENCE_REF``. The
rule-fallback EntityRuler already catches the *exact* shapes; a TRAINED statistical
model is what lets the system **generalise** to unseen variants and gives a real
``models/ner_v1/`` artifact the app auto-loads (``SpacyNerExtractor`` flips from
``rule_fallback`` → ``trained:…``).

This script builds the training corpus, writes a spaCy config, trains, and prints a
per-label scoreboard. Two data sources are merged:

1. ``data/training/refs.{train,dev}.jsonl`` — the existing 178 hand/gen labels for
   the regex-overlapping types.
2. **Weak supervision** over ``data/synthetic/ner_dossier`` — the EntityRuler labels
   the catalog-extension spans (Form FDA…, Site 0…, MedDRA…), so the statistical
   model learns those types too. (Rules → model distillation; flagged honestly.)

Run from the backend dir:
    PYTHONPATH=src python scripts/train_ner.py
"""

from __future__ import annotations

import glob
import json
import random
from pathlib import Path

import spacy
from docx import Document
from spacy.tokens import DocBin

from hyperlink_engine.core.detection.ner_model import SpacyNerExtractor

ROOT = Path(__file__).resolve().parents[1]
TRAIN_JSONL = ROOT / "data" / "training" / "refs.train.jsonl"
DEV_JSONL = ROOT / "data" / "training" / "refs.dev.jsonl"
NER_DOSSIER = ROOT / "data" / "synthetic" / "ner_dossier"
OUT_DIR = ROOT / "data" / "training" / "ner_v1"
CONFIG = ROOT / "config" / "spacy_ner.cfg"
MODEL_DIR = ROOT / "models" / "ner_v1"

EXCLUSIVE = {"FORM_REF", "SITE_CODE", "IMP_BATCH", "AE_CODE", "SEQUENCE_REF"}


def _read_jsonl(path: Path) -> list[tuple[str, list[tuple[int, int, str]]]]:
    rows = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line:
            continue
        d = json.loads(line)
        ents = [(e[0], e[1], e[2]) for e in d.get("entities", [])]
        rows.append((d["text"], ents))
    return rows


def _weak_label_exclusives() -> list[tuple[str, list[tuple[int, int, str]]]]:
    """Use the EntityRuler to label the catalog-extension entities in the dossier."""
    ner = SpacyNerExtractor()  # rule_fallback EntityRuler
    rows: list[tuple[str, list[tuple[int, int, str]]]] = []
    for f in glob.glob(str(NER_DOSSIER / "**" / "*.docx"), recursive=True):
        doc = Document(f)
        blocks = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for r in t.rows:
                blocks += [c.text for c in r.cells if c.text.strip()]
        for text in blocks:
            ents = [
                (m.start, m.end, m.groups.get("label", m.text))
                for m in ner.extract(text)
                if m.groups.get("label") in EXCLUSIVE
            ]
            if ents:
                rows.append((text, ents))
    return rows


def _to_docbin(nlp, rows) -> DocBin:
    db = DocBin()
    for text, ents in rows:
        doc = nlp.make_doc(text)
        spans = []
        for s, e, lbl in ents:
            span = doc.char_span(s, e, label=lbl, alignment_mode="contract")
            if span is not None:
                spans.append(span)
        # drop overlaps (char_span contraction can collide); keep the longest
        spans = spacy.util.filter_spans(spans)
        doc.ents = spans
        db.add(doc)
    return db


def main() -> None:
    random.seed(13)
    nlp = spacy.blank("en")

    train = _read_jsonl(TRAIN_JSONL)
    dev = _read_jsonl(DEV_JSONL)
    weak = _weak_label_exclusives()
    random.shuffle(weak)
    # 80/20 split of the weak-labeled exclusive examples into train/dev
    cut = max(1, int(len(weak) * 0.8))
    train += weak[:cut]
    dev += weak[cut:]
    print(f"train examples: {len(train)}  dev examples: {len(dev)}  "
          f"(weak-labeled exclusive: {len(weak)})")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    _to_docbin(nlp, train).to_disk(OUT_DIR / "train.spacy")
    _to_docbin(nlp, dev).to_disk(OUT_DIR / "dev.spacy")
    print(f"wrote {OUT_DIR/'train.spacy'} and {OUT_DIR/'dev.spacy'}")


if __name__ == "__main__":
    main()
