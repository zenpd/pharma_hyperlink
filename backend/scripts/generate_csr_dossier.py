"""Generate a 4-study CSR dossier for the cross-document hyperlink demo.

Creates four study folders, each containing four documents that reference
each other within the study and across studies:

    data/synthetic/csr_dossier/
      SP-2026-001/
        csr-sp-2026-001-body.docx       (CSR clinical study report)
        protocol-sp-2026-001.docx
        sap-sp-2026-001.docx            (statistical analysis plan)
        listings-sp-2026-001.docx
      SP-2026-002/ ... (same 4 doc types)
      SP-2026-003/ ...
      SP-2026-004/ ...

The CSR body of each study cross-references:
  * its own protocol / SAP / listings (intra-study), and
  * the other three studies' CSRs (cross-study, cross-folder).

These references are written with the canonical study-id form ("SP-2026-002")
and doc-type words ("Protocol", "SAP", "Listings", "CSR") so the detection +
target-resolution layers route each hyperlink to the correct document.

Run::

    poetry run python -m scripts.generate_csr_dossier
    # or:  python scripts/generate_csr_dossier.py --out data/synthetic/csr_dossier
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

STUDIES = ["SP-2026-001", "SP-2026-002", "SP-2026-003", "SP-2026-004"]

STUDY_TITLES = {
    "SP-2026-001": "A Phase 1 Single-Ascending-Dose PK Study of Solzumab",
    "SP-2026-002": "A Phase 2a Efficacy and Safety Study of Solzumab",
    "SP-2026-003": "A Phase 2b Dose-Finding Study of Solzumab",
    "SP-2026-004": "A Phase 3 Pivotal Efficacy Study of Solzumab",
}


def _slug(study: str) -> str:
    """SP-2026-001 → sp-2026-001 (lowercase, used in filenames)."""
    return study.lower()


def _others(study: str) -> list[str]:
    return [s for s in STUDIES if s != study]


def _table(caption: str, rows: list[list[str]]) -> dict:
    """A table block: a caption paragraph followed by a real Word table.

    The caption (e.g. 'Table 14.2.1.1 - Primary Efficacy Summary') lets the
    detector/snippet locate the table by its number, and the rows give a real
    table that 'Table 14.2.1.1' references can resolve to and preview.
    """
    return {"caption": caption, "rows": rows}


def _write(path: Path, title: str, paragraphs: list) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for para in paragraphs:
        # Table block — {"caption": str, "rows": [[...], ...]}
        if isinstance(para, dict) and "rows" in para:
            if para.get("caption"):
                doc.add_paragraph(para["caption"])
            rows = para["rows"]
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                try:
                    tbl.style = "Light Grid Accent 1"
                except Exception:  # noqa: BLE001 — style is cosmetic only
                    pass
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        tbl.cell(ri, ci).text = str(val)
            continue
        if para.startswith("# "):
            doc.add_heading(para[2:], level=1)
        elif para.startswith("## "):
            doc.add_heading(para[3:], level=2)
        elif para == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(para)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


# ─────────────────────────────────────────────────────────────────────────────
# Per-document content builders
# ─────────────────────────────────────────────────────────────────────────────


def _nct(study: str) -> str:
    """Deterministic ClinicalTrials.gov ID per study (NCT + 8 digits)."""
    n = study.split("-")[-1]  # "001"
    return f"NCT0482799{int(n)}"  # NCT04827991 .. NCT04827994


def _csr_body(study: str) -> list[str]:
    o = _others(study)
    return [
        f"Clinical Study Report — {STUDY_TITLES[study]} ({study}).",
        "",
        "# 1. Introduction and Study Overview",
        f"This report presents the results of study {study}. The full study "
        f"design is described in Protocol {study} Section 6.1. The pre-specified "
        f"analyses are defined in the SAP {study}.",
        f"This trial is registered as {_nct(study)} on ClinicalTrials.gov. It was "
        f"conducted in accordance with ICH E6(R2) GCP guidelines and reported per "
        f"ICH E3 guidance. Sponsor reporting follows FDA guidance at "
        f"https://www.fda.gov/regulatory-information/search-fda-guidance-documents.",
        "",
        "# 2.3 Demographic and Baseline Characteristics",
        f"A total of {[447, 450, 455, 460][int(study[-3:]) - 1]} subjects were enrolled across multiple sites. "
        f"Refer to Listings {study} for individual subject data. Key baseline characteristics are presented in Table 14.1.1.1.",
        _table(
            "Table 14.1.1.1 - Subject Demographics (Intent-to-Treat Population)",
            [
                ["Characteristic", "Solzumab 200 mg", "Solzumab 100 mg", "Placebo", "Total"],
                ["N enrolled", "150", "149", "148", "447"],
                ["Age (years), mean (SD)", "52.3 (12.1)", "53.1 (11.8)", "52.7 (12.4)", "52.7 (12.1)"],
                ["Male, n (%)", "89 (59.3)", "88 (59.1)", "88 (59.5)", "265 (59.3)"],
                ["Race: White, n (%)", "125 (83.3)", "124 (83.2)", "123 (83.1)", "372 (83.2)"],
            ],
        ),
        "",
        "# 2.5 Clinical Overview and Efficacy",
        f"Efficacy results are summarized in Section 2.5.3 and Table 14.2.1.1. "
        f"Subject-level source data are provided in Listings {study}. The primary endpoint "
        f"analysis compares all treatment arms against placebo using logistic regression.",
        _table(
            "Table 14.2.1.1 - Primary Efficacy Summary (ITT Population)",
            [
                ["Treatment Arm", "N", "Responders (%)", "95% CI", "p-value"],
                ["Solzumab 200 mg", "150", "92 (61.3%)", "(53.0–69.6%)", "0.0003"],
                ["Solzumab 100 mg", "149", "77 (51.7%)", "(43.3–60.0%)", "0.0156"],
                ["Placebo", "148", "56 (37.8%)", "(29.8–46.2%)", "—"],
            ],
        ),
        f"For the pooled safety population, cross-reference CSR {o[0]} Section 2.7 "
        f"and CSR {o[1]} Section 2.7.4.",
        "",
        "# 2.7 Clinical Summary and Safety Analysis",
        f"Integrated efficacy follows the methodology of CSR {o[2]} Section 5.3.5. "
        f"Adverse events are tabulated in Table 14.3.1.2 and detailed in Listings {study}. "
        f"The safety population consists of all enrolled subjects.",
        _table(
            "Table 14.3.1.2 - Treatment-Emergent Adverse Events by System Organ Class",
            [
                ["System Organ Class", "Solzumab (n=299)", "Placebo (n=148)", "Total (n=447)"],
                ["Any TEAE", "223 (74.6%)", "98 (66.2%)", "321 (71.8%)"],
                ["Gastrointestinal disorders", "72 (24.1%)", "31 (20.9%)", "103 (23.1%)"],
                ["Infections and infestations", "51 (17.1%)", "22 (14.9%)", "73 (16.3%)"],
                ["Nervous system disorders", "38 (12.7%)", "14 (9.5%)", "52 (11.6%)"],
                ["Serious AEs", "12 (4.0%)", "7 (4.7%)", "19 (4.3%)"],
            ],
        ),
        _table(
            "Table 14.3.1.3 - Treatment-Related Adverse Events (Safety Population)",
            [
                ["Event", "Solzumab (n=299)", "Placebo (n=148)"],
                ["Related TEAE", "156 (52.2%)", "31 (20.9%)"],
                ["Mild–Moderate", "142 (47.5%)", "28 (18.9%)"],
                ["Severe", "14 (4.7%)", "3 (2.0%)"],
                ["Leading to discontinuation", "8 (2.7%)", "2 (1.4%)"],
            ],
        ),
        f"Deaths and serious adverse events are reconciled against Protocol {study} "
        f"Section 9.5 and the SAP {study} Section 4.2.",
        "",
        "# 5.3 Clinical Study Reports",
        f"This CSR {study} should be read together with CSR {o[0]}, CSR {o[1]}, "
        f"and CSR {o[2]} for the complete clinical picture of the Solzumab program. "
        f"Cross-references within this report are denoted as 'Section X.Y.Z' or "
        f"'Table N.N.N.N' and redirect to the specific location.",
    ]


def _protocol(study: str) -> list[str]:
    return [
        f"Clinical Trial Protocol — {study}.",
        "",
        "# 1. Protocol Summary",
        f"Protocol {study} is a randomized, double-blind, placebo-controlled trial "
        f"of Solzumab in patients with the target condition. The study design, "
        f"safety monitoring, and statistical approach are detailed in sections below.",
        "",
        "# 6.1 Study Design",
        f"This protocol governs study {study}. Endpoints align with the analysis "
        f"defined in SAP {study}. Results are reported in CSR {study} Section 2.5.",
        _table(
            "Table 6.1.1 - Study Design Parameters",
            [
                ["Parameter", "Value"],
                ["Type", "Randomized, double-blind, placebo-controlled"],
                ["Duration", "12 weeks treatment + 4 weeks follow-up"],
                ["Primary Endpoint", "Proportion of responders at Week 12"],
                ["Secondary Endpoints", "Time to response, quality of life, safety"],
                ["Target N", "~450 (power analysis per SAP {study})"],
            ],
        ),
        "",
        "# 7. Pharmacokinetics",
        f"PK sampling and analysis are specified in Section 7.2 of this protocol "
        f"and cross-referenced in SAP {study} Section 5.1. All subjects provide "
        f"PK samples at protocol-defined timepoints.",
        "",
        "# 9.5 Safety Assessments and Monitoring",
        f"Safety data feed Table 14.3.1.2 of CSR {study} and the subject "
        f"Listings {study}. An independent Data Safety Monitoring Board meets "
        f"every 8 weeks to review blinded safety data and interim efficacy signals.",
        _table(
            "Table 9.5.1 - Safety Assessment Schedule",
            [
                ["Assessment", "Screening", "Baseline", "Weekly", "End of Study"],
                ["Physical Exam", "X", "X", "", "X"],
                ["Clinical Lab", "X", "X", "X", "X"],
                ["AE Monitoring", "", "X", "X", "X"],
                ["ECG", "X", "", "", "X"],
            ],
        ),
    ]


def _sap(study: str) -> list[str]:
    return [
        f"Statistical Analysis Plan — {study}.",
        "",
        "# 1. Overview and Objectives",
        f"This Statistical Analysis Plan (SAP {study}) defines all statistical "
        f"analyses for protocol {study}. It was finalized before database lock "
        f"and approved by the independent statistician. Key analysis populations, "
        f"methods, and outputs are defined below.",
        "",
        "# 4.2 Analysis Populations",
        f"Analyses defined here are reported in CSR {study} Section 2.7. The "
        f"design assumptions trace to Protocol {study} Section 6.1.",
        _table(
            "Table 4.2.1 - Definition of Analysis Populations",
            [
                ["Population", "Description", "N (est.)"],
                ["Intention-to-Treat (ITT)", "All randomized subjects", "~450"],
                ["Per-Protocol (PP)", "Completers with ≥80% compliance", "~400"],
                ["Safety (SAF)", "All subjects who received ≥1 dose", "~450"],
                ["PK-evaluable", "All with valid PK samples", "~440"],
            ],
        ),
        "",
        "# 5.1 Statistical Analysis Methods",
        f"Outputs map to Table 14.2.1.1 and Table 14.3.1.2 of CSR {study}, "
        f"with supporting subject-level data in Listings {study}.",
        _table(
            "Table 5.1.1 - Primary and Secondary Endpoint Analyses",
            [
                ["Endpoint", "Population", "Method", "Significance Level"],
                ["Primary (responder rate)", "ITT", "Logistic regression vs placebo", "α=0.05"],
                ["Secondary (time to response)", "ITT", "Kaplan-Meier, log-rank test", "Exploratory"],
                ["Safety (TEAE incidence)", "SAF", "Descriptive (% and 95% CI)", "Exploratory"],
                ["PK (AUC, Cmax)", "PK-evaluable", "ANOVA, dose-linearity", "Exploratory"],
            ],
        ),
        "",
        "# 6. Sensitivity and Subgroup Analyses",
        f"Sensitivity analyses per protocol amendment and baseline risk stratification "
        f"subgroups are detailed in Section 6.2. Cross-references to Protocol {study} "
        f"Section 10.2 and CSR {study} Appendix 16.1.",
    ]


def _listings(study: str) -> list[str]:
    return [
        f"Subject Data Listings — {study}.",
        "",
        "# 16.1 Baseline and Demographics Listings",
        f"Individual subject demographics and baseline characteristics supporting "
        f"Table 14.1.1.1 in CSR {study}. Data are sorted by randomization order "
        f"and treatment group within CSR {study} Section 2.3.",
        "",
        "# 16.2 Subject Efficacy Listings",
        f"These listings support CSR {study} Section 2.5 and the efficacy tables "
        f"(Table 14.2.1.1) defined in SAP {study} Section 5.1. Individual subject "
        f"efficacy assessments by visit and responder status are included.",
        _table(
            "Listing 16.2.1 - Subject Efficacy Assessment Summary (excerpt)",
            [
                ["Subject ID", "Treatment", "Baseline Score", "Week 4 Score", "Week 12 Score", "Responder"],
                ["S001-001", "Solzumab 200 mg", "42", "28", "15", "Yes"],
                ["S001-002", "Placebo", "41", "38", "36", "No"],
                ["S001-003", "Solzumab 100 mg", "43", "32", "20", "Yes"],
            ],
        ),
        "",
        "# 16.3 Subject Safety Listings",
        f"Adverse events by subject, with MedDRA coding to version 27.0 as specified "
        f"in SAP {study} Section 5.1. Complete SAE narratives and follow-up outcomes "
        f"are included. Table 14.3.1.2 in CSR {study} summarizes incidence by SOC.",
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Ambiguous-reference block (opt-in via --ambiguous)
# ─────────────────────────────────────────────────────────────────────────────
#
# These sentences deliberately use *bare dotted numbers* next to a context cue
# ("see", "described in", "per") but WITHOUT a "Section"/"Table" prefix. The
# detector matches them with the low-confidence SECTION_REF_DOTTED_V1 pattern
# (confidence 0.55), which sits below the 0.7 LLM threshold — so each one is
# escalated to Ollama for disambiguation. Roughly 3 such references per doc, so
# 16 docs produce about 48 Ollama calls. The default dossier never includes
# this block, so its output stays byte-identical.


def _ambiguous_block(study: str) -> list[str]:
    o = _others(study)
    return [
        "",
        "# 9.9 Supplementary Cross-References",
        f"For the dosing rationale, see 2.5.3 and the methods described in 5.3.5 "
        f"of the broader Solzumab program ({study}).",
        f"Integrated safety for the pooled population is summarized per 14.2.1 and "
        f"should be read together with the {o[0]} findings noted earlier.",
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Orchestration
# ─────────────────────────────────────────────────────────────────────────────


def generate(out_dir: Path, *, ambiguous: bool = False) -> list[Path]:
    """Write the 4×4 CSR dossier.

    When ``ambiguous`` is True, each document gets an extra block of bare
    dotted-number references that fall below the LLM confidence threshold and
    therefore route through Ollama. When False (the default) the output is
    unchanged from the original demo dossier.
    """
    written: list[Path] = []
    for study in STUDIES:
        slug = _slug(study)
        folder = out_dir / study
        docs = {
            f"csr-{slug}-body.docx": (f"CSR {study}", _csr_body(study)),
            f"protocol-{slug}.docx": (f"Protocol {study}", _protocol(study)),
            f"sap-{slug}.docx": (f"SAP {study}", _sap(study)),
            f"listings-{slug}.docx": (f"Listings {study}", _listings(study)),
        }
        for filename, (title, paras) in docs.items():
            content = paras + _ambiguous_block(study) if ambiguous else paras
            path = folder / filename
            _write(path, title, content)
            written.append(path)
    return written


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate the 4-study CSR demo dossier.")
    parser.add_argument(
        "--out",
        type=Path,
        default=None,
        help="Output directory. Default: data/synthetic/csr_dossier, or "
        "data/synthetic/csr_ollama_dossier when --ambiguous is set.",
    )
    parser.add_argument(
        "--ambiguous",
        action="store_true",
        help="Seed each document with ~3 sub-threshold references that route "
        "through Ollama (default off — keeps the standard dossier unchanged).",
    )
    args = parser.parse_args()

    out = args.out or Path(
        "data/synthetic/csr_ollama_dossier" if args.ambiguous else "data/synthetic/csr_dossier"
    )

    written = generate(out, ambiguous=args.ambiguous)
    mode = "ambiguous (Ollama-triggering)" if args.ambiguous else "standard"
    print(f"Generated {len(written)} {mode} documents across {len(STUDIES)} study folders:")
    for p in written:
        print(f"  {p}")
    if args.ambiguous:
        print(
            "\nEach doc carries ~3 bare-number references below the 0.7 confidence "
            "floor -> ~48 Ollama calls total. Check output/.../llm_calls.jsonl after a run."
        )
    print("\nUpload one study folder (4 docs) - or all 16 - via the Pipeline screen.")


if __name__ == "__main__":
    main()
