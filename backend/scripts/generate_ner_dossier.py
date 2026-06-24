"""Generate the NER-exercise dossier for demonstrating the NER detection layer.

The existing csr_ollama_dossier documents only trigger the regex layer because
every reference (SP-2026-XXX, Section X.Y.Z, Table N.N.N) is caught by the
regex catalog with confidence 0.92–0.99, which beats NER's 0.80.

This dossier deliberately embeds entity types that are EXCLUSIVELY detected by
the NER EntityRuler layer (no matching regex pattern exists):

  * FORM_REF   — FDA regulatory form numbers (Form FDA 1572 / Form FDA-3674A)
  * SITE_CODE  — Clinical site/centre identifiers (Site 001, Site US-042)
  * IMP_BATCH  — Investigational product batch/lot numbers (Batch No. A2024-001)
  * AE_CODE    — MedDRA-coded adverse-event terms (MedDRA PT: Hepatotoxicity)
  * SEQUENCE_REF — eCTD submission sequence refs (Sequence 0001)

These appear alongside the standard Study-ID/Section/Table refs so the
Detection Trace screen shows all three columns populated:
    regex_only   NER_triggered   LLM_triggered(optional)

Cross-reference map (9 NER-exclusive refs + 9 standard refs per dossier):

  Protocol AMD-2025-001 (Phase 2 Dose-Finding Protocol)
      standard: SP-2026-101 Section 3.2, Table 5.1, Listing 16.1
      NER-only: Site US-001, Form FDA 1572, MedDRA PT: Dizziness,
                Sequence 0001

  CSR AMD-2025-002 (Phase 2 CSR referencing the protocol + CRF)
      standard: AMD-2025-001 Section 4.2, Figure 3, Appendix 16.1
      NER-only: Batch No. B2025-044, Site 007, MedDRA SOC: Nervous system disorders,
                Form FDA-3674A

  SAP AMD-2025-003 (Statistical Analysis Plan)
      standard: AMD-2025-001 Section 5.1, Table 7.2, Listing 16.2
      NER-only: Site US-014, Lot C2025-099, MedDRA PT: Nausea,
                Sequence 0002

Usage::

    cd hyperlink-engine/backend
    python scripts/generate_ner_dossier.py
    python scripts/generate_ner_dossier.py --out data/synthetic/ner_dossier

Output::

    data/synthetic/ner_dossier/
        m5/53-clin-stud-rep/
            protocol-amd-2025-001.docx
            csr-amd-2025-002.docx
            sap-amd-2025-003.docx
        index.xml
        MANIFEST.txt
"""

from __future__ import annotations

import argparse
import textwrap
from datetime import date
from pathlib import Path

# ── docx writer ──────────────────────────────────────────────────────────────

def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


def _make_doc(title: str, content: list) -> object:
    """Create a python-docx Document with a title and body sections/tables.

    content can contain:
    - tuples (heading, body_text) for text sections
    - dicts with "caption", "rows" for tables
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date + sponsor
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Amida Pharma Ltd  ·  {date.today().isoformat()}").italic = True
    doc.add_paragraph()

    for item in content:
        # Table block
        if isinstance(item, dict) and "rows" in item:
            if item.get("caption"):
                doc.add_paragraph(item["caption"])
            rows = item["rows"]
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                try:
                    tbl.style = "Light Grid Accent 1"
                except Exception:  # style is cosmetic only
                    pass
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        tbl.cell(ri, ci).text = str(val)
            doc.add_paragraph()
            continue

        # Text section (heading, body)
        heading, body = item
        doc.add_heading(heading, level=2)
        for line in body.strip().split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        doc.add_paragraph()

    return doc


# ── document content ─────────────────────────────────────────────────────────

def _table(caption: str, rows: list[list[str]]) -> dict:
    """A table block: caption paragraph + real Word table."""
    return {"caption": caption, "rows": rows}


PROTOCOL_CONTENT = [
    ("1. Introduction and Background",
     textwrap.dedent("""\
         This protocol AMD-2025-001 describes a Phase 2 dose-finding study of
         Amidazetil in adult patients with moderate hepatic impairment (MedDRA PT:
         Hepatic impairment) conducted under IND NCT00342178.
         The study design follows the principles outlined in ICH E8(R1) and
         Module 2.5 of the CTD dossier, Sequence 0001.
         Investigators must complete Form FDA 1572 (Statement of Investigator)
         before enrolling any subject at their site.""")),

    ("2. Study Objectives",
     textwrap.dedent("""\
         Primary: Evaluate PK parameters of Amidazetil 50 mg and 100 mg in Section 3.2.
         Secondary: Assess safety endpoints per Table 5.1 of this protocol.
         Exploratory: Characterise MedDRA PT: Dizziness events across doses.
         All listings of individual subject data are compiled in Listing 16.1.""")),

    ("3. Investigational Plan and Methodology",
     textwrap.dedent("""\
         3.1 Study Design
         This is a randomised, double-blind, placebo-controlled study.
         Subjects are enrolled across six clinical sites.  Site US-001 will
         serve as the coordinating centre.

         3.2 Investigational Medicinal Product
         Amidazetil film-coated tablets (50 mg and 100 mg) are supplied by Amida
         Pharma Ltd.  Each shipment is accompanied by a Certificate of Analysis
         referencing the assigned batch.  The first dispensed batch is
         Batch No. A2025-018.""")),

    _table(
        "Table 5.1 - Planned Dose Regimens and Cohorts",
        [
            ["Cohort", "Dose (mg)", "Formulation", "N Subjects", "Duration (days)"],
            ["1", "50", "Film-coated tablet", "18", "14"],
            ["2", "100", "Film-coated tablet", "18", "14"],
            ["3", "Placebo", "Matching tablet", "12", "14"],
        ],
    ),

    ("3.3 Informed Consent and Regulatory",
     textwrap.dedent("""\
         Signed ICF is collected by site personnel per GCP guidelines before any
         study procedure.  Investigators must hold a valid Form FDA-3674A on file.
         Site US-001 (coordinating) and Site 007 (enrollment lead) are responsible
         for regulatory compliance at enrolled sites.""")),

    ("4. Subject Eligibility",
     textwrap.dedent("""\
         Inclusion: Adults 18–75 years, Child-Pugh A/B hepatic impairment
         (MedDRA SOC: Hepatobiliary disorders), ECOG PS 0–1.
         Exclusion: Active malignancy, renal impairment (eGFR <30 mL/min/1.73 m²),
         use of strong CYP3A4 inhibitors within 14 days.

         The statistical analysis of eligibility screen failures is described
         in SAP AMD-2025-003, Section 5.1 and Table 7.2.""")),

    ("5. Cross-References",
     textwrap.dedent("""\
         This protocol should be read in conjunction with:
           • CSR AMD-2025-002 for clinical results including Figure 3 and Appendix 16.1.
           • SAP AMD-2025-003 for the full statistical methodology (Sequence 0002).
           • MedDRA PT: Nausea is the primary tolerability endpoint; coding details
             are in CSR AMD-2025-002.
           • Site 007 enrolled the largest cohort; site-level data are tabulated in
             CSR AMD-2025-002 Table 6.3.
           • Batch No. B2025-044 and Lot C2025-099 CoAs are on file per Form FDA 1572.""")),
]

CSR_CONTENT = [
    ("1. Study Synopsis",
     textwrap.dedent("""\
         Study AMD-2025-002 evaluated Amidazetil 50 mg and 100 mg vs placebo in 96
         patients with hepatic impairment.  The study was conducted under protocol
         AMD-2025-001, Section 4.2, and registered as NCT00342178.
         All investigators completed Form FDA 1572 prior to enrolment.
         Site US-014 served as primary co-ordinator and Site US-001 as regulatory lead.""")),

    ("2. Investigational Product and Batches",
     textwrap.dedent("""\
         2.1 Product Details
         Amidazetil (IMP) was supplied as film-coated tablets (50 mg and 100 mg).
         Study drug was dispensed from two manufacturing lots.  Lot C2025-099 was
         used for Cohort 1 (50 mg) and Batch No. B2025-044 for Cohort 2 (100 mg).
         Certificates of analysis are filed in Appendix 16.1.

         2.2 Randomisation and Blinding
         Randomisation codes were generated per the procedure in SAP AMD-2025-003,
         Section 5.1.  Site 007 received blinded kits numbered 7001–7024.
         Sequence 0001 approval was received on 2025-01-15.""")),

    _table(
        "Table 6.3 - Demographics and Baseline Characteristics (Safety Population)",
        [
            ["Characteristic", "Amidazetil 50 mg (N=32)", "Amidazetil 100 mg (N=32)", "Placebo (N=12)"],
            ["Age (years), mean (SD)", "54.2 (11.8)", "55.1 (10.3)", "53.8 (12.1)"],
            ["Male, n (%)", "18 (56%)", "19 (59%)", "7 (58%)"],
            ["Weight (kg), mean (SD)", "78.5 (14.2)", "79.3 (13.9)", "77.1 (15.3)"],
            ["Child-Pugh A/B, n", "28/4", "27/5", "10/2"],
        ],
    ),

    ("3. Safety Results",
     textwrap.dedent("""\
         3.1 Adverse Events
         Treatment-emergent adverse events were coded using MedDRA version 27.0.
         The most common events were MedDRA PT: Nausea (18%), MedDRA PT: Dizziness
         (12%), and MedDRA PT: Headache (9%).  No events of MedDRA PT: Hepatotoxicity
         were reported.
         System organ class distribution is shown in Figure 3.
         Complete AE listings are in Appendix 16.1.

         3.2 Serious Adverse Events
         One SAE of MedDRA SOC: Nervous system disorders was reported at Site US-014.
         Details are in Appendix 16.2 of this report and cross-referenced to the
         site file for Site US-014 held at the sponsor.
         The case was reviewed by the Data Safety Monitoring Board per Form FDA-3674A guidance.""")),

    _table(
        "Table 7.2 - Summary of Treatment-Emergent Adverse Events",
        [
            ["System Organ Class", "Amidazetil (N=64)", "Placebo (N=12)", "Total (N=76)"],
            ["Any TEAE", "48 (75%)", "8 (67%)", "56 (74%)"],
            ["Gastrointestinal disorders", "22 (34%)", "3 (25%)", "25 (33%)"],
            ["Nervous system disorders", "18 (28%)", "2 (17%)", "20 (26%)"],
            ["Serious AEs", "2 (3%)", "0 (0%)", "2 (3%)"],
        ],
    ),

    ("4. Efficacy Results",
     textwrap.dedent("""\
         4.1 Primary Endpoint
         The primary endpoint (PK AUC 0–24 h) was analysed as specified in
         SAP AMD-2025-003, Table 7.2.  Results are summarised in Table 6.3 and
         individual data are in Listing 16.1 of this report.

         4.2 Secondary Endpoints
         Secondary PK parameters are described in protocol AMD-2025-001,
         Section 3.2, and are tabulated in Table 7.2 of this CSR.""")),

    ("5. Cross-References",
     textwrap.dedent("""\
         This CSR references the following documents:
           • Protocol AMD-2025-001 for the full study design (especially Section 4.2).
           • SAP AMD-2025-003 for the statistical analysis framework.
           • Batch No. B2025-044 and Lot C2025-099 CoAs are filed in Appendix 16.1.
           • Site 007 and Site US-014 site files held at the sponsor.
           • All AEs coded to MedDRA SOC: Hepatobiliary disorders are listed in Appendix 16.2.
           • Form FDA 1572 copies are on file per ICH E6(R2) requirement.
           • Sequence 0001 (initial) and Sequence 0002 (amendment) submitted to FDA.""")),
]

SAP_CONTENT = [
    ("1. Introduction",
     textwrap.dedent("""\
         This Statistical Analysis Plan (SAP AMD-2025-003) governs all analyses
         for protocol AMD-2025-001, Sequence 0002.  It was finalised before
         database lock and approved by the independent statistician.
         Registration: NCT00342178.  This plan aligns with ICH E3 and ICH E9 guidance.""")),

    ("2. Study Design and Endpoints",
     textwrap.dedent("""\
         2.1 Primary Endpoint
         The primary PK endpoint is AUC 0–24 h as specified in Section 5.1.
         Analysis populations are defined per protocol AMD-2025-001, Section 3.2.

         2.2 Safety Endpoints
         All adverse events are coded to MedDRA PT using MedDRA version 27.0.
         MedDRA PT: Nausea and MedDRA PT: Dizziness are prespecified as
         tolerability endpoints of special interest.

         2.3 Site Data
         Site-level summaries will be produced for all enrolled sites.  Site US-001
         and Site 007 are expected to contribute the largest populations.""")),

    ("3. Statistical Methods",
     textwrap.dedent("""\
         3.1 Primary Analysis
         The primary analysis uses a mixed-effects model as described in Table 7.2.
         Sensitivity analyses are listed in Listing 16.2.

         3.2 Safety Analysis
         Incidence rates for MedDRA SOC: Hepatobiliary disorders and
         MedDRA PT: Dizziness are presented in Figure 3.
         Exposure data from Batch No. A2025-018 and Lot C2025-099 are used to
         derive dose-normalised PK metrics.""")),

    _table(
        "Table 7.2 - Statistical Analysis Methods and Model Parameters",
        [
            ["Analysis", "Method", "Population", "Primary Endpoint"],
            ["Primary PK", "Mixed-effects model", "PK-evaluable", "AUC 0-24h"],
            ["Secondary PK", "ANOVA", "PK-evaluable", "Cmax, Tmax"],
            ["Safety", "Descriptive statistics", "Safety", "Incidence of AEs"],
            ["Efficacy", "Logistic regression", "ITT", "Response rate"],
        ],
    ),

    ("3.3 Subject-Level Listings",
     textwrap.dedent("""\
         All individual AE data are compiled in Listing 16.2 per ICH E3 guidance.
         Site-level listings for Site US-001 and Site US-014 are in Appendix 16.1.
         MedDRA coding is version 27.0 for all adverse events.""")),

    ("4. Data Management",
     textwrap.dedent("""\
         4.1 Data Capture
         Data are captured via electronic CRF.  Site 007 uses a legacy paper CRF
         per protocol amendment.  Form FDA-3674A waivers are on file for two
         investigators at Site US-014.

         4.2 eCTD Submission
         This document will be submitted as part of Sequence 0001 and Sequence 0002
         of the eCTD dossier.  Module 5.3.3 will contain the CSR and appendices.""")),

    ("5. Cross-References",
     textwrap.dedent("""\
         This SAP cross-references:
           • Protocol AMD-2025-001 Section 3.2 and Section 5.1 for endpoint definitions.
           • CSR AMD-2025-002 Table 6.3 and Figure 3 for observed results.
           • Listing 16.1 and Listing 16.2 in CSR AMD-2025-002 for subject-level data.
           • Sequence 0001 (initial submission) and Sequence 0002 (amendment).
           • Batch No. A2025-018 CoA (Appendix 16.1 of the protocol).
           • Lot C2025-099 CoA (Appendix 16.1 of the CSR).
           • All MedDRA PT: Headache events are reconciled in Appendix 16.2.""")),
]


# ── generator ────────────────────────────────────────────────────────────────

DOCUMENTS: list[tuple[str, str, list]] = [
    (
        "protocol-amd-2025-001.docx",
        "Protocol AMD-2025-001: Phase 2 Dose-Finding Study of Amidazetil",
        PROTOCOL_CONTENT,
    ),
    (
        "csr-amd-2025-002.docx",
        "Clinical Study Report AMD-2025-002: Amidazetil Phase 2",
        CSR_CONTENT,
    ),
    (
        "sap-amd-2025-003.docx",
        "Statistical Analysis Plan AMD-2025-003: Amidazetil Phase 2",
        SAP_CONTENT,
    ),
]


def _ectd_xml(docs: list[str]) -> str:
    leaves = "\n".join(
        f'        <leaf ID="ID-{i+1:03d}" xlink:href="m5/53-clin-stud-rep/{d}">'
        f"\n            <title>{d}</title>\n        </leaf>"
        for i, d in enumerate(docs)
    )
    return textwrap.dedent(f"""\
        <?xml version="1.0" encoding="UTF-8"?>
        <ectd xmlns:xlink="http://www.w3.org/1999/xlink">
          <sequence>0002</sequence>
          <tableOfContents>
            <section ID="m5" title="Module 5 - Clinical">
              <section ID="m53" title="5.3 Reports of Clinical Studies">
                <section ID="m533" title="5.3.3 Reports of Controlled Clinical Studies">
    {leaves}
                </section>
              </section>
            </section>
          </tableOfContents>
        </ectd>
        """)


def generate(out_dir: Path) -> None:
    leaf_dir = out_dir / "m5" / "53-clin-stud-rep"
    leaf_dir.mkdir(parents=True, exist_ok=True)

    if not _docx_available():
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        return

    generated: list[str] = []
    for filename, title, sections in DOCUMENTS:
        path = leaf_dir / filename
        doc = _make_doc(title, sections)
        doc.save(str(path))
        print(f"  wrote {path}")
        generated.append(filename)

    # eCTD stub
    xml_path = out_dir / "index.xml"
    xml_path.write_text(_ectd_xml(generated), encoding="utf-8")
    print(f"  wrote {xml_path}")

    # Manifest
    manifest = out_dir / "MANIFEST.txt"
    manifest.write_text(
        f"NER Exercise Dossier — Amida Pharma / Amidazetil Phase 2\n"
        f"Generated: {date.today().isoformat()}\n\n"
        "Documents:\n" + "\n".join(f"  {f}" for f in generated) + "\n\n"
        "NER-exclusive entity types embedded:\n"
        "  FORM_REF    — Form FDA 1572, Form FDA-3674A\n"
        "  SITE_CODE   — Site US-001, Site 007, Site US-014\n"
        "  IMP_BATCH   — Batch No. A2025-018, Batch No. B2025-044, Lot C2025-099\n"
        "  AE_CODE     — MedDRA PT: Nausea, MedDRA PT: Dizziness, MedDRA SOC: ...\n"
        "  SEQUENCE_REF— Sequence 0001, Sequence 0002\n\n"
        "Standard entity types (caught by regex, will show as regex in trace):\n"
        "  STUDY_ID    — AMD-2025-001, AMD-2025-002, AMD-2025-003, NCT00342178\n"
        "  SECTION_REF — Section 3.2, Section 4.2, Section 5.1\n"
        "  TABLE_REF   — Table 5.1, Table 6.3, Table 7.2\n"
        "  FIGURE_REF  — Figure 3\n"
        "  LISTING_REF — Listing 16.1, Listing 16.2\n"
        "  APPENDIX_REF— Appendix 16.1, Appendix 16.2\n",
        encoding="utf-8",
    )
    print(f"  wrote {manifest}")
    print(f"\nDone — {len(generated)} documents in {out_dir}")


def main() -> None:
    p = argparse.ArgumentParser(description="Generate NER-exercise dossier")
    p.add_argument(
        "--out",
        default="data/synthetic/ner_dossier",
        help="Output directory (default: data/synthetic/ner_dossier)",
    )
    args = p.parse_args()
    out = Path(args.out)
    print(f"Generating NER dossier -> {out.resolve()}")
    generate(out)


if __name__ == "__main__":
    main()
