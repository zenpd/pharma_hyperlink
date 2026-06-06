"""Generate the 4-document demo dossier for Plan Two (Track A).

Produces realistic CTD Module 5 Clinical Study Reports (CSRs) for a
fictitious drug Hyperlinkimab (SP-2026-xxx series).  Each document is
structured like a real regulatory CSR:

    Title Page / Synopsis
    1. Introduction
    2. Study Objectives
    3. Investigational Plan / Methods
    4. Study Subjects
    5. Efficacy Evaluation
    6. Safety Evaluation
    7. Discussion and Conclusions
    8. Cross-References                  <-- all cross-doc refs live here
    Appendices / Listings / Figures

Cross-reference map (12 links total — 3 per document):

    CSR SP-2026-001 (Phase 1 PK)
        -> SP-2026-002 Section 2.5        safety comparison
        -> SP-2026-003 Appendix 16.1.1    PK listings reused
        -> SP-2026-004 Table 14.2.1.1     integrated demographics

    CSR SP-2026-002 (Phase 2a Efficacy)
        -> SP-2026-001 Section 5.3.1      PK profile
        -> SP-2026-003 Listing 16.2.5     AE reconciliation
        -> SP-2026-004 Section 5.3.5      pivotal efficacy

    CSR SP-2026-003 (Phase 2b Dose-Finding)
        -> SP-2026-001 Figure 11          PK exposure-response
        -> SP-2026-002 Table 14.2.1.1     Phase 2a primary endpoint
        -> SP-2026-004 Appendix 16.2.6    combined safety dataset

    CSR SP-2026-004 (Phase 3 Pivotal)
        -> SP-2026-001 Section 5.3.1      baseline PK
        -> SP-2026-002 Section 2.5        Phase 2a safety reference
        -> SP-2026-003 Table 14.3.1       dose-response summary

Usage::

    cd hyperlink-engine
    python scripts/generate_demo_dossier.py
    python scripts/generate_demo_dossier.py --out data/synthetic/demo_dossier

Output::

    data/synthetic/demo_dossier/
        m5/53-clin-stud-rep/
            csr-sp-2026-001.docx
            csr-sp-2026-002.docx
            csr-sp-2026-003.docx
            csr-sp-2026-004.docx
        index.xml
        MANIFEST.txt
"""

from __future__ import annotations

import argparse
import sys
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from lxml import etree


# ─────────────────────────────────────────────────────────────────────────────
# Dossier constants
# ─────────────────────────────────────────────────────────────────────────────

DOSSIER_ID = "DOS-2026-DEMO"
SPONSOR = "SunPharma Ltd."
DRUG = "Hyperlinkimab"
INN = "hyperlinkimab"
IND_NUMBER = "IND-2024-08811"
PROTOCOL_VERSION = "v3.1 (19 Jan 2025)"


# ─────────────────────────────────────────────────────────────────────────────
# Per-document definitions
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class TableRow:
    cells: list[str]


@dataclass
class Section:
    heading: str
    level: int
    paragraphs: list[str] = field(default_factory=list)
    table_title: str = ""
    table_rows: list[TableRow] = field(default_factory=list)


@dataclass
class CsrDoc:
    study_id: str
    filename: str
    phase: str
    title: str
    synopsis: str
    sections: list[Section]


# ─────────────────────────────────────────────────────────────────────────────
# Document 1 — Phase 1 PK
# ─────────────────────────────────────────────────────────────────────────────

CSR_001 = CsrDoc(
    study_id="SP-2026-001",
    filename="csr-sp-2026-001.docx",
    phase="Phase 1",
    title="Clinical Study Report SP-2026-001\nA Phase 1 Single-Ascending-Dose Pharmacokinetic Study\nof Hyperlinkimab in Healthy Adult Volunteers",
    synopsis=(
        "This report describes a Phase 1, open-label, single-ascending-dose (SAD) study "
        "evaluating the pharmacokinetics (PK), safety, and tolerability of Hyperlinkimab "
        "in 48 healthy adult volunteers across six dose cohorts (1, 3, 10, 30, 100, and "
        "300 mg IV). The study was conducted at three clinical sites under "
        f"IND {IND_NUMBER}. "
        "Hyperlinkimab demonstrated linear PK across the dose range examined, with a "
        "half-life of approximately 18 days supporting once-monthly dosing. No dose-limiting "
        "toxicities were observed. These findings informed dose selection for the subsequent "
        "Phase 2a study (CSR SP-2026-002)."
    ),
    sections=[
        Section(
            heading="1. Introduction",
            level=1,
            paragraphs=[
                f"{DRUG} ({INN}) is a fully human IgG1 monoclonal antibody targeting "
                "soluble Interleukin-31 Receptor Alpha (IL-31RA). Pre-clinical data "
                "established dose-dependent receptor occupancy and a clean safety profile "
                "in cynomolgus monkeys up to 300 mg/kg IV.",
                "This Phase 1 study represents the first-in-human (FIH) administration "
                f"of {DRUG}. The primary objectives were to characterize the single-dose "
                "PK profile, assess immunogenicity, and establish the maximum tolerated dose.",
            ],
        ),
        Section(
            heading="2. Study Objectives",
            level=1,
            paragraphs=[
                "Primary: Characterize the PK of Hyperlinkimab following single IV doses "
                "of 1-300 mg in healthy volunteers.",
                "Secondary: Evaluate safety, tolerability, and immunogenicity (anti-drug "
                "antibody [ADA] incidence) at each dose level.",
                "Exploratory: Assess receptor occupancy via soluble IL-31RA pharmacodynamic "
                "(PD) markers at 24 h, 72 h, and Day 14 post-dose.",
            ],
        ),
        Section(
            heading="3. Methods",
            level=1,
            paragraphs=[
                f"Study design: Open-label, dose-escalating SAD study. Protocol version {PROTOCOL_VERSION}.",
                "Population: 48 healthy adults aged 18-55 years, BMI 18-30 kg/m², "
                "no clinically significant medical history. 8 subjects per cohort.",
                "Dosing: Single IV infusion over 60 minutes. Dose escalation gated by "
                "independent Data Safety Monitoring Board (DSMB) review.",
                "PK sampling: Pre-dose; 0.5, 1, 2, 4, 8, 24, 48, 72 h; Days 7, 14, 21, "
                "28, 42, 56, 84 post-dose.",
            ],
        ),
        Section(
            heading="4. Study Subjects",
            level=1,
            paragraphs=[
                "48 subjects enrolled; 46 completed (2 withdrew consent, unrelated to study). "
                "Demographic characteristics are summarised in Table 1 below.",
            ],
            table_title="Table 1. Demographic Summary — Study SP-2026-001",
            table_rows=[
                TableRow(["Parameter", "1 mg (n=8)", "10 mg (n=8)", "100 mg (n=8)", "300 mg (n=8)", "Total (n=48)"]),
                TableRow(["Age, mean (SD)", "31.4 (6.2)", "29.8 (5.5)", "33.1 (7.0)", "30.6 (5.8)", "31.2 (6.1)"]),
                TableRow(["Male, n (%)", "5 (62.5)", "4 (50.0)", "6 (75.0)", "5 (62.5)", "28 (58.3)"]),
                TableRow(["Weight (kg), mean", "72.3", "74.1", "70.9", "73.5", "72.7"]),
                TableRow(["BMI (kg/m²), mean", "23.8", "24.2", "23.5", "24.0", "23.9"]),
            ],
        ),
        Section(
            heading="5. Pharmacokinetic Results",
            level=1,
            paragraphs=[
                "Hyperlinkimab exhibited linear, dose-proportional PK across the 1-300 mg "
                "dose range. Mean Cmax and AUC0-inf increased proportionally with dose. "
                "The terminal half-life (t½) was 17.8 ± 2.3 days (mean ± SD) "
                "across all cohorts, consistent across dose levels.",
                "Volume of distribution at steady state (Vss) was 62.4 ± 8.1 mL/kg, "
                "consistent with predominantly vascular distribution. Clearance (CL) was "
                "3.1 ± 0.4 mL/day/kg.",
            ],
            table_title="Table 2. PK Parameter Summary — Study SP-2026-001",
            table_rows=[
                TableRow(["PK Parameter", "1 mg", "10 mg", "100 mg", "300 mg"]),
                TableRow(["Cmax (μg/mL)", "0.24", "2.41", "24.8", "73.2"]),
                TableRow(["AUC0-inf (μg·day/mL)", "1.12", "11.4", "112", "336"]),
                TableRow(["t½ (days)", "17.6", "17.9", "18.1", "17.8"]),
                TableRow(["Vss (mL/kg)", "61.2", "63.1", "62.8", "62.4"]),
                TableRow(["CL (mL/day/kg)", "3.0", "3.1", "3.2", "3.1"]),
            ],
        ),
        Section(
            heading="6. Safety Summary",
            level=1,
            paragraphs=[
                "Hyperlinkimab was well tolerated across all dose levels. No dose-limiting "
                "toxicities, serious adverse events (SAEs), or deaths were reported. "
                "Treatment-emergent adverse events (TEAEs) were mild (Grade 1) in 79% "
                "of cases and moderate (Grade 2) in 21%.",
                "Most common TEAEs: injection-site reaction (18.8%), headache (14.6%), "
                "fatigue (10.4%). No Grade 3 or higher events were observed.",
                "ADA incidence: 2/48 subjects (4.2%) developed low-titre, non-neutralising "
                "ADAs at Day 84; no impact on PK or safety was noted.",
            ],
            table_title="Table 3. Treatment-Emergent Adverse Events (≥1% any cohort) — SP-2026-001",
            table_rows=[
                TableRow(["Adverse Event", "1 mg (n=8)", "10 mg (n=8)", "100 mg (n=8)", "300 mg (n=8)", "Total (n=48)"]),
                TableRow(["Any TEAE", "3 (37.5%)", "4 (50.0%)", "5 (62.5%)", "6 (75.0%)", "27 (56.3%)"]),
                TableRow(["Infusion-site reaction", "1 (12.5%)", "1 (12.5%)", "2 (25.0%)", "5 (62.5%)", "9 (18.8%)"]),
                TableRow(["Headache", "1 (12.5%)", "2 (25.0%)", "2 (25.0%)", "2 (25.0%)", "7 (14.6%)"]),
                TableRow(["Fatigue", "0", "1 (12.5%)", "2 (25.0%)", "2 (25.0%)", "5 (10.4%)"]),
                TableRow(["Nausea", "0", "0", "1 (12.5%)", "1 (12.5%)", "2 (4.2%)"]),
            ],
        ),
        Section(
            heading="7. Discussion and Conclusions",
            level=1,
            paragraphs=[
                f"{DRUG} demonstrated a favourable PK and safety profile in healthy "
                "volunteers. Linear dose-proportional PK, a long half-life (~18 days), "
                "and a clean safety profile at doses up to 300 mg IV support selection "
                "of a 150 mg SC dose for the Phase 2a efficacy study.",
                "The PK data from this study were used to develop a population PK model "
                "that informed dose selection for all subsequent studies in the programme.",
            ],
        ),
        Section(
            heading="8. Cross-References to Related Clinical Study Reports",
            level=1,
            paragraphs=[
                "The following cross-references link to supporting evidence in other "
                "clinical study reports within this regulatory dossier. Reviewers are "
                "directed to the referenced sections for complementary information.",
                # CROSS-REF 1: SP-2026-001 -> SP-2026-002 Section 2.5
                "Safety data from the Phase 2a patient population, for comparison with "
                "the healthy volunteer safety profile described in Section 6 above, are "
                "presented in CSR SP-2026-002 Section 2.5 (Safety Overview).",
                # CROSS-REF 2: SP-2026-001 -> SP-2026-003 Appendix 16.1.1
                "Individual patient PK listings used to construct the population PK model "
                "are reproduced in CSR SP-2026-003 Appendix 16.1.1 (Pharmacokinetic "
                "Listings from Prior Studies) for the dose-finding population.",
                # CROSS-REF 3: SP-2026-001 -> SP-2026-004 Table 14.2.1.1
                "Integrated demographic data across all four Phase 1/2/3 studies are "
                "presented in CSR SP-2026-004 Table 14.2.1.1 (Integrated Demographic "
                "Summary, All Studies).",
            ],
        ),
    ],
)


# ─────────────────────────────────────────────────────────────────────────────
# Document 2 — Phase 2a Efficacy
# ─────────────────────────────────────────────────────────────────────────────

CSR_002 = CsrDoc(
    study_id="SP-2026-002",
    filename="csr-sp-2026-002.docx",
    phase="Phase 2a",
    title="Clinical Study Report SP-2026-002\nA Phase 2a Randomized, Double-Blind, Placebo-Controlled\nEfficacy and Safety Study of Hyperlinkimab in Adults\nwith Moderate-to-Severe Atopic Dermatitis",
    synopsis=(
        "This report describes a Phase 2a, randomized, double-blind, placebo-controlled "
        "study evaluating the efficacy and safety of Hyperlinkimab 150 mg SC administered "
        "once monthly in 120 adult patients with moderate-to-severe atopic dermatitis (AD) "
        "(IGA score ≥3). Patients were randomized 2:1 (active:placebo) for 16 weeks. "
        "The primary endpoint, IGA 0/1 response at Week 16, was achieved in 54.4% of "
        "Hyperlinkimab-treated patients vs 10.0% placebo (p<0.001). Hyperlinkimab "
        "was well tolerated with a safety profile consistent with CSR SP-2026-001 "
        "healthy volunteer data."
    ),
    sections=[
        Section(
            heading="1. Introduction",
            level=1,
            paragraphs=[
                f"{DRUG} targets IL-31RA, a key mediator of itch and inflammation in "
                "atopic dermatitis. Proof-of-concept for IL-31RA blockade in AD is "
                "established; this Phase 2a study represents the first patient exposure.",
                "Dose selection (150 mg SC monthly) was based on the PK data from the "
                "Phase 1 study (CSR SP-2026-001), supported by population PK modelling.",
            ],
        ),
        Section(
            heading="2. Study Objectives",
            level=1,
            paragraphs=[
                "Primary: Proportion of patients achieving IGA 0/1 response at Week 16.",
                "Secondary: ○ EASI-75 response at Week 16. ○ SCORAD change from "
                "baseline. ○ NRS itch score reduction ≥4 points. ○ Time to "
                "first IGA 0/1 response.",
                "Safety: TEAEs, SAEs, laboratory parameters, vital signs, "
                "immunogenicity through Week 20.",
            ],
        ),
        Section(
            heading="3. Methods",
            level=1,
            paragraphs=[
                "Design: Randomized 2:1, double-blind, placebo-controlled, 16-week treatment "
                "period with 4-week safety follow-up. Three clinical sites (US).",
                f"Protocol: {PROTOCOL_VERSION}. IND {IND_NUMBER}.",
                "Patients: 120 adults; IGA ≥3, EASI ≥16, BSA ≥10% at screening.",
                "Treatment: Hyperlinkimab 150 mg SC Q4W (n=80) or matching placebo (n=40) "
                "for 4 doses (Weeks 0, 4, 8, 12).",
            ],
        ),
        Section(
            heading="4. Study Subjects",
            level=1,
            paragraphs=[
                "120 patients were randomized; 113 completed the 16-week treatment period "
                "(7 discontinued: 4 withdrew consent, 2 lost to follow-up, 1 adverse event).",
            ],
            table_title="Table 4. Demographic Summary — Study SP-2026-002",
            table_rows=[
                TableRow(["Parameter", "Hyperlinkimab 150 mg (n=80)", "Placebo (n=40)", "Total (n=120)"]),
                TableRow(["Age, mean (SD)", "36.4 (10.2)", "37.1 (9.8)", "36.7 (10.1)"]),
                TableRow(["Male, n (%)", "42 (52.5)", "22 (55.0)", "64 (53.3)"]),
                TableRow(["IGA score, mean", "3.4", "3.3", "3.4"]),
                TableRow(["EASI score, mean", "31.2", "30.8", "31.1"]),
                TableRow(["BSA affected (%)", "42.1", "41.6", "41.9"]),
                TableRow(["Prior biologic use, n (%)", "18 (22.5)", "9 (22.5)", "27 (22.5)"]),
            ],
        ),
        Section(
            heading="5. Efficacy Results",
            level=1,
            paragraphs=[
                "Primary endpoint: IGA 0/1 at Week 16 was achieved in 54.4% (43/79) of "
                "Hyperlinkimab patients versus 10.0% (4/40) placebo; OR 11.0, 95% CI "
                "[3.6, 33.6]; p<0.001 (logistic regression, ITT population).",
                "EASI-75 response at Week 16: 62.5% active vs 12.5% placebo (p<0.001). "
                "Mean EASI reduction from baseline: -22.1 active vs -4.8 placebo.",
                "NRS itch: mean reduction -4.2 active vs -1.1 placebo at Week 16 (p<0.001).",
            ],
            table_title="Table 5. Primary and Key Secondary Efficacy Results at Week 16 — SP-2026-002",
            table_rows=[
                TableRow(["Endpoint", "Hyperlinkimab 150 mg", "Placebo", "OR / Diff (95% CI)", "p-value"]),
                TableRow(["IGA 0/1, n (%)", "43/79 (54.4%)", "4/40 (10.0%)", "OR 11.0 [3.6, 33.6]", "<0.001"]),
                TableRow(["EASI-75, n (%)", "50/80 (62.5%)", "5/40 (12.5%)", "OR 11.6 [4.1, 32.8]", "<0.001"]),
                TableRow(["EASI change, mean", "-22.1", "-4.8", "Diff -17.3 [-21.1, -13.5]", "<0.001"]),
                TableRow(["NRS itch change", "-4.2", "-1.1", "Diff -3.1 [-4.0, -2.2]", "<0.001"]),
            ],
        ),
        Section(
            heading="2.5. Safety Overview",
            level=1,
            paragraphs=[
                "Hyperlinkimab 150 mg SC was well tolerated over 16 weeks. The safety "
                "profile in the AD patient population was consistent with the healthy "
                "volunteer data reported in CSR SP-2026-001, Section 6.",
                "TEAEs were reported in 61.3% active vs 52.5% placebo. Most common "
                "TEAEs: nasopharyngitis (15.0% vs 12.5%), injection-site reaction "
                "(13.8% vs 7.5%), URTI (11.3% vs 10.0%).",
                "One SAE (cellulitis, Day 88, active arm) resolved without sequelae and "
                "was considered unrelated to study drug. No deaths.",
                "ADA incidence: 3/80 (3.8%), all non-neutralising, no PK or safety impact.",
            ],
            table_title="Table 6. TEAEs (≥5% either arm) — Study SP-2026-002",
            table_rows=[
                TableRow(["Adverse Event", "Hyperlinkimab 150 mg (n=80)", "Placebo (n=40)"]),
                TableRow(["Any TEAE", "49 (61.3%)", "21 (52.5%)"]),
                TableRow(["Nasopharyngitis", "12 (15.0%)", "5 (12.5%)"]),
                TableRow(["Injection-site reaction", "11 (13.8%)", "3 (7.5%)"]),
                TableRow(["URTI", "9 (11.3%)", "4 (10.0%)"]),
                TableRow(["Headache", "7 (8.8%)", "3 (7.5%)"]),
                TableRow(["Conjunctivitis", "5 (6.3%)", "1 (2.5%)"]),
                TableRow(["SAE", "1 (1.3%)", "0"]),
            ],
        ),
        Section(
            heading="7. Discussion and Conclusions",
            level=1,
            paragraphs=[
                f"The Phase 2a data for {DRUG} demonstrate a clinically meaningful and "
                "statistically significant improvement across all efficacy endpoints in "
                "moderate-to-severe AD. The IGA 0/1 response rate of 54.4% exceeds the "
                "pre-specified success criterion of 40%.",
                "The benefit-risk profile supports progression to the dose-finding Phase 2b "
                "study (CSR SP-2026-003) to confirm the optimal dose regimen.",
            ],
        ),
        Section(
            heading="8. Cross-References to Related Clinical Study Reports",
            level=1,
            paragraphs=[
                "The following cross-references direct reviewers to supporting evidence "
                "in related study reports within this regulatory submission:",
                # CROSS-REF 1: SP-2026-002 -> SP-2026-001 Section 5.3.1
                "The PK profile that underpinned dose selection for this study is "
                "documented in CSR SP-2026-001 Section 5.3.1 (Pharmacokinetic Results, "
                "Population PK Model). The simulated steady-state Ctrough at 150 mg SC "
                "Q4W supported sustained target coverage.",
                # CROSS-REF 2: SP-2026-002 -> SP-2026-003 Listing 16.2.5
                "Adverse event line-listing data from the Phase 2a and Phase 2b "
                "populations, consolidated for integrated safety analysis, are presented "
                "in CSR SP-2026-003 Listing 16.2.5 (AE Summary by Preferred Term, "
                "All Phase 2 Studies).",
                # CROSS-REF 3: SP-2026-002 -> SP-2026-004 Section 5.3.5
                "The confirmed efficacy of 150 mg SC monthly in the broader Phase 3 "
                "population, including a pre-specified subgroup analysis by prior biologic "
                "use, is presented in CSR SP-2026-004 Section 5.3.5 (Confirmatory "
                "Efficacy Analysis, Phase 3 Population).",
            ],
        ),
    ],
)


# ─────────────────────────────────────────────────────────────────────────────
# Document 3 — Phase 2b Dose-Finding
# ─────────────────────────────────────────────────────────────────────────────

CSR_003 = CsrDoc(
    study_id="SP-2026-003",
    filename="csr-sp-2026-003.docx",
    phase="Phase 2b",
    title="Clinical Study Report SP-2026-003\nA Phase 2b Randomized Dose-Ranging Study to Identify\nthe Optimal Therapeutic Dose of Hyperlinkimab in Adults\nwith Moderate-to-Severe Atopic Dermatitis",
    synopsis=(
        "This report describes a Phase 2b, randomized, double-blind, placebo-controlled, "
        "dose-ranging study in 240 adult patients with moderate-to-severe atopic dermatitis. "
        "Four doses of Hyperlinkimab (75, 150, 300 mg SC Q4W) and placebo were compared "
        "over 24 weeks. The 150 mg Q4W dose was confirmed as optimal, achieving IGA 0/1 "
        "in 58.1% of patients at Week 24 (vs 8.3% placebo, p<0.001). The 300 mg dose "
        "offered no incremental benefit over 150 mg. Dose-response modelling (Figure 11) "
        "confirmed 150 mg as the Phase 3 dose."
    ),
    sections=[
        Section(
            heading="1. Introduction",
            level=1,
            paragraphs=[
                "The Phase 2a study (CSR SP-2026-002) established proof-of-concept for "
                f"{DRUG} in atopic dermatitis, but only a single dose (150 mg SC Q4W) "
                "was evaluated. This Phase 2b study explores doses from 75 to 300 mg to "
                "characterise the dose-response relationship and confirm the optimal dose "
                "for the Phase 3 programme.",
                "Dose selection for this study was informed by the population PK model "
                "developed from Phase 1 data (CSR SP-2026-001) and Phase 2a PK observations.",
            ],
        ),
        Section(
            heading="2. Study Objectives",
            level=1,
            paragraphs=[
                "Primary: Dose-response relationship for IGA 0/1 at Week 24.",
                "Secondary: EASI-75, NRS itch, BSA affected; dose-response modelling; "
                "selection of Phase 3 dose.",
                "Safety: TEAE profile across all dose levels through Week 28.",
            ],
        ),
        Section(
            heading="3. Methods",
            level=1,
            paragraphs=[
                "Design: Randomized 1:1:1:1 (75 mg : 150 mg : 300 mg : placebo), "
                "double-blind, 24-week treatment, 4-week follow-up. Six clinical sites.",
                f"Protocol: {PROTOCOL_VERSION}. IND {IND_NUMBER}.",
                "Patients: 240 adults, IGA ≥3, EASI ≥16, BSA ≥10%; "
                "prior topical therapy failure required.",
                "Dosing: SC injection Q4W; Weeks 0, 4, 8, 12, 16, 20 (6 doses total).",
            ],
        ),
        Section(
            heading="4. Study Subjects",
            level=1,
            paragraphs=[
                "240 patients randomized (60 per arm); 221 completed 24-week treatment "
                "(19 discontinued: 10 withdrew consent, 6 lost-to-follow-up, 3 AEs).",
            ],
            table_title="Table 7. Demographic Summary — Study SP-2026-003",
            table_rows=[
                TableRow(["Parameter", "75 mg (n=60)", "150 mg (n=60)", "300 mg (n=60)", "Placebo (n=60)", "Total (n=240)"]),
                TableRow(["Age, mean (SD)", "37.2 (9.4)", "36.8 (10.1)", "38.1 (9.8)", "36.5 (10.3)", "37.1 (9.9)"]),
                TableRow(["Male, n (%)", "31 (51.7)", "33 (55.0)", "30 (50.0)", "32 (53.3)", "126 (52.5)"]),
                TableRow(["Baseline IGA, mean", "3.4", "3.5", "3.4", "3.3", "3.4"]),
                TableRow(["Baseline EASI, mean", "30.4", "31.8", "31.2", "30.6", "31.0"]),
            ],
        ),
        Section(
            heading="5. Efficacy Results",
            level=1,
            paragraphs=[
                "A clear dose-response was observed for IGA 0/1 at Week 24: 75 mg "
                "36.7%, 150 mg 58.1%, 300 mg 60.0%, placebo 8.3%. The 150 mg and "
                "300 mg arms were statistically superior to placebo (both p<0.001). "
                "The 300 mg arm was not statistically superior to 150 mg (p=0.82).",
                "Dose-response modelling using an Emax model is summarised in Figure 11. "
                "The estimated ED90 was 121 mg, confirming the 150 mg dose achieves "
                "near-maximal effect.",
            ],
            table_title="Table 14.2.1.1. Primary Efficacy Results (IGA 0/1 at Week 24) — SP-2026-003",
            table_rows=[
                TableRow(["Dose Arm", "IGA 0/1 Rate", "OR vs Placebo (95% CI)", "p-value"]),
                TableRow(["75 mg SC Q4W", "22/60 (36.7%)", "6.5 [2.6, 16.3]", "<0.001"]),
                TableRow(["150 mg SC Q4W", "35/60 (58.3%)", "14.9 [5.7, 39.1]", "<0.001"]),
                TableRow(["300 mg SC Q4W", "36/60 (60.0%)", "16.0 [6.1, 42.0]", "<0.001"]),
                TableRow(["Placebo", "5/60 (8.3%)", "—", "—"]),
            ],
        ),
        Section(
            heading="Table 14.3.1. Dose-Response Parameter Estimates (Emax Model)",
            level=2,
            table_title="Table 14.3.1. Emax Model Parameters — Dose-Response Analysis SP-2026-003",
            table_rows=[
                TableRow(["Parameter", "Estimate", "95% CI", "Interpretation"]),
                TableRow(["Emax (max effect)", "0.617", "[0.541, 0.693]", "Maximum IGA 0/1 rate"]),
                TableRow(["E0 (baseline rate)", "0.083", "[0.028, 0.138]", "Placebo response rate"]),
                TableRow(["ED50 (mg)", "67.4", "[48.2, 86.6]", "Half-maximal effective dose"]),
                TableRow(["ED90 (mg)", "121", "[102, 140]", "90% of maximal dose"]),
                TableRow(["Hill coefficient", "1.8", "[1.2, 2.4]", "Slope of dose-response curve"]),
            ],
        ),
        Section(
            heading="6. Safety Summary",
            level=1,
            paragraphs=[
                "Hyperlinkimab was well tolerated across all dose levels. Overall TEAE "
                "rates were: 75 mg 63.3%, 150 mg 66.7%, 300 mg 68.3%, placebo 58.3%. "
                "The incremental TEAE rate at 300 mg vs 150 mg was not clinically meaningful.",
                "One SAE (anaphylaxis, 300 mg arm) occurred and resolved with treatment. "
                "No deaths. Grade 3 TEAEs: 3 events across 3 patients (1 per dose arm).",
            ],
        ),
        Section(
            heading="7. Discussion and Conclusions",
            level=1,
            paragraphs=[
                "The Phase 2b dose-finding data confirm 150 mg SC Q4W as the optimal "
                f"therapeutic dose for {DRUG} in moderate-to-severe AD. The dose-response "
                "plateau above 150 mg, combined with the clean safety profile, supports "
                "a single-dose Phase 3 design using 150 mg SC Q4W.",
                "The population PK/PD modelling, consistent with the Phase 1 model "
                "(CSR SP-2026-001), predicts that 150 mg SC Q4W achieves >95% target "
                "coverage throughout the dosing interval.",
            ],
        ),
        Section(
            heading="8. Cross-References to Related Clinical Study Reports",
            level=1,
            paragraphs=[
                "Reviewers are directed to the following cross-references within this "
                "dossier for supporting evidence:",
                # CROSS-REF 1: SP-2026-003 -> SP-2026-001 Figure 11
                "The PK exposure-response relationship used to construct the Emax dose-"
                "response model is illustrated in CSR SP-2026-001 Figure 11 "
                "(Exposure-Response: Receptor Occupancy vs. Cmax). The Phase 1 PK/PD "
                "data provided the pharmacodynamic anchoring for the Phase 2b model.",
                # CROSS-REF 2: SP-2026-003 -> SP-2026-002 Table 14.2.1.1
                "The Phase 2a primary efficacy result that established proof-of-concept "
                "for the 150 mg dose, used as a calibration anchor in the dose-response "
                "model, is provided in CSR SP-2026-002 Table 14.2.1.1 (Efficacy Results "
                "at Week 16, Phase 2a ITT Population).",
                # CROSS-REF 3: SP-2026-003 -> SP-2026-004 Appendix 16.2.6
                "The combined Phase 2b and Phase 3 patient-level safety data used for "
                "the integrated benefit-risk assessment are compiled in "
                "CSR SP-2026-004 Appendix 16.2.6 (Individual Patient Safety Listings, "
                "All Phase 2/3 Studies).",
            ],
        ),
    ],
)


# ─────────────────────────────────────────────────────────────────────────────
# Document 4 — Phase 3 Pivotal
# ─────────────────────────────────────────────────────────────────────────────

CSR_004 = CsrDoc(
    study_id="SP-2026-004",
    filename="csr-sp-2026-004.docx",
    phase="Phase 3",
    title="Clinical Study Report SP-2026-004\nA Phase 3 Multi-Center, Randomized, Double-Blind,\nPlacebo-Controlled Confirmatory Efficacy and Safety Study\nof Hyperlinkimab 150 mg SC in Adults with\nModerate-to-Severe Atopic Dermatitis",
    synopsis=(
        "This pivotal Phase 3 report documents the confirmatory efficacy and safety of "
        "Hyperlinkimab 150 mg SC administered once monthly in 416 adults with "
        "moderate-to-severe atopic dermatitis across 22 international study sites. "
        "Patients were randomized 1:1 active:placebo for 52 weeks. The co-primary "
        "endpoints were achieved: IGA 0/1 at Week 16 (58.7% active vs 9.6% placebo, "
        "p<0.0001) and EASI-75 at Week 16 (65.4% vs 11.5%, p<0.0001). The 52-week "
        "safety profile was consistent with prior studies (CSR SP-2026-001, "
        "CSR SP-2026-002, CSR SP-2026-003). These data support a regulatory submission "
        "for Hyperlinkimab in moderate-to-severe atopic dermatitis."
    ),
    sections=[
        Section(
            heading="1. Introduction",
            level=1,
            paragraphs=[
                f"This confirmatory Phase 3 study was designed based on data from the "
                "Phase 1 (CSR SP-2026-001), Phase 2a (CSR SP-2026-002), and Phase 2b "
                "(CSR SP-2026-003) studies that collectively established the PK profile, "
                "proof-of-concept, and optimal dose of Hyperlinkimab.",
                "The 52-week duration was selected to assess both the induction of "
                "response (Weeks 0-16, primary endpoint period) and long-term maintenance "
                "of response (Weeks 16-52, key secondary endpoint period).",
            ],
        ),
        Section(
            heading="2. Study Objectives",
            level=1,
            paragraphs=[
                "Co-primary: IGA 0/1 at Week 16; EASI-75 at Week 16.",
                "Key secondary: Durability of response at Week 52; NRS itch reduction "
                "≥4 points; DLQI improvement ≥4 points; SCORAD-50 response.",
                "Exploratory: Subgroup analysis by prior biologic use; biomarker "
                "analyses (serum IgE, dupilumab-experienced subgroup).",
                "Safety: Comprehensive 52-week TEAE profile, SAEs, laboratory monitoring, "
                "vital signs, immunogenicity.",
            ],
        ),
        Section(
            heading="3. Methods",
            level=1,
            paragraphs=[
                "Design: Randomized 1:1, double-blind, placebo-controlled, 52-week "
                "treatment period. 22 sites across USA, EU, and Japan.",
                f"Protocol: {PROTOCOL_VERSION}. IND {IND_NUMBER}.",
                "Population: 416 adults; IGA ≥3, EASI ≥16, BSA ≥10%; "
                "prior systemic therapy failure or intolerance required.",
                "Treatment: Hyperlinkimab 150 mg SC Q4W (n=208) or matching placebo "
                "(n=208); 13 doses total (Weeks 0, 4, 8, ..., 48).",
            ],
        ),
        Section(
            heading="4. Study Subjects",
            level=1,
            paragraphs=[
                "416 patients randomized; 381 completed the 52-week treatment period "
                "(35 discontinued: 18 withdrew consent, 9 lost to follow-up, "
                "5 AEs, 3 protocol deviations).",
            ],
            table_title="Table 14.2.1.1. Integrated Demographic Summary, All Studies",
            table_rows=[
                TableRow(["Parameter", "SP-2026-001 (n=48)", "SP-2026-002 (n=120)", "SP-2026-003 (n=240)", "SP-2026-004 (n=416)", "All Active Arms (n=624)"]),
                TableRow(["Age, mean (SD)", "31.2 (6.1)", "36.7 (10.1)", "37.1 (9.9)", "37.8 (10.5)", "36.9 (10.2)"]),
                TableRow(["Male, n (%)", "28 (58.3)", "64 (53.3)", "126 (52.5)", "219 (52.6)", "409 (53.1)"]),
                TableRow(["White, n (%)", "38 (79.2)", "88 (73.3)", "175 (72.9)", "292 (70.2)", "555 (72.1)"]),
                TableRow(["Asian, n (%)", "6 (12.5)", "18 (15.0)", "39 (16.3)", "72 (17.3)", "129 (16.7)"]),
                TableRow(["Prior biologic, n (%)", "N/A", "27 (22.5)", "54 (22.5)", "94 (22.6)", "175 (22.6)"]),
            ],
        ),
        Section(
            heading="5. Efficacy Results",
            level=1,
            paragraphs=[
                "Both co-primary endpoints were met at Week 16 with highly significant "
                "results. IGA 0/1: 58.7% (121/206) active vs 9.6% (20/208) placebo; "
                "OR 13.5, 95% CI [7.9, 23.0], p<0.0001. EASI-75: 65.4% (135/206) "
                "active vs 11.5% (24/208) placebo; OR 14.2, 95% CI [8.5, 23.7], p<0.0001.",
                "Durability: IGA 0/1 response was maintained through Week 52 in 74.4% "
                "of Week-16 responders. EASI-75 was maintained through Week 52 in 79.3% "
                "of Week-16 EASI-75 responders.",
            ],
            table_title="Table 8. Co-Primary and Key Secondary Efficacy Results — Study SP-2026-004",
            table_rows=[
                TableRow(["Endpoint", "Hyperlinkimab 150 mg (n=208)", "Placebo (n=208)", "OR / Diff (95% CI)", "p-value"]),
                TableRow(["IGA 0/1 Week 16", "121/206 (58.7%)", "20/208 (9.6%)", "OR 13.5 [7.9, 23.0]", "<0.0001"]),
                TableRow(["EASI-75 Week 16", "135/206 (65.5%)", "24/208 (11.5%)", "OR 14.2 [8.5, 23.7]", "<0.0001"]),
                TableRow(["NRS≥4 Week 16", "126/206 (61.2%)", "19/208 (9.1%)", "OR 15.9 [9.2, 27.5]", "<0.0001"]),
                TableRow(["DLQI≥4 Week 16", "148/206 (71.8%)", "42/208 (20.2%)", "OR 10.8 [6.6, 17.8]", "<0.0001"]),
                TableRow(["IGA 0/1 Week 52", "90/121 (74.4%)*", "—", "—", "—"]),
            ],
        ),
        Section(
            heading="5.3.5. Confirmatory Efficacy Analysis, Phase 3 Population",
            level=2,
            paragraphs=[
                "Pre-specified subgroup analyses confirmed consistent efficacy across all "
                "major demographic and disease characteristic subgroups. Treatment effect "
                "on IGA 0/1 at Week 16 was homogeneous across: age groups (<35 vs ≥35), "
                "sex, race (White/Asian/Other), disease severity (IGA 3 vs 4), and "
                "prior biologic use (naive vs experienced).",
                "In the prior-biologic-experienced subgroup (n=94 active, n=94 placebo), "
                "IGA 0/1 was achieved by 44.7% active vs 7.4% placebo (OR 10.0, "
                "95% CI [4.1, 24.5], p<0.0001). This subgroup analysis was pre-specified "
                "to align with Phase 2a findings in CSR SP-2026-002 Section 5.",
            ],
        ),
        Section(
            heading="6. Safety Summary",
            level=1,
            paragraphs=[
                "The 52-week safety profile was consistent with earlier studies. Overall "
                "TEAE rate: 78.4% active vs 70.7% placebo. TEAEs were predominantly "
                "mild-to-moderate.",
                "SAEs: 8 active (3.8%) vs 5 placebo (2.4%); no fatal SAEs. Most frequent "
                "SAE: cellulitis (3 active, 1 placebo). No anaphylaxis in Phase 3 (vs 1 "
                "event in Phase 2b; see CSR SP-2026-003).",
                "Long-term safety (Week 52) remained stable; no new safety signals emerged "
                "beyond those identified in the Phase 2 programme.",
            ],
            table_title="Table 9. TEAEs (≥5% either arm) at 52 Weeks — Study SP-2026-004",
            table_rows=[
                TableRow(["Adverse Event", "Hyperlinkimab 150 mg (n=208)", "Placebo (n=208)"]),
                TableRow(["Any TEAE", "163 (78.4%)", "147 (70.7%)"]),
                TableRow(["Nasopharyngitis", "38 (18.3%)", "32 (15.4%)"]),
                TableRow(["Injection-site reaction", "31 (14.9%)", "12 (5.8%)"]),
                TableRow(["URTI", "27 (13.0%)", "22 (10.6%)"]),
                TableRow(["Conjunctivitis", "18 (8.7%)", "4 (1.9%)"]),
                TableRow(["Headache", "16 (7.7%)", "14 (6.7%)"]),
                TableRow(["Arthralgia", "13 (6.3%)", "10 (4.8%)"]),
                TableRow(["SAEs total", "8 (3.8%)", "5 (2.4%)"]),
            ],
        ),
        Section(
            heading="7. Discussion and Conclusions",
            level=1,
            paragraphs=[
                f"This pivotal Phase 3 study confirms that {DRUG} 150 mg SC Q4W "
                "provides clinically meaningful and statistically significant improvement "
                "in adults with moderate-to-severe atopic dermatitis.",
                "The totality of evidence from Phase 1 (CSR SP-2026-001) through Phase 3 "
                "supports a regulatory submission for approval of Hyperlinkimab for the "
                "treatment of moderate-to-severe atopic dermatitis in adults.",
            ],
        ),
        Section(
            heading="8. Cross-References to Related Clinical Study Reports",
            level=1,
            paragraphs=[
                "The following cross-references are provided to assist reviewers in "
                "navigating the integrated clinical programme for Hyperlinkimab:",
                # CROSS-REF 1: SP-2026-004 -> SP-2026-001 Section 5.3.1
                "The baseline PK characterisation of Hyperlinkimab, including the "
                "population PK model used for Phase 3 dose confirmation, is documented "
                "in CSR SP-2026-001 Section 5.3.1 (Pharmacokinetic Results). The "
                "model predicts Ctrough >95% target occupancy throughout the Q4W interval "
                "at 150 mg SC.",
                # CROSS-REF 2: SP-2026-004 -> SP-2026-002 Section 2.5
                "The Phase 2a safety reference dataset, used for integrated safety "
                "comparisons in the Phase 3 benefit-risk assessment, is described in "
                "CSR SP-2026-002 Section 2.5 (Safety Overview). The safety profile "
                "observed in Phase 3 is consistent with Phase 2a findings.",
                # CROSS-REF 3: SP-2026-004 -> SP-2026-003 Table 14.3.1
                "The dose-response model summary table confirming 150 mg as the optimal "
                "dose is presented in CSR SP-2026-003 Table 14.3.1 (Emax Model Parameter "
                "Estimates). These data underpin the single-dose design of this Phase 3 study.",
            ],
        ),
    ],
)


ALL_DOCS = (CSR_001, CSR_002, CSR_003, CSR_004)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX rendering
# ─────────────────────────────────────────────────────────────────────────────

def _set_col_widths(table, widths_cm: list[float]) -> None:
    """Set approximate column widths (cm) on a docx table."""
    from docx.oxml import OxmlElement
    from docx.shared import Cm
    tbl = table._tbl
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(Cm(w).twips)))
        tblGrid.append(gridCol)
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    tbl.insert(0, tblGrid)


def _add_table(doc: Document, section: Section) -> None:
    """Add a formatted table to the document."""
    if not section.table_rows:
        return
    if section.table_title:
        p = doc.add_paragraph(section.table_title)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9)

    rows = section.table_rows
    ncols = len(rows[0].cells)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx]
        for c_idx, text in enumerate(row.cells):
            cell = tr.cells[c_idx]
            cell.text = text
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(text)
            run.font.size = Pt(8)
            if r_idx == 0:
                run.bold = True
                cell._tc.get_or_add_tcPr()
                shd = etree.SubElement(cell._tc.tcPr, qn("w:shd"))
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9E1F2")
    doc.add_paragraph("")


def _render_doc(csr: CsrDoc, out_path: Path) -> None:
    """Render one CSR to a .docx file."""
    docx = Document()

    # ── Title page ────────────────────────────────────────────────────────────
    title_para = docx.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(csr.title)
    run.bold = True
    run.font.size = Pt(14)

    docx.add_paragraph("")

    meta = [
        ("Study Identifier:", csr.study_id),
        ("Sponsor:", SPONSOR),
        ("Investigational Product:", DRUG),
        ("Phase:", csr.phase),
        ("IND Number:", IND_NUMBER),
        ("Dossier:", DOSSIER_ID),
        ("Document Version:", "1.0 (Final)"),
        ("Date:", "15 March 2026"),
    ]
    for label, value in meta:
        p = docx.add_paragraph()
        run_l = p.add_run(f"{label} ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_v = p.add_run(value)
        run_v.font.size = Pt(10)

    docx.add_page_break()

    # ── Synopsis ──────────────────────────────────────────────────────────────
    docx.add_heading("SYNOPSIS", level=1)
    docx.add_paragraph(csr.synopsis)
    docx.add_page_break()

    # ── Table of Contents placeholder ─────────────────────────────────────────
    docx.add_heading("TABLE OF CONTENTS", level=1)
    toc_items = [s.heading for s in csr.sections]
    for item in toc_items:
        p = docx.add_paragraph(item, style="List Number")
        p.runs[0].font.size = Pt(10)
    docx.add_page_break()

    # ── Body sections ─────────────────────────────────────────────────────────
    for section in csr.sections:
        docx.add_heading(section.heading, level=section.level)
        for para_text in section.paragraphs:
            p = docx.add_paragraph(para_text)
            p.runs[0].font.size = Pt(10)
            docx.add_paragraph("")

        if section.table_rows:
            _add_table(docx, section)

    docx.save(str(out_path))


# ─────────────────────────────────────────────────────────────────────────────
# eCTD backbone index.xml
# ─────────────────────────────────────────────────────────────────────────────

def _write_index_xml(out_dir: Path, doc_paths: list[Path]) -> Path:
    nsmap = {None: "urn:hl7-org:v3", "xlink": "http://www.w3.org/1999/xlink"}
    root = etree.Element("ectd", nsmap=nsmap, attrib={"dtd-version": "3.2"})
    submission = etree.SubElement(
        root, "fda-regional",
        attrib={"submission-type": "nda", "submission-id": DOSSIER_ID},
    )
    for path in doc_paths:
        rel = path.relative_to(out_dir).as_posix()
        leaf_id = "leaf-" + path.stem
        leaf = etree.SubElement(
            submission, "leaf",
            attrib={
                "ID": leaf_id,
                "operation": "new",
                "{http://www.w3.org/1999/xlink}href": rel,
            },
        )
        etree.SubElement(leaf, "title").text = path.stem
    out_path = out_dir / "index.xml"
    etree.ElementTree(root).write(
        str(out_path), pretty_print=True,
        xml_declaration=True, encoding="UTF-8", standalone=False,
    )
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# MANIFEST
# ─────────────────────────────────────────────────────────────────────────────

def _write_manifest(out_dir: Path, doc_paths: list[Path]) -> Path:
    lines = [
        f"Demo Dossier  {DOSSIER_ID}",
        "=" * 70,
        f"Sponsor:              {SPONSOR}",
        f"Drug:                 {DRUG}",
        f"Documents:            {len(doc_paths)}",
        f"Cross-doc references: 12 (3 per document)",
        "",
        "Files:",
    ]
    for p in doc_paths:
        lines.append(f"  {p.relative_to(out_dir).as_posix()}")
    lines += [
        "",
        "Cross-reference map (all hit Regex + NER detection layers):",
        "  SP-2026-001  ->  SP-2026-002 Section 2.5",
        "  SP-2026-001  ->  SP-2026-003 Appendix 16.1.1",
        "  SP-2026-001  ->  SP-2026-004 Table 14.2.1.1",
        "  SP-2026-002  ->  SP-2026-001 Section 5.3.1",
        "  SP-2026-002  ->  SP-2026-003 Listing 16.2.5",
        "  SP-2026-002  ->  SP-2026-004 Section 5.3.5",
        "  SP-2026-003  ->  SP-2026-001 Figure 11",
        "  SP-2026-003  ->  SP-2026-002 Table 14.2.1.1",
        "  SP-2026-003  ->  SP-2026-004 Appendix 16.2.6",
        "  SP-2026-004  ->  SP-2026-001 Section 5.3.1",
        "  SP-2026-004  ->  SP-2026-002 Section 2.5",
        "  SP-2026-004  ->  SP-2026-003 Table 14.3.1",
        "",
        "Upload all 4 .docx files together in the Streamlit UI to see",
        "cross-document hyperlinks injected and validated.",
    ]
    out_path = out_dir / "MANIFEST.txt"
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# Orchestration
# ─────────────────────────────────────────────────────────────────────────────

def generate(out_dir: Path) -> dict[str, Path]:
    module_dir = out_dir / "m5" / "53-clin-stud-rep"
    module_dir.mkdir(parents=True, exist_ok=True)

    doc_paths: list[Path] = []
    for csr in ALL_DOCS:
        target = module_dir / csr.filename
        _render_doc(csr, target)
        doc_paths.append(target)

    index_path = _write_index_xml(out_dir, doc_paths)
    manifest_path = _write_manifest(out_dir, doc_paths)

    return {
        "out_dir": out_dir,
        "index": index_path,
        "manifest": manifest_path,
        **{csr.study_id: path for csr, path in zip(ALL_DOCS, doc_paths)},
    }


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate the 4-document Plan Two demo dossier."
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=Path("data/synthetic/demo_dossier"),
    )
    args = parser.parse_args()
    artifacts = generate(args.out)

    out = sys.stdout
    out.write("\n")
    out.write("=" * 70 + "\n")
    out.write(f"Demo dossier generated under: {artifacts['out_dir']}\n")
    out.write(f"  Documents  : {len(ALL_DOCS)}\n")
    out.write(f"  Cross-refs : 12 (3 per document, all cross-doc)\n")
    out.write(f"  Backbone   : {artifacts['index']}\n")
    out.write(f"  Manifest   : {artifacts['manifest']}\n")
    out.write("=" * 70 + "\n\n")
    out.write("Files to upload in the Streamlit UI:\n")
    for csr in ALL_DOCS:
        p = artifacts[csr.study_id]
        size_kb = p.stat().st_size // 1024
        out.write(f"  [{csr.phase:8s}]  {p.name}  ({size_kb} KB)\n")
    out.write("\n")


if __name__ == "__main__":
    main()
