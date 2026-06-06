"""Generate a synthetic CTD dossier for POC testing.

Produces:
    data/synthetic/
        m2/
            2-5-clin-overview/2-5-clin-overview.docx
            2-7-clin-summary/2-7-1-summary-bio.docx
            ...
        m5/
            5-3-1-bio-stud-rep/<study>.docx
            ...
        index.xml          # eCTD backbone v3.2-ish (mock)

Each generated document contains hundreds of intentionally placed reference
strings drawn from the pattern catalog so the detection engine has rich
material to chew on. The mix is tuned for ~25 references per document,
yielding ~500 across a 20-document synthetic Module 2.

Usage:
    python -m scripts.bootstrap_synthetic_data --out data/synthetic --docs 20
"""

from __future__ import annotations

import argparse
import random
import textwrap
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from lxml import etree


SPONSOR_PREFIXES = ["SP", "SUNP", "SPL", "ABC", "XYZ", "MED"]
SECTION_NUMBERS = [
    "2.5.1", "2.5.2", "2.5.3", "2.5.4", "2.5.5",
    "2.7.1", "2.7.2", "2.7.3", "2.7.4",
    "5.3.1", "5.3.2", "5.3.3", "5.3.4", "5.3.5",
    "3.2.P.5", "3.2.S.4",
]
TABLE_NUMBERS = ["14.2.1.1", "14.2.2.1", "11.4", "14.1.1", "5.3-1", "2.7.3-1", "16.1.1"]
FIGURE_NUMBERS = ["11", "14.2.1", "5", "2.7-1", "5.3-2"]
LISTING_NUMBERS = ["16.2.5", "16.1.1", "16.2.7", "16.2.5.1", "14-1"]
APPENDIX_NUMBERS = ["16.1.1", "16.2.5", "16.1.4", "16.2"]
MODULE_NUMBERS = ["5.3.1", "5.3.5", "2.5.3", "2.7.4", "3.2.P.5", "1.3.1"]

# ── Extended patterns for NER / Ollama trigger documents ────────────────────
AMBIGUOUS_NUMBERS = ["14", "42", "25", "11", "2", "3", "5", "7", "9", "15"]
IMPLICIT_SECTIONS = ["(2.5.3)", "(3.2.1)", "(14.2)", "section 2-5", "sec. 3.2"]
IMPLICIT_TABLES = ["chart 11", "table in 14.2", "displayed in 14", "shown in the table"]
IMPLICIT_FIGURES = ["figure 11", "chart 5.3", "plot 2.7"]
AMBIGUOUS_REFS = ["the previous report", "the study mentioned above", "CSR section", "the results shown here"]
CONTEXTUAL_NUMBERS = ["Patient 25", "Subject 42", "Center 14"]
IMPLICIT_MODULES = ["M2 ref", "Mod 2", "Module-5", "from Module 5"]
ABBREVIATIONS = ["CSR", "ISS", "ISE", "PK", "BA", "BE", "BCS"]


@dataclass
class StudyMeta:
    study_id: str
    nct_id: str
    sponsor: str


def _random_study() -> StudyMeta:
    sponsor = random.choice(SPONSOR_PREFIXES)
    year = random.randint(2020, 2025)
    seq = random.randint(1, 999)
    nct = f"NCT{random.randint(10000000, 99999999):08d}"
    return StudyMeta(
        study_id=f"{sponsor}-{year}-{seq:03d}",
        nct_id=nct,
        sponsor=sponsor,
    )


def _sentence(study: StudyMeta, sentence_type: str = "standard") -> str:
    """Build a single sentence containing 1–3 references.

    Args:
        study: StudyMeta with study_id, nct_id, sponsor
        sentence_type: one of "standard", "ambiguous", "contextual"
    """
    if sentence_type == "ambiguous":
        templates = [
            "In study {ambig_num}, the results were significant.",
            "As mentioned in {ambig_num}, the dose was escalated.",
            "Patient {ambig_num} experienced adverse events.",
            "{ambig_num} subjects completed the trial.",
            "The CSR section discusses protocol {ambig_num} amendments.",
            "Refer to {ambig_num} for details on the primary endpoint.",
            "In {ambig_num}, both treatment arms showed improvements.",
            "The results from {ambig_num} are tabulated in the appendix.",
            "Section {ambig_num} and {ambig_num2} together cover efficacy.",
            "As shown in table {ambig_num}, dosing was weight-based.",
        ]
        return random.choice(templates).format(
            ambig_num=random.choice(AMBIGUOUS_NUMBERS),
            ambig_num2=random.choice(AMBIGUOUS_NUMBERS),
        )
    elif sentence_type == "contextual":
        templates = [
            "Results in {implicit_section} (displayed below) were consistent.",
            "See {implicit_table} for the detailed pharmacokinetic profile.",
            "The {implicit_figure} demonstrates dose-response relationship.",
            "As per {ambig_ref}, we adopted the revised protocol.",
            "Patient {contextual_num} enrolled from {implicit_module}.",
            "{abbreviation} results are covered in {implicit_section}.",
            "The {abbreviation} section {implicit_section} details methodology.",
            "Cross-reference to {implicit_module} for the regulatory position.",
            "Data from {contextual_num} and {contextual_num2} support efficacy.",
            "Per {implicit_section}, the {abbreviation} was acceptable.",
        ]
        return random.choice(templates).format(
            implicit_section=random.choice(IMPLICIT_SECTIONS),
            implicit_table=random.choice(IMPLICIT_TABLES),
            implicit_figure=random.choice(IMPLICIT_FIGURES),
            ambig_ref=random.choice(AMBIGUOUS_REFS),
            contextual_num=random.choice(CONTEXTUAL_NUMBERS),
            contextual_num2=random.choice(CONTEXTUAL_NUMBERS),
            implicit_module=random.choice(IMPLICIT_MODULES),
            abbreviation=random.choice(ABBREVIATIONS),
        )
    else:  # standard
        templates = [
            "Demographic data are summarized in Table {table} of Section {section}.",
            "As described in Section {section}, study {study} demonstrated efficacy.",
            "Refer to Module {module} for the full CSR of {study} ({nct}).",
            "See Figure {figure} and Listing {listing} for adverse-event distributions.",
            "Per §{section}, the safety profile was acceptable; details in Appendix {appendix}.",
            "Pharmacokinetic results for {study} are tabulated in Table {table}.",
            "The {study} ({nct}) cohort is described in Section {section} and Module {module}.",
            "Listing {listing} (Appendix {appendix}) enumerates serious adverse events.",
            "Figure {figure} shows dose-response data referenced in §{section}.",
            "Cross-references to Module {module} and Section {section} are documented in Table {table}.",
        ]
        tmpl = random.choice(templates)
        return tmpl.format(
            section=random.choice(SECTION_NUMBERS),
            table=random.choice(TABLE_NUMBERS),
            figure=random.choice(FIGURE_NUMBERS),
            listing=random.choice(LISTING_NUMBERS),
            appendix=random.choice(APPENDIX_NUMBERS),
            module=random.choice(MODULE_NUMBERS),
            study=study.study_id,
            nct=study.nct_id,
        )


def _document_body(
    study: StudyMeta, sentences_per_doc: int, sentence_type: str = "standard"
) -> list[str]:
    """Generate document body with configurable sentence type.

    Args:
        study: StudyMeta with study_id, nct_id, sponsor
        sentences_per_doc: number of sentences to generate
        sentence_type: one of "standard", "ambiguous", "contextual"
    """
    paragraphs: list[str] = []
    paragraphs.append(f"Clinical Overview for {study.study_id} ({study.nct_id})")
    paragraphs.append("")
    paragraphs.append("Background")
    # Group sentences into paragraphs of 3–5 sentences each
    bucket: list[str] = []
    for _ in range(sentences_per_doc):
        bucket.append(_sentence(study, sentence_type=sentence_type))
        if len(bucket) >= random.randint(3, 5):
            paragraphs.append(" ".join(bucket))
            bucket = []
    if bucket:
        paragraphs.append(" ".join(bucket))
    return paragraphs


CTD_LAYOUT: list[tuple[str, str, str]] = [
    # ── Original 20 documents (standard references) ──────────────────────────
    ("m2/2-5-clin-overview/2-5-clin-overview.docx", "2.5 Clinical Overview", "standard"),
    ("m2/2-7-clin-summary/2-7-1-summary-bio.docx", "2.7.1 Summary of Biopharmaceutic Studies", "standard"),
    ("m2/2-7-clin-summary/2-7-2-summary-clin-pharm.docx", "2.7.2 Summary of Clinical Pharmacology", "standard"),
    ("m2/2-7-clin-summary/2-7-3-summary-clin-efficacy.docx", "2.7.3 Summary of Clinical Efficacy", "standard"),
    ("m2/2-7-clin-summary/2-7-4-summary-clin-safety.docx", "2.7.4 Summary of Clinical Safety", "standard"),
    ("m5/5-3-1-bio-stud-rep/study-001.docx", "5.3.1 BA/BE Study Report", "standard"),
    ("m5/5-3-1-bio-stud-rep/study-002.docx", "5.3.1 BA/BE Study Report", "standard"),
    ("m5/5-3-3-pk-stud-rep/study-pk-01.docx", "5.3.3 PK Study Report", "standard"),
    ("m5/5-3-5-efficacy-safety/study-eff-01.docx", "5.3.5 Efficacy/Safety Study Report", "standard"),
    ("m5/5-3-5-efficacy-safety/study-eff-02.docx", "5.3.5 Efficacy/Safety Study Report", "standard"),
    ("m3/3-2-quality/3-2-s-drug-substance.docx", "3.2.S Drug Substance", "standard"),
    ("m3/3-2-quality/3-2-p-drug-product.docx", "3.2.P Drug Product", "standard"),
    ("m4/4-2-stud-rep/nonclin-01.docx", "4.2 Non-Clinical Study Reports", "standard"),
    ("m4/4-2-stud-rep/nonclin-02.docx", "4.2 Non-Clinical Study Reports", "standard"),
    ("m1/us/1-3-1-letter.docx", "1.3.1 Cover Letter (US)", "standard"),
    ("m1/eu/1-3-1-letter.docx", "1.3.1 Cover Letter (EU)", "standard"),
    ("m2/2-3-quality-overall.docx", "2.3 Quality Overall Summary", "standard"),
    ("m2/2-4-nonclin-overview.docx", "2.4 Nonclinical Overview", "standard"),
    ("m2/2-6-nonclin-summary.docx", "2.6 Nonclinical Summary", "standard"),
    ("m5/5-3-7-case-report-forms.docx", "5.3.7 Case Report Forms", "standard"),
    # ── New 10 documents (5 ambiguous to trigger Ollama, 5 contextual to trigger NER) ──
    ("m2/2-5-ambiguous-refs-01.docx", "2.5 Ambiguous References (Test Set)", "ambiguous"),
    ("m2/2-7-ambiguous-refs-02.docx", "2.7 Ambiguous References CSR", "ambiguous"),
    ("m3/3-2-ambiguous-refs-03.docx", "3.2 Ambiguous Quality Refs", "ambiguous"),
    ("m4/4-2-ambiguous-refs-04.docx", "4.2 Ambiguous Nonclinical Refs", "ambiguous"),
    ("m5/5-3-ambiguous-refs-05.docx", "5.3 Ambiguous Study Refs", "ambiguous"),
    ("m1/1-3-contextual-ner-01.docx", "1.3 Contextual NER Regional Refs", "contextual"),
    ("m2/2-3-contextual-ner-02.docx", "2.3 Contextual NER Quality Refs", "contextual"),
    ("m2/2-6-contextual-ner-03.docx", "2.6 Contextual NER Study Refs", "contextual"),
    ("m5/5-3-1-contextual-ner-04.docx", "5.3.1 Contextual NER BA/BE Refs", "contextual"),
    ("m5/5-3-5-contextual-ner-05.docx", "5.3.5 Contextual NER Efficacy Refs", "contextual"),
]


def generate_documents(out_dir: Path, num_docs: int, sentences_per_doc: int) -> list[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    rng = random.Random(42)
    random.seed(42)  # deterministic for reproducibility

    selected = CTD_LAYOUT[:num_docs] if num_docs <= len(CTD_LAYOUT) else CTD_LAYOUT
    created: list[Path] = []

    for layout_entry in selected:
        # Handle both old (2-tuple) and new (3-tuple) layout entries
        if len(layout_entry) == 3:
            rel_path, title, sentence_type = layout_entry
        else:
            rel_path, title = layout_entry
            sentence_type = "standard"

        study = _random_study()
        target = out_dir / rel_path
        target.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(f"Document path: {rel_path}")
        doc.add_paragraph("")

        for para in _document_body(study, sentences_per_doc, sentence_type=sentence_type):
            if not para:
                doc.add_paragraph("")
            elif len(para) < 80 and not para.endswith("."):
                doc.add_heading(para, level=2)
            else:
                doc.add_paragraph(para)

        doc.save(str(target))
        created.append(target)
        _ = rng  # placeholder; reserved if we want per-doc randomization later
    return created


def generate_index_xml(out_dir: Path, docs: list[Path]) -> Path:
    """Build a minimal eCTD backbone index.xml referencing every generated leaf."""
    nsmap = {
        None: "urn:hl7-org:v3",
        "xlink": "http://www.w3.org/1999/xlink",
    }
    root = etree.Element(
        "ectd",
        nsmap=nsmap,
        attrib={"dtd-version": "3.2"},
    )
    submission = etree.SubElement(
        root,
        "fda-regional",
        attrib={"submission-type": "ind", "submission-id": "synthetic-0001"},
    )

    for path in docs:
        rel = path.relative_to(out_dir).as_posix()
        leaf = etree.SubElement(
            submission,
            "leaf",
            attrib={
                "ID": f"leaf-{abs(hash(rel)) % 10**8}",
                "operation": "new",
                "{http://www.w3.org/1999/xlink}href": rel,
            },
        )
        title = etree.SubElement(leaf, "title")
        title.text = path.stem

    index = out_dir / "index.xml"
    tree = etree.ElementTree(root)
    tree.write(
        str(index),
        pretty_print=True,
        xml_declaration=True,
        encoding="UTF-8",
        standalone=False,
    )
    return index


def generate_manifest(out_dir: Path, docs: list[Path]) -> Path:
    manifest = out_dir / "MANIFEST.txt"
    lines = [
        "Synthetic CTD dossier",
        "=====================",
        f"Generated by scripts/bootstrap_synthetic_data.py",
        f"Document count: {len(docs)}",
        "",
        "Files:",
    ]
    for path in docs:
        rel = path.relative_to(out_dir).as_posix()
        lines.append(f"  - {rel}")
    manifest.write_text("\n".join(lines), encoding="utf-8")
    return manifest


def main() -> None:
    parser = argparse.ArgumentParser(description=textwrap.dedent(__doc__ or ""))
    parser.add_argument(
        "--out",
        type=Path,
        default=Path("data/synthetic"),
        help="Output directory (default: data/synthetic)",
    )
    parser.add_argument(
        "--docs",
        type=int,
        default=30,
        help="Number of documents to generate (default: 30; max: 30 in catalog)",
    )
    parser.add_argument(
        "--sentences-per-doc",
        type=int,
        default=25,
        help="Sentences per document — drives reference density (default: 25)",
    )
    args = parser.parse_args()

    if args.docs > len(CTD_LAYOUT):
        print(f"Warning: {args.docs} docs requested, but only {len(CTD_LAYOUT)} in catalog. Capping to {len(CTD_LAYOUT)}.")
        args.docs = len(CTD_LAYOUT)

    docs = generate_documents(args.out, args.docs, args.sentences_per_doc)
    index_path = generate_index_xml(args.out, docs)
    manifest_path = generate_manifest(args.out, docs)

    total_refs_estimate = len(docs) * args.sentences_per_doc * 2  # ~2 refs/sentence
    ambiguous_count = sum(1 for rel_path, _, stype in CTD_LAYOUT[:args.docs] if stype == "ambiguous")
    contextual_count = sum(1 for rel_path, _, stype in CTD_LAYOUT[:args.docs] if stype == "contextual")
    standard_count = args.docs - ambiguous_count - contextual_count

    print(f"\n{'='*70}")
    print(f"Generated {len(docs)} documents under {args.out}")
    print(f"  - {standard_count} standard (regex-friendly)")
    print(f"  - {ambiguous_count} ambiguous (Ollama triggers)")
    print(f"  - {contextual_count} contextual (NER triggers)")
    print(f"Wrote eCTD backbone: {index_path}")
    print(f"Wrote manifest:      {manifest_path}")
    print(f"Estimated references embedded: ~{total_refs_estimate}")
    print(f"Expected Ollama calls: ~{ambiguous_count * 8} (8–15 per ambiguous doc)")
    print(f"Expected NER calls: ~{contextual_count * 13} (13–20 per contextual doc)")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()
