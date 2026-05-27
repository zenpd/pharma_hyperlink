"""Unit tests for validation/cross_module_integrity.py (W6.3)."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from hyperlink_engine.graph.backbone_graph import BackboneGraph
from hyperlink_engine.models import (
    BackboneLeaf,
    BackboneSnapshot,
    CrossModuleLink,
    DocumentProvenance,
    LeafOperation,
    LinkStatus,
)
from hyperlink_engine.validation.cross_module_integrity import (
    CircularReference,
    CrossModuleAudit,
    IntegrityIssue,
    detect_circular_refs,
    detect_orphans,
    run_full_audit,
    to_link_records,
    validate_cross_module_link,
    validate_cross_module_links,
)


def _provenance() -> DocumentProvenance:
    return DocumentProvenance(
        source_path=Path("index.xml"), sha256="0" * 64, file_size_bytes=10
    )


def _snap(*leaves: BackboneLeaf) -> BackboneSnapshot:
    return BackboneSnapshot(
        provenance=_provenance(),
        schema_version="v3.2",
        region="us",
        sequence_number="0001",
        leaves=list(leaves),
    )


def _leaf(leaf_id: str, relpath: str, module: str = "m2.5") -> BackboneLeaf:
    return BackboneLeaf(
        leaf_id=leaf_id,
        relative_path=Path(relpath),
        module=module,
        operation=LeafOperation.NEW,
    )


def _link(
    src_id: str,
    dst_id: str,
    *,
    src_module: str = "m2.5",
    dst_module: str = "m5.3.1",
    anchor: str | None = None,
    relative_uri: str = "../../m5/x.docx",
) -> CrossModuleLink:
    return CrossModuleLink(
        source_leaf_id=src_id,
        target_leaf_id=dst_id,
        source_module=src_module,
        target_module=dst_module,
        relative_uri=relative_uri,
        anchor=anchor,
        confidence=0.9,
    )


def _make_docx_with_bookmark(path: Path, bookmark_name: str) -> None:
    doc = Document()
    p = doc.add_paragraph("anchor here")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "1")
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "1")
    p._p.insert(0, start)
    p._p.append(end)
    doc.save(str(path))


# ── Per-link validation ────────────────────────────────────────────────


def test_validate_ok_when_target_exists(tmp_path: Path) -> None:
    src = _leaf("S", "m2/2-5/source.docx")
    dst = _leaf("T", "m5/5-3/target.docx", module="m5.3.1")
    snap = _snap(src, dst)
    (tmp_path / "m5" / "5-3").mkdir(parents=True)
    Document().save(str(tmp_path / "m5" / "5-3" / "target.docx"))
    report = validate_cross_module_link(_link("S", "T"), snap, base_dir=tmp_path)
    assert report.is_ok
    assert report.issue == IntegrityIssue.OK


def test_validate_target_leaf_missing_from_backbone(tmp_path: Path) -> None:
    snap = _snap(_leaf("S", "m2/x.docx"))
    report = validate_cross_module_link(_link("S", "MISSING"), snap, base_dir=tmp_path)
    assert report.issue == IntegrityIssue.TARGET_LEAF_MISSING_FROM_BACKBONE


def test_validate_source_leaf_missing_from_backbone(tmp_path: Path) -> None:
    snap = _snap(_leaf("T", "m5/y.docx"))
    report = validate_cross_module_link(_link("MISSING", "T"), snap, base_dir=tmp_path)
    assert report.issue == IntegrityIssue.SOURCE_LEAF_MISSING


def test_validate_target_file_missing(tmp_path: Path) -> None:
    src = _leaf("S", "m2/x.docx")
    dst = _leaf("T", "m5/missing.docx", module="m5.3.1")
    snap = _snap(src, dst)
    report = validate_cross_module_link(_link("S", "T"), snap, base_dir=tmp_path)
    assert report.issue == IntegrityIssue.TARGET_FILE_MISSING
    assert "missing.docx" in (report.detail or "")


def test_validate_self_reference_rejected(tmp_path: Path) -> None:
    snap = _snap(_leaf("S", "m2/x.docx"))
    report = validate_cross_module_link(_link("S", "S"), snap, base_dir=tmp_path)
    assert report.issue == IntegrityIssue.SELF_REFERENCE


def test_validate_anchor_must_exist(tmp_path: Path) -> None:
    src = _leaf("S", "m2/source.docx")
    dst = _leaf("T", "m5/target.docx", module="m5.3.1")
    snap = _snap(src, dst)
    (tmp_path / "m5").mkdir(parents=True)
    _make_docx_with_bookmark(tmp_path / "m5" / "target.docx", "real_anchor")
    # Valid anchor.
    ok = validate_cross_module_link(
        _link("S", "T", anchor="real_anchor"), snap, base_dir=tmp_path
    )
    assert ok.is_ok
    # Missing anchor.
    broken = validate_cross_module_link(
        _link("S", "T", anchor="missing_anchor"), snap, base_dir=tmp_path
    )
    assert broken.issue == IntegrityIssue.ANCHOR_NOT_FOUND


# ── Batch validation ───────────────────────────────────────────────────


def test_validate_batch_returns_one_report_per_input(tmp_path: Path) -> None:
    src = _leaf("S", "m2/x.docx")
    dst_ok = _leaf("T1", "m5/exists.docx", module="m5.3.1")
    dst_missing = _leaf("T2", "m5/ghost.docx", module="m5.3.1")
    snap = _snap(src, dst_ok, dst_missing)
    (tmp_path / "m5").mkdir(parents=True)
    Document().save(str(tmp_path / "m5" / "exists.docx"))
    links = [_link("S", "T1"), _link("S", "T2")]
    reports = validate_cross_module_links(links, snap, base_dir=tmp_path)
    assert len(reports) == 2
    assert reports[0].is_ok
    assert reports[1].issue == IntegrityIssue.TARGET_FILE_MISSING


# ── Circular reference detection ───────────────────────────────────────


def test_detect_no_cycles_when_clean() -> None:
    snap = _snap(_leaf("A", "m2/a.docx"), _leaf("B", "m5/b.docx", module="m5.3.1"))
    graph = BackboneGraph.from_snapshot(snap)
    graph.add_reference("A", "B")
    cycles = detect_circular_refs(graph)
    assert cycles == []


def test_detect_two_node_cycle() -> None:
    snap = _snap(_leaf("A", "m2/a.docx"), _leaf("B", "m5/b.docx", module="m5.3.1"))
    graph = BackboneGraph.from_snapshot(snap)
    graph.add_reference("A", "B")
    graph.add_reference("B", "A")
    cycles = detect_circular_refs(graph)
    assert len(cycles) == 1
    leaf_ids = set(cycles[0].leaf_path)
    assert leaf_ids == {"A", "B"}
    assert cycles[0].length >= 2


def test_detect_three_node_cycle() -> None:
    snap = _snap(
        _leaf("A", "m2/a.docx"),
        _leaf("B", "m5/b.docx", module="m5.3.1"),
        _leaf("C", "m2/c.docx", module="m2.7"),
    )
    graph = BackboneGraph.from_snapshot(snap)
    graph.add_reference("A", "B")
    graph.add_reference("B", "C")
    graph.add_reference("C", "A")
    cycles = detect_circular_refs(graph)
    assert len(cycles) == 1
    assert set(cycles[0].leaf_path) == {"A", "B", "C"}


def test_circular_reference_dataclass_length() -> None:
    cycle = CircularReference(leaf_path=["A", "B", "C", "A"])
    assert cycle.length == 4


# ── Orphan detection ───────────────────────────────────────────────────


def test_detect_orphans_returns_unreferenced_leaves() -> None:
    snap = _snap(
        _leaf("A", "m2/a.docx"),
        _leaf("B", "m5/b.docx", module="m5.3.1"),
        _leaf("ORPH", "m2/orphan.docx"),
    )
    graph = BackboneGraph.from_snapshot(snap)
    graph.add_reference("A", "B")
    assert detect_orphans(graph) == ["ORPH"]


def test_detect_orphans_can_ignore_module_prefixes() -> None:
    snap = _snap(
        _leaf("COVER", "m1/us/cover.docx", module="m1"),
        _leaf("A", "m2/a.docx"),
        _leaf("B", "m5/b.docx", module="m5.3.1"),
    )
    graph = BackboneGraph.from_snapshot(snap)
    graph.add_reference("A", "B")
    # m1 leaves are legitimately orphan (cover letter, application form).
    assert detect_orphans(graph, ignore_modules={"m1"}) == []


# ── Full audit ─────────────────────────────────────────────────────────


def test_full_audit_aggregates_links_cycles_and_orphans(tmp_path: Path) -> None:
    src = _leaf("S", "m2/x.docx")
    dst = _leaf("T", "m5/y.docx", module="m5.3.1")
    orphan = _leaf("ORPH", "m2/orph.docx")
    snap = _snap(src, dst, orphan)
    (tmp_path / "m5").mkdir(parents=True)
    Document().save(str(tmp_path / "m5" / "y.docx"))
    graph = BackboneGraph.from_snapshot(snap)
    graph.add_reference("S", "T")
    graph.add_reference("T", "S")  # cycle!

    audit = run_full_audit(
        links=[_link("S", "T")],
        snapshot=snap,
        base_dir=tmp_path,
        graph=graph,
    )
    assert isinstance(audit, CrossModuleAudit)
    assert audit.ok_count == 1
    assert audit.failure_count == 0
    assert audit.has_failures is True  # cycle present
    assert audit.orphan_leaf_ids == ["ORPH"]
    assert audit.cycles


def test_full_audit_no_graph_skips_cycle_and_orphan_checks(tmp_path: Path) -> None:
    src = _leaf("S", "m2/x.docx")
    dst = _leaf("T", "m5/y.docx", module="m5.3.1")
    snap = _snap(src, dst)
    (tmp_path / "m5").mkdir(parents=True)
    Document().save(str(tmp_path / "m5" / "y.docx"))
    audit = run_full_audit(
        links=[_link("S", "T")], snapshot=snap, base_dir=tmp_path
    )
    assert audit.ok_count == 1
    assert audit.cycles == []
    assert audit.orphan_leaf_ids == []
    assert audit.has_failures is False


# ── CSV bridge ─────────────────────────────────────────────────────────


def test_to_link_records_maps_issues_to_statuses(tmp_path: Path) -> None:
    src = _leaf("S", "m2/x.docx")
    dst = _leaf("T", "m5/missing.docx", module="m5.3.1")
    snap = _snap(src, dst)
    reports = validate_cross_module_links([_link("S", "T")], snap, base_dir=tmp_path)
    records = to_link_records(reports)
    assert len(records) == 1
    assert records[0].status == LinkStatus.BROKEN
    assert records[0].source_doc == "S"
    assert records[0].target_doc == "T"
    assert "cross-module" in records[0].link_location_descriptor


def test_to_link_records_marks_self_ref_as_suspicious(tmp_path: Path) -> None:
    snap = _snap(_leaf("S", "m2/x.docx"))
    reports = validate_cross_module_links([_link("S", "S")], snap, base_dir=tmp_path)
    records = to_link_records(reports)
    assert records[0].status == LinkStatus.SUSPICIOUS


def test_to_link_records_handles_ok(tmp_path: Path) -> None:
    src = _leaf("S", "m2/x.docx")
    dst = _leaf("T", "m5/y.docx", module="m5.3.1")
    snap = _snap(src, dst)
    (tmp_path / "m5").mkdir(parents=True)
    Document().save(str(tmp_path / "m5" / "y.docx"))
    reports = validate_cross_module_links([_link("S", "T")], snap, base_dir=tmp_path)
    records = to_link_records(reports)
    assert records[0].status == LinkStatus.OK
