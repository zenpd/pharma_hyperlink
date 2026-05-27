"""Unit tests for parsing/docx_parser.py."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import RGBColor

from hyperlink_engine.parsing.docx_parser import (
    candidate_blue_runs,
    parse_docx,
)


@pytest.fixture
def styled_docx(tmp_path: Path) -> Path:
    path = tmp_path / "styled.docx"
    doc = Document()
    doc.core_properties.title = "Clinical Overview"
    doc.core_properties.author = "Publishing Bot"

    p1 = doc.add_paragraph()
    p1.add_run("Plain text. ")
    bold_run = p1.add_run("Bold portion. ")
    bold_run.bold = True
    blue_run = p1.add_run("Blue text should look like a link.")
    blue_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # pure blue

    p2 = doc.add_paragraph("Reference to Section 2.5.3 here.")

    # Table with one cell of content
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"

    doc.save(str(path))
    return path


def test_parse_captures_metadata(styled_docx: Path) -> None:
    parsed = parse_docx(styled_docx)
    assert parsed.title == "Clinical Overview"
    assert parsed.author == "Publishing Bot"
    assert parsed.paragraph_count >= 2  # 2 body paragraphs + table cells


def test_parse_captures_run_styling(styled_docx: Path) -> None:
    parsed = parse_docx(styled_docx)
    first_para = parsed.paragraphs[0]
    styles = [r.style for r in first_para.runs]

    assert any(s.bold for s in styles), "expected a bold run"
    assert any(s.color_rgb == "0000FF" for s in styles), "expected a blue run"
    # Plain run should be neither bold nor blue
    assert any(
        not s.bold and (s.color_rgb is None or s.color_rgb != "0000FF") for s in styles
    )


def test_char_offsets_are_monotonic(styled_docx: Path) -> None:
    parsed = parse_docx(styled_docx)
    for para in parsed.paragraphs:
        prev_end = 0
        for run in para.runs:
            assert run.char_offset_in_paragraph >= prev_end
            prev_end = run.char_offset_in_paragraph + len(run.text)


def test_blue_run_flagged_as_anomaly_candidate(styled_docx: Path) -> None:
    parsed = parse_docx(styled_docx)
    candidates = candidate_blue_runs(parsed)
    assert candidates, "expected at least one blue-text-no-link candidate"
    # And the candidate is in the first paragraph
    p_idx, _r_idx, run = candidates[0]
    assert p_idx == 0
    assert run.style.color_rgb == "0000FF"
    assert not run.style.is_hyperlink


def test_table_paragraphs_marked_in_table(styled_docx: Path) -> None:
    parsed = parse_docx(styled_docx)
    table_paragraphs = [p for p in parsed.paragraphs if p.in_table]
    assert table_paragraphs, "expected table paragraphs to be tagged"
    for para in table_paragraphs:
        assert para.table_coords is not None
        assert len(para.table_coords) == 3


def test_existing_hyperlink_marks_run(tmp_path: Path) -> None:
    """A run wrapped in <w:hyperlink> is flagged with style.is_hyperlink=True."""
    from docx import Document

    from hyperlink_engine.injection.docx_linker import DocxLinker
    from hyperlink_engine.models import RunLocation

    src = tmp_path / "src.docx"
    out = tmp_path / "linked.docx"
    doc = Document()
    doc.add_paragraph("Click here for details.")
    doc.save(str(src))

    linker = DocxLinker(src, out)
    linker.add_external_link(
        RunLocation(paragraph_index=0, run_index=0, char_start=0, char_end=10),
        url="https://example.org",
    )
    linker.save()

    parsed = parse_docx(out)
    hyperlink_runs = [r for p in parsed.paragraphs for r in p.runs if r.style.is_hyperlink]
    assert hyperlink_runs, "expected at least one run to be flagged as a hyperlink"
    # And the existing-hyperlinks list captured the target
    assert any("example.org" in url for url in parsed.existing_hyperlinks)


def test_blueish_helper_thresholds() -> None:
    """The _is_blueish helper correctly classifies edge cases."""
    from hyperlink_engine.parsing.docx_parser import _is_blueish

    assert _is_blueish("0000FF", tolerance=40) is True
    assert _is_blueish("0000FE", tolerance=40) is True
    assert _is_blueish("000080", tolerance=40) is True  # navy
    assert _is_blueish("FFFFFF", tolerance=40) is False  # white
    assert _is_blueish("000000", tolerance=40) is False  # black
    assert _is_blueish("808080", tolerance=40) is False  # grey
    assert _is_blueish(None, tolerance=40) is False
    assert _is_blueish("XYZXYZ", tolerance=40) is False  # garbage hex
