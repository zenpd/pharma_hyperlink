"""Unit tests for injection/docx_linker.py.

Exercises hyperlink injection on a small programmatically-built .docx
fixture so tests run fast and do not require external sample files.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from hyperlink_engine.injection.docx_linker import DocxInjectionError, DocxLinker
from hyperlink_engine.models import RunLocation


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Build a tiny .docx with one paragraph containing one run."""
    path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("See Section 2.5.3 for the rationale.")
    doc.add_paragraph("Refer to Module 5.3.1 and Table 14.2.1 below.")
    doc.save(str(path))
    return path


def _load_hyperlinks(path: Path) -> list[dict[str, str | None]]:
    """Inspect the saved .docx and return one dict per hyperlink found."""
    doc = Document(str(path))
    rels = doc.part.rels
    found: list[dict[str, str | None]] = []
    for para in doc.paragraphs:
        for hyperlink in para._element.iter(qn("w:hyperlink")):  # type: ignore[attr-defined]
            r_id = hyperlink.get(qn("r:id"))
            anchor = hyperlink.get(qn("w:anchor"))
            text_parts = [t.text or "" for t in hyperlink.iter(qn("w:t"))]
            target = rels[r_id].target_ref if r_id and r_id in rels else None
            found.append(
                {
                    "text": "".join(text_parts),
                    "anchor": anchor,
                    "external_target": target,
                }
            )
    return found


# ── External URL injection ───────────────────────────────────────────────


def test_add_external_link_writes_w_hyperlink(sample_docx: Path, tmp_path: Path) -> None:
    output = tmp_path / "linked.docx"
    linker = DocxLinker(sample_docx, output)

    # Run 0 text is "See Section 2.5.3 for the rationale."
    # Link the substring "Section 2.5.3" at char [4,17)
    linker.add_external_link(
        location=RunLocation(paragraph_index=0, run_index=0, char_start=4, char_end=17),
        url="https://example.com/sec-2.5.3",
    )
    linker.save()

    links = _load_hyperlinks(output)
    assert any(
        link["text"] == "Section 2.5.3"
        and link["external_target"] == "https://example.com/sec-2.5.3"
        for link in links
    )


def test_external_link_preserves_surrounding_text(sample_docx: Path, tmp_path: Path) -> None:
    output = tmp_path / "linked.docx"
    linker = DocxLinker(sample_docx, output)
    linker.add_external_link(
        location=RunLocation(paragraph_index=0, run_index=0, char_start=4, char_end=17),
        url="https://example.com/sec-2.5.3",
    )
    linker.save()

    # Concatenated paragraph text must be unchanged
    doc = Document(str(output))
    assert doc.paragraphs[0].text == "See Section 2.5.3 for the rationale."


# ── Internal bookmark injection ──────────────────────────────────────────


def test_add_internal_link_uses_w_anchor(sample_docx: Path, tmp_path: Path) -> None:
    output = tmp_path / "linked.docx"
    linker = DocxLinker(sample_docx, output)
    # Para 1 text: "Refer to Module 5.3.1 and Table 14.2.1 below."
    # "Module 5.3.1" is at char [9, 21)
    linker.add_internal_link(
        location=RunLocation(paragraph_index=1, run_index=0, char_start=9, char_end=21),
        anchor="m5_3_1",
    )
    linker.save()

    links = _load_hyperlinks(output)
    assert any(link["text"] == "Module 5.3.1" and link["anchor"] == "m5_3_1" for link in links)


# ── Multiple injections in one paragraph ─────────────────────────────────


def test_multiple_links_same_paragraph(sample_docx: Path, tmp_path: Path) -> None:
    output = tmp_path / "linked.docx"
    linker = DocxLinker(sample_docx, output)
    # Para 1 has both "Module 5.3.1" and "Table 14.2.1"
    linker.add_internal_link(
        location=RunLocation(paragraph_index=1, run_index=0, char_start=9, char_end=21),
        anchor="m5_3_1",
    )
    linker.add_internal_link(
        location=RunLocation(paragraph_index=1, run_index=0, char_start=26, char_end=38),
        anchor="t14_2_1",
    )
    linker.save()

    links = _load_hyperlinks(output)
    anchors = {link["anchor"] for link in links if link["anchor"]}
    assert {"m5_3_1", "t14_2_1"} <= anchors


# ── Display-text override ────────────────────────────────────────────────


def test_display_text_override(sample_docx: Path, tmp_path: Path) -> None:
    output = tmp_path / "linked.docx"
    linker = DocxLinker(sample_docx, output)
    linker.add_external_link(
        location=RunLocation(paragraph_index=0, run_index=0, char_start=4, char_end=17),
        url="https://example.com/sec-2.5.3",
        display_text="custom anchor text",
    )
    linker.save()

    links = _load_hyperlinks(output)
    assert any(link["text"] == "custom anchor text" for link in links)
    # The paragraph reflows with replaced text
    doc = Document(str(output))
    assert "custom anchor text" in doc.paragraphs[0].text


# ── Error handling ───────────────────────────────────────────────────────


def test_invalid_paragraph_raises(sample_docx: Path, tmp_path: Path) -> None:
    output = tmp_path / "linked.docx"
    linker = DocxLinker(sample_docx, output)
    linker.add_external_link(
        location=RunLocation(paragraph_index=99, run_index=0, char_start=0, char_end=1),
        url="https://example.com",
    )
    with pytest.raises(DocxInjectionError, match="Paragraph index 99 out of range"):
        linker.save()


def test_invalid_char_span_raises(sample_docx: Path, tmp_path: Path) -> None:
    output = tmp_path / "linked.docx"
    linker = DocxLinker(sample_docx, output)
    linker.add_external_link(
        location=RunLocation(paragraph_index=0, run_index=0, char_start=0, char_end=999),
        url="https://example.com",
    )
    with pytest.raises(DocxInjectionError, match="exceeds run length"):
        linker.save()


def test_missing_source_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        DocxLinker(tmp_path / "nope.docx", tmp_path / "out.docx")


# ── Idempotence: input file not mutated ──────────────────────────────────


def test_source_file_unchanged(sample_docx: Path, tmp_path: Path) -> None:
    output = tmp_path / "linked.docx"
    original_bytes = sample_docx.read_bytes()
    linker = DocxLinker(sample_docx, output)
    linker.add_external_link(
        location=RunLocation(paragraph_index=0, run_index=0, char_start=4, char_end=17),
        url="https://example.com",
    )
    linker.save()
    assert sample_docx.read_bytes() == original_bytes


# ── Pending counter ──────────────────────────────────────────────────────


def test_pending_count(sample_docx: Path, tmp_path: Path) -> None:
    linker = DocxLinker(sample_docx, tmp_path / "out.docx")
    assert linker.pending_count == 0
    linker.add_external_link(
        location=RunLocation(paragraph_index=0, run_index=0, char_start=0, char_end=3),
        url="https://example.com",
    )
    assert linker.pending_count == 1
