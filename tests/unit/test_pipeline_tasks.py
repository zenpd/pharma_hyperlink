"""Unit tests for pipeline/tasks.py (W7.1)."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from hyperlink_engine.models import LinkKind, LinkStatus
from hyperlink_engine.pipeline.cache import ExtractorConfig
from hyperlink_engine.pipeline.celery_app import PIPELINE_STAGES, reset_app
from hyperlink_engine.pipeline.tasks import (
    DocPipelineResult,
    _resolve_target,
    detect_references,
    get_task,
    ingest_document,
    inject_links,
    process_document,
    register_celery_tasks,
    validate_links,
    write_per_doc_report,
    _registered_tasks,
)


# ── Test fixtures ───────────────────────────────────────────────────────


def _make_docx_with_refs(path: Path, text: str = "See Section 2.5.3 and Table 1") -> None:
    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))


@pytest.fixture(autouse=True)
def _reset_celery_state() -> None:
    reset_app()
    _registered_tasks.clear()
    yield
    reset_app()
    _registered_tasks.clear()


# ── Stage 1: ingest_document ────────────────────────────────────────────


def test_ingest_document_returns_sha256(tmp_path: Path) -> None:
    src = tmp_path / "x.docx"
    _make_docx_with_refs(src)
    record = ingest_document(str(src))
    assert record["source_path"] == str(src)
    assert len(record["sha256"]) == 64
    assert record["file_size_bytes"] > 0


def test_ingest_document_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        ingest_document(str(tmp_path / "ghost.docx"))


def test_ingest_document_is_deterministic(tmp_path: Path) -> None:
    src = tmp_path / "x.docx"
    _make_docx_with_refs(src, "same content")
    a = ingest_document(str(src))
    b = ingest_document(str(src))
    assert a["sha256"] == b["sha256"]


# ── Stage 2: detect_references ─────────────────────────────────────────


def test_detect_references_picks_up_section_refs(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx_with_refs(src, "See Section 2.5.3 for details")
    ingest = ingest_document(str(src))
    detection = detect_references(
        ingest, extractor_config=ExtractorConfig.regex_only().__dict__
    )
    labels = {d["label"] for d in detection["detections"]}
    assert "SECTION_REF" in labels


def test_detect_references_carries_ingest_record(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx_with_refs(src, "Section 1")
    ingest = ingest_document(str(src))
    detection = detect_references(ingest, extractor_config={})
    assert detection["ingest"] == ingest


def test_detect_references_skips_empty_runs(tmp_path: Path) -> None:
    src = tmp_path / "blank.docx"
    doc = Document()
    doc.add_paragraph("")
    doc.add_paragraph("Section 2.5.3 here")
    doc.save(str(src))
    ingest = ingest_document(str(src))
    detection = detect_references(ingest, extractor_config={})
    # First paragraph (empty) yields nothing; second contributes the ref.
    assert any(d["label"] == "SECTION_REF" for d in detection["detections"])


# ── _resolve_target ─────────────────────────────────────────────────────


def test_resolve_target_section_ref() -> None:
    det = {
        "label": "SECTION_REF",
        "groups": {"num": "2.5.3"},
        "text": "Section 2.5.3",
        "pattern_id": "SECTION_REF_LABELED_V1",
    }
    kind, target = _resolve_target(det)
    assert kind == LinkKind.INTERNAL_BOOKMARK
    assert target == "section_ref_2_5_3"


def test_resolve_target_study_id_nct_goes_external() -> None:
    det = {
        "label": "STUDY_ID",
        "groups": {},
        "text": "NCT01234567",
        "pattern_id": "STUDY_ID_NCT_V1",
    }
    kind, target = _resolve_target(det)
    assert kind == LinkKind.EXTERNAL_URL
    assert target.startswith("https://clinicaltrials.gov/")


def test_resolve_target_study_id_other_goes_internal() -> None:
    det = {
        "label": "STUDY_ID",
        "groups": {},
        "text": "SP-2024-001",
        "pattern_id": "STUDY_ID_SPONSOR_V1",
    }
    kind, target = _resolve_target(det)
    assert kind == LinkKind.INTERNAL_BOOKMARK
    assert target == "study_SP_2024_001"


def test_resolve_target_ctd_leaf_with_sub() -> None:
    det = {
        "label": "CTD_LEAF",
        "groups": {"mod": "5", "sub": "3.1"},
        "text": "Module 5.3.1",
        "pattern_id": "CTD_LEAF_MODULE_V1",
    }
    kind, target = _resolve_target(det)
    assert kind == LinkKind.INTERNAL_BOOKMARK
    assert target == "m5_3_1"


def test_resolve_target_unknown_label_uses_text() -> None:
    det = {
        "label": "WEIRD",
        "groups": {},
        "text": "Section X",
        "pattern_id": "X",
    }
    kind, target = _resolve_target(det)
    assert kind == LinkKind.INTERNAL_BOOKMARK
    assert target == "Section X"


# ── Stage 3: inject_links ──────────────────────────────────────────────


def test_inject_links_writes_output_docx(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx_with_refs(src, "See Section 2.5.3")
    ingest = ingest_document(str(src))
    detection = detect_references(ingest, extractor_config={})
    out = tmp_path / "doc.linked.docx"
    injection = inject_links(detection, output_path=str(out))
    assert out.exists()
    assert injection["output_path"] == str(out)
    assert injection["probes"]
    assert all("kind" in p for p in injection["probes"])


def test_inject_links_never_mutates_source(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx_with_refs(src, "Section 2.5.3 ref")
    orig_size = src.stat().st_size
    ingest = ingest_document(str(src))
    detection = detect_references(ingest, extractor_config={})
    inject_links(detection, output_path=str(tmp_path / "out.docx"))
    assert src.stat().st_size == orig_size  # source untouched


# ── Stage 4: validate_links ────────────────────────────────────────────


def test_validate_links_returns_records(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx_with_refs(src, "Section 2.5.3 ref")
    ingest = ingest_document(str(src))
    detection = detect_references(ingest, extractor_config={})
    injection = inject_links(detection, output_path=str(tmp_path / "out.docx"))
    validation = validate_links(injection)
    assert validation["link_records"]
    # Internal-bookmark links should resolve OK because inject_links
    # also queues the bookmarks.
    statuses = {r["status"] for r in validation["link_records"]}
    assert "ok" in statuses


# ── Stage 5: write_per_doc_report ──────────────────────────────────────


def test_per_doc_report_csv_round_trip(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx_with_refs(src, "Section 2.5.3 ref")
    ingest = ingest_document(str(src))
    detection = detect_references(ingest, extractor_config={})
    injection = inject_links(detection, output_path=str(tmp_path / "out.docx"))
    validation = validate_links(injection)
    report_path = tmp_path / "report.csv"
    write_per_doc_report(validation, output_path=str(report_path))
    assert report_path.exists()
    content = report_path.read_text(encoding="utf-8")
    assert "source_doc" in content  # CSV header


# ── End-to-end process_document ────────────────────────────────────────


def test_process_document_runs_all_five_stages(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx_with_refs(src, "See Section 2.5.3 and Table 1")
    out = tmp_path / "out.docx"
    report = tmp_path / "report.csv"
    result = process_document(
        src,
        output_path=out,
        report_path=report,
        extractor_config=ExtractorConfig.regex_only(),
    )
    assert isinstance(result, DocPipelineResult)
    assert out.exists()
    assert report.exists()
    assert result.detection_count > 0
    assert result.total_links == result.detection_count
    assert result.duration_seconds > 0


def test_process_document_doc_with_no_refs_returns_zero_counts(tmp_path: Path) -> None:
    src = tmp_path / "boring.docx"
    _make_docx_with_refs(src, "Nothing interesting here.")
    result = process_document(
        src,
        output_path=tmp_path / "out.docx",
        report_path=tmp_path / "rep.csv",
        extractor_config=ExtractorConfig.regex_only(),
    )
    assert result.detection_count == 0
    assert result.total_links == 0
    assert result.broken_count == 0


def test_process_document_idempotent(tmp_path: Path) -> None:
    src = tmp_path / "doc.docx"
    _make_docx_with_refs(src, "Section 2.5.3 reference")
    out1 = tmp_path / "out1.docx"
    out2 = tmp_path / "out2.docx"
    a = process_document(
        src,
        output_path=out1,
        report_path=tmp_path / "r1.csv",
        extractor_config=ExtractorConfig.regex_only(),
    )
    b = process_document(
        src,
        output_path=out2,
        report_path=tmp_path / "r2.csv",
        extractor_config=ExtractorConfig.regex_only(),
    )
    assert a.detection_count == b.detection_count
    assert a.total_links == b.total_links


# ── Celery task registration ───────────────────────────────────────────

# These tests require the optional ``celery`` package.
try:
    import celery as _celery_mod  # noqa: F401

    _has_celery = True
except ImportError:
    _has_celery = False

_skip_no_celery = pytest.mark.skipif(not _has_celery, reason="celery not installed")


@_skip_no_celery
def test_register_celery_tasks_creates_one_per_stage() -> None:
    tasks = register_celery_tasks()
    # One task per stage.
    for stage in PIPELINE_STAGES:
        matches = [name for name in tasks if f".{stage}." in name]
        assert matches, f"no task registered for stage {stage}"


@_skip_no_celery
def test_register_celery_tasks_is_idempotent() -> None:
    a = register_celery_tasks()
    b = register_celery_tasks()
    assert a.keys() == b.keys()
    for key in a:
        assert a[key] is b[key]


@_skip_no_celery
def test_get_task_returns_callable() -> None:
    register_celery_tasks()
    task = get_task("ingestion", "ingest_document")
    assert callable(task)


@_skip_no_celery
def test_get_task_unknown_raises() -> None:
    register_celery_tasks()
    with pytest.raises(KeyError):
        get_task("ingestion", "no_such_action")


@_skip_no_celery
def test_celery_task_runs_eagerly(tmp_path: Path) -> None:
    """In eager mode, .delay() should execute synchronously and return a result."""
    from hyperlink_engine.pipeline.celery_app import make_celery_app
    make_celery_app(eager=True)
    register_celery_tasks()
    src = tmp_path / "x.docx"
    _make_docx_with_refs(src, "Section 2.5.3")
    task = get_task("ingestion", "ingest_document")
    result = task.delay(str(src))
    record = result.get(timeout=5)  # eager → already done
    assert record["sha256"]
    assert record["source_path"] == str(src)

